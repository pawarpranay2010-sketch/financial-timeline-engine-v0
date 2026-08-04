# App.py
"""
Financial Timeline Engine -- Institutional Financial Intelligence Platform

Section map (Phase 12 -- modular organization within a single file):
  1. Config & Session State
  2. Parsing            (file ingestion: PDF/DOCX/XLSX/CSV/TXT)
  3. AI Providers        (Google AI Studio -> Groq -> OpenRouter fallback chain)
  4. Document Processing  (chunking, summarization, hierarchical merge, caching)
  5. Timeline             (event extraction)
  6. Analytics            (Phase 3/5 institutional intelligence modules)
  7. Export               (DOCX / PDF / JSON / CSV / Excel / Markdown)
  8. Dashboard             (visualization)
  9. Future Modules        (Phase 11 architecture scaffold)
 10. Live Market Intelligence (Phase 6 architecture scaffold)
 11. Main App / UI
 12. Auth
"""
import streamlit as st
import requests
import io
import html
import json
import os
import re
import urllib.parse
from datetime import datetime, timezone
import pandas as pd
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from xml.sax.saxutils import escape as xml_escape
from google import genai
from google.genai import types

from core import (
    SETTINGS,
    FUTURE_MODULES,
    GROUNDING_RULE,
    DEFAULT_SESSION_STATE,
    get_secret,
    get_provider_health,
    provider_logger,
    hash_text,
    CacheManager,
    retry,
    extract_json,
    contains_error_marker,
)

# AI Executive Gateway -- replaces the hard-coded provider chain with
# workload-aware routing, deterministic failover, and Redis quota tracking.
# Initialised lazily so the app still starts if the gateway module has an
# import-time issue (unlikely, but keeps startup robust).
_ai_executive = None
from ingestion import (
    extract_multiple,
    merge_document_text,
    chunk_text,
    needs_chunking,
    document_statistics,
    print_statistics,
)
# TODO: verify this matches your actual backend/module3_controller.py --
# assumed path per your latest instructions (flat backend/ package, no
# module3/ subfolder) and that it exports a top-level `run_module3`.
from backend.module3_controller import run_module3

# =============================================================================
# SECTION 1: Config & Session State
# =============================================================================
# All tunable constants (model IDs, timeouts, retry policy, chunking sizes,
# the future-module roadmap) now live in core/config.py as `SETTINGS`, a
# single typed EngineSettings instance -- see that module's docstring for
# why. The names below are kept as module-level aliases so every existing
# reference further down in this file (PRIMARY_MODEL, CHUNK_SIZE, etc.)
# continues to work completely unchanged.
st.set_page_config(page_title="Financial Timeline Engine", layout="wide")

PRIMARY_MODEL = SETTINGS.primary_model
FALLBACK_MODEL = SETTINGS.fallback_model
OPENROUTER_MODELS = list(SETTINGS.openrouter_models)
GROQ_MODELS = list(SETTINGS.groq_models)

GROQ_TIMEOUT_SECONDS = SETTINGS.groq_timeout_seconds
OPENROUTER_TIMEOUT_SECONDS = SETTINGS.openrouter_timeout_seconds
PROVIDER_RETRY_ATTEMPTS = SETTINGS.provider_retry_attempts
PROVIDER_RETRY_DELAY_SECONDS = SETTINGS.provider_retry_delay_seconds

CHUNK_SIZE = SETTINGS.chunk_size
CHUNK_OVERLAP = SETTINGS.chunk_overlap
MERGE_BATCH_SIZE = SETTINGS.merge_batch_size

LIVE_INTELLIGENCE_API_KEY_NAME = SETTINGS.live_intelligence_api_key_name

# --- Session state init ---
# Default shape now lives in core/constants.py (DEFAULT_SESSION_STATE) so
# it can be reused by future non-Streamlit entry points (tests, backend
# warm-up) without importing this UI module.
for _key, _default in DEFAULT_SESSION_STATE.items():
    if _key not in st.session_state:
        st.session_state[_key] = _default


def _hash_text(text):
    """Thin backward-compatible alias for core.utilities.hash_text, kept
    so every call site below (`_hash_text(...)`) needs no further edits."""
    return hash_text(text)


def _cached_call(cache_name, cache_key, compute_fn):
    """Thin backward-compatible alias over core.utilities.CacheManager,
    preserving the exact (cache_name, cache_key, compute_fn) call
    signature used throughout this file. Builds a fresh CacheManager over
    the relevant st.session_state[cache_name] dict on each call -- cheap,
    since CacheManager itself holds no state beyond the mapping reference."""
    cache_store = st.session_state.setdefault(cache_name, {})
    return CacheManager(cache_store).get_or_compute(cache_key, compute_fn)


def _log_provider_event(stage, provider, status, detail=""):
    """Thin backward-compatible alias over the shared
    core.logging.provider_logger instance."""
    return provider_logger.log(stage, provider, status, detail)


def _retry(fn, attempts=PROVIDER_RETRY_ATTEMPTS, delay=PROVIDER_RETRY_DELAY_SECONDS):
    """Thin backward-compatible alias over core.utilities.retry."""
    return retry(fn, attempts=attempts, delay=delay)


def _debug_stage(label, text):
    """Pipeline debug instrumentation (added while diagnosing the
    "always insufficient information" + export traceback bugs): records
    the length, a first-1000-character preview, and whether an error
    marker is present for one intermediate pipeline output. Entries
    accumulate in st.session_state['pipeline_debug_log'] for the current
    "Generate Timeline Report" run (reset at the top of that button
    handler) and are rendered in a "Pipeline Debug Trace" expander so the
    exact stage where real content is lost -- or silently replaced by an
    error/placeholder -- is visible directly in the UI, without needing
    server log access."""
    text = text or ""
    st.session_state.setdefault("pipeline_debug_log", []).append({
        "stage": label,
        "length": len(text),
        "error_marker_detected": contains_error_marker(text),
        "preview": text[:1000],
    })

# =============================================================================
# SECTION 2: Parsing (file ingestion)
# =============================================================================
# All document parsing (PDF/DOCX/XLSX/CSV/TXT), chunking, caching, and
# per-file statistics now live in the ingestion/ package (parser.py,
# chunking.py, cache.py, statistics.py, extraction.py) and are imported
# above. There is exactly one document ingestion pipeline --
# ingestion.extract_multiple() / ingestion.merge_document_text() -- and no
# parsing logic is duplicated here. The legacy in-file extraction function
# that used to live in this section has been removed.


# =============================================================================
# SECTION 3: AI Providers (canonical gateway — backend/gateway/)
# =============================================================================
# backend/gateway/ is the SINGLE canonical AI provider/fallback system for
# the whole application (workload-aware routing, capability detection,
# provider health, deterministic failover, normalized responses, graceful
# no-provider behavior, env-vars -> Streamlit-secrets configuration).
#
# The four legacy functions below are retained ONLY as thin compatibility
# wrappers so existing callers/tests keep working; none of them implements
# its own provider chain anymore. Every real generation request flows
# through call_ai_with_fallback() -> AIExecutive.generate() ->
# ProviderManager/router/adapters. A missing optional provider (e.g.
# OpenRouter) is skipped by the gateway's failover and can NEVER surface as
# a raw provider-specific error string.

NO_ELIGIBLE_PROVIDER_MESSAGE = (
    "No eligible AI provider is configured. Add an AI provider key to continue."
)


def call_google_ai_studio(prompt_text, system_prompt=None, temperature=None):
    """Thin compatibility wrapper — generation now flows through the
    canonical AI gateway (backend/gateway). Kept only so existing
    callers/tests keep working; checks the Google key, then delegates
    to `call_ai_with_fallback` (the app's single provider chain)."""
    api_key = get_secret("GOOGLE_API_KEY", "")
    if not api_key:
        raise ValueError("Missing Google Key")
    return call_ai_with_fallback(prompt_text, system_prompt=system_prompt, temperature=temperature)


def call_groq_engine(prompt_text, system_prompt=None, temperature=None):
    """Thin compatibility wrapper — generation now flows through the
    canonical AI gateway (backend/gateway). Kept only so existing
    callers/tests keep working; checks the Groq key, then delegates
    to `call_ai_with_fallback` (the app's single provider chain)."""
    api_key = get_secret("GROQ_API_KEY", "")
    if not api_key:
        raise ValueError("Missing Groq Key")
    return call_ai_with_fallback(prompt_text, system_prompt=system_prompt, temperature=temperature)


def _openrouter_request(prompt_text, model_id, system_prompt=None, temperature=None):
    """Thin compatibility wrapper — generation now flows through the
    canonical AI gateway (backend/gateway), which skips unavailable
    providers deterministically. Returns (success: bool, content_or_error).
    A missing OpenRouter key yields the shared graceful no-provider
    message — never a raw provider-specific error string."""
    api_key = get_secret("OPENROUTER_API_KEY", "")
    if not api_key:
        return False, NO_ELIGIBLE_PROVIDER_MESSAGE
    result = call_ai_with_fallback(prompt_text, system_prompt=system_prompt, temperature=temperature)
    if result.startswith(("❌", "🔴", "⚠️")):
        return False, result
    return True, result


def call_openrouter_engine(prompt_text, system_prompt=None, temperature=None):
    """Thin compatibility wrapper — generation now flows through the
    canonical AI gateway (backend/gateway). Kept only so existing
    callers/tests keep working."""
    if system_prompt is None:
        system_prompt = (
            "You are an elite Wall Street financial research analyst. Generate "
            "structured multi-section corporate reports with key dates, events, "
            "and milestones."
        )
    _, result = _openrouter_request(prompt_text, PRIMARY_MODEL, system_prompt=system_prompt, temperature=temperature)
    return result


def _get_ai_executive():
    """Lazy singleton: initialises the AI Executive gateway on first use.
    Returns None if the module isn't available (graceful no-provider
    behavior downstream)."""
    global _ai_executive
    if _ai_executive is None:
        try:
            from backend.gateway import AIExecutive
            _ai_executive = AIExecutive()
        except Exception:
            _ai_executive = None
    return _ai_executive


def get_canonical_provider_status():
    """Per-provider status from the canonical ProviderManager
    (backend/gateway). Returns {slug: "available" |
    "configured_unavailable" | "not_configured"} for every registered
    provider, so the UI can distinguish 🟢 available / 🟡 configured but
    unavailable / ⚪ not configured instead of claiming live connectivity
    from a key alone."""
    executive = _get_ai_executive()
    if executive is None:
        return {}
    pm = executive.provider_manager
    status = {}
    for name in pm.DEFAULT_PRIORITY:
        keyed = pm.key_status().get(name, False)
        adapter = pm.get(name)
        if keyed and adapter is not None and adapter.health_check():
            status[name] = "available"
        elif keyed:
            status[name] = "configured_unavailable"
        else:
            status[name] = "not_configured"
    return status


def _detect_task_type(system_prompt, prompt_text):
    """Heuristic: guess the workload type from prompt content so the AI
    Executive router can select the best provider. Returns a task type
    string compatible with backend.gateway.router.* constants.

    Rules (checked in order):
    - system_prompt mentions JSON / array / object → "structured"
    - prompt asks for "investment memo" / "financial" → "financial"
    - prompt is very short (< 200 chars) → "simple"
    - otherwise → "financial" (safe default for this application)
    """
    combined = ((system_prompt or "") + " " + (prompt_text or "")).lower()
    json_keywords = ["json", "return only valid json", "return a json", "as json",
                     "return only a single valid json object", "expected_type"]
    if any(kw in combined for kw in json_keywords):
        return "structured"
    memo_keywords = ["investment memo", "investment research", "analyze the document summary",
                     "financial analysis", "financial statements", "analyze financial",
                     "institutional", "key financial events", "market movements",
                     "rag", "retrieved", "retrieved data", "retrieved documents",
                     "based on the retrieved", "based only on the following",
                     "source documents", "provided context"]
    if any(kw in combined for kw in memo_keywords):
        return "financial"
    if len(prompt_text or "") < 200:
        return "simple"
    return "financial"


def call_ai_with_fallback(prompt_text, system_prompt=None, temperature=None):
    """SINGLE canonical AI entry point for the entire application.

    Delegates exclusively to the AI Executive gateway (backend/gateway/):
        prompt -> AIExecutive.generate -> Router (workload-aware)
               -> ProviderManager + adapters (Google -> Groq -> OpenRouter
                  -> NVIDIA -> ... deterministic failover)
               -> normalized response -> content

    There is no second provider chain anywhere in app.py. Missing optional
    providers are skipped by the gateway's failover, so a missing
    OPENROUTER_API_KEY can never block an available Groq key and can never
    surface as a raw error string. If no eligible provider exists (or the
    gateway is unavailable), the shared graceful message is returned.
    """
    executive = _get_ai_executive()
    if executive is not None:
        try:
            task_type = _detect_task_type(system_prompt, prompt_text)
            response = executive.generate(
                prompt=prompt_text,
                system_prompt=system_prompt or "",
                temperature=temperature or 0.3,
                task_type=task_type,
            )
            if response.success:
                st.session_state["ai_connected"] = True
                st.session_state["ai_provider_used"] = f"{response.provider}/{response.model}"
                _log_provider_event("call_ai_with_fallback", response.provider, "success",
                                    f"task={task_type}, model={response.model}, latency={response.latency_ms}ms")
                return response.content
            # Gateway returned an error -- log and return the graceful
            # no-provider message (never a raw provider error string).
            _log_provider_event("call_ai_with_fallback", "gateway", "failed",
                                f"task={task_type}, error={(response.error or 'unknown')[:100]}")
        except Exception as e:
            _log_provider_event("call_ai_with_fallback", "gateway", "failed",
                                f"exception={type(e).__name__}: {str(e)[:100]}")

    # No eligible provider / gateway unavailable: graceful and machine-
    # detectable (⚠️ is in ERROR_RESPONSE_MARKERS) so downstream pipeline
    # stages degrade instead of treating this as generated content.
    st.session_state["ai_connected"] = False
    st.session_state["ai_provider_used"] = ""
    return f"⚠️ {NO_ELIGIBLE_PROVIDER_MESSAGE}"
  
# =============================================================================
# SECTION 4: Document Processing (chunking, summarization, hierarchical merge)
# =============================================================================
# GROUNDING_RULE now lives in core/constants.py and is imported above; kept
# as a plain reference here (no reassignment) so every use further down in
# this file is unchanged.


# =============================================================================
# SECTION 4: Document Processing (chunking, summarization, hierarchical merge)
# =============================================================================
# GROUNDING_RULE now lives in core/constants.py and is imported above; kept
# as a plain reference here (no reassignment) so every use further down in
# this file is unchanged.
#
# chunk_text() and needs_chunking() are no longer defined here -- they are
# imported from the ingestion package above (ingestion/chunking.py), which
# is now the single source of chunking logic for the whole application.
# This eliminates the duplicate chunking implementation that used to exist
# in both app.py and ingestion/chunking.py.


def summarize_single_document(document_text, file_name):
    """Summarizes one document/chunk into a concise institutional
    financial summary (500-1000 words), explicitly instructed to
    preserve numbers, tables, dates, timeline events, management
    commentary, and financial-statement figures rather than smoothing
    them into vague prose (Phase 2)."""
    if not document_text or not document_text.strip():
        return f"⚠️ No extractable text found in '{file_name}'. Nothing to summarize."

    system_prompt = (
        "You are an elite institutional financial analyst. Produce concise, "
        "precise, and professional document summaries suitable for a "
        "buy-side investment research desk. Focus on material facts: exact "
        "figures, tables (rendered as clear text), dates, timeline events, "
        "management commentary, and financial statement line items. Avoid "
        "filler language and avoid restating the obvious. " + GROUNDING_RULE
    )

    summarization_prompt = f"""Summarize the following document into a single, coherent institutional financial summary.

Requirements:
- Length: 500 to 1000 words.
- Tone: professional, analytical, institutional-grade.
- Preserve exact figures, tables, dates, timeline events, management commentary, and financial statement data verbatim -- do not round or paraphrase numbers.
- Do not use markdown headers or bullet lists; write in clear prose paragraphs (tables may be rendered as simple "Label: Value" lines within the prose).

Source Document Name: {file_name}

Source Document Text:
{document_text}

Return only the summary text, with no preamble or meta-commentary."""

    return call_ai_with_fallback(summarization_prompt, system_prompt=system_prompt, temperature=0.3)


def _merge_summary_batch(document_summaries):
    """Merges a single batch (<= MERGE_BATCH_SIZE) of summaries into one
    coherent master summary. Internal helper for merge_document_summaries;
    see that function for the hierarchical/recursive batching logic.

    ROOT-CAUSE FIX for "every intelligence module says insufficient
    information": an individual chunk/document summarization call can
    fail and return one of this app's own rendered error strings (e.g.
    "❌ OpenRouter Connection Failed...", see call_ai_with_fallback /
    summarize_single_document) instead of raising. Previously that error
    string was indistinguishable from a real summary by the time it
    reached here, so it was fed straight into the merge prompt as if it
    were genuine document content. When enough (or all) inputs to a merge
    were error strings, the merge model had nothing real to work with and,
    per its own grounding instructions, correctly reported that no
    information was available -- and *that* message became the
    master_summary, so every downstream intelligence field also
    (correctly, given that garbage input) reported insufficient
    information. This was the exact stage where real extracted text was
    being replaced by a fallback message.

    Fix: error-marked inputs are filtered out before building the merge
    prompt (the prompt text itself is completely unchanged), and if
    literally everything in this batch failed, a clear error is returned
    immediately instead of asking the AI to "merge" garbage.
    """
    valid_summaries = [d for d in document_summaries if not contains_error_marker(d.get("summary", ""))]
    failed_summaries = [d for d in document_summaries if contains_error_marker(d.get("summary", ""))]

    for doc in failed_summaries:
        _log_provider_event(
            "_merge_summary_batch", "pipeline", "excluded",
            f"Excluded '{doc.get('file_name')}' from merge -- its summary was an "
            f"error message, not real content: {str(doc.get('summary', ''))[:200]}"
        )

    if not valid_summaries:
        error_result = (
            "❌ Unable to generate a summary: every source document/chunk in this "
            "batch failed AI summarization. Check the Provider Health & Activity "
            "Log above for the underlying provider errors, then try again."
        )
        _debug_stage("merge_batch_result (all inputs failed)", error_result)
        return error_result

    system_prompt = (
        "You are an elite institutional financial analyst. Merge multiple "
        "per-document summaries into a single, coherent master summary "
        "suitable for a buy-side investment research desk. Reconcile "
        "overlapping information rather than repeating it, but preserve "
        "every distinct financial metric, table, date, timeline event, "
        "management commentary, and strategic insight found across the "
        "source summaries. Do not drop material information for brevity. "
        + GROUNDING_RULE
    )

    combined_summaries_text = ""
    for doc in valid_summaries:
        file_name = doc.get("file_name", "Unknown Document")
        summary = doc.get("summary", "")
        combined_summaries_text += f"\n--- Summary of: {file_name} ---\n{summary}\n"

    merge_prompt = f"""Below are individual summaries of separate financial documents (or document sections). Merge them into one coherent institutional financial master summary.

Requirements:
- Reconcile and consolidate overlapping information; do not repeat the same fact twice.
- Preserve all distinct financial metrics, tables, dates/timelines, management commentary, risk factors, controversies, and strategic implications mentioned across the summaries.
- Organize the merged summary in clear prose paragraphs.
- Attribute conflicting figures or claims across documents where relevant, rather than silently picking one.
- Do not fabricate information not present in the source summaries.

Individual Summaries:
{combined_summaries_text}

Return only the merged master summary text, with no preamble or meta-commentary."""

    result = call_ai_with_fallback(merge_prompt, system_prompt=system_prompt, temperature=0.3)
    _debug_stage("merge_batch_result", result)
    return result


def merge_document_summaries(document_summaries):
    """Merges per-document (or per-chunk) summaries into one coherent
    institutional financial master summary.

    Phase 2/7 -- hierarchical/recursive merge: if there are more than
    MERGE_BATCH_SIZE summaries, they are merged in batches first, and the
    batch-level merged summaries are then recursively merged again. This
    lets the pipeline handle an arbitrarily large number of documents or
    chunks without ever sending an oversized single merge request to any
    provider (which is what caused Groq's earlier HTTP 413s)."""
    if not document_summaries or len(document_summaries) == 0:
        return "⚠️ No document summaries available to merge."

    if len(document_summaries) <= MERGE_BATCH_SIZE:
        return _merge_summary_batch(document_summaries)

    batch_merged = []
    for i in range(0, len(document_summaries), MERGE_BATCH_SIZE):
        batch = document_summaries[i:i + MERGE_BATCH_SIZE]
        batch_summary = _merge_summary_batch(batch)
        batch_merged.append({"file_name": f"Batch {i // MERGE_BATCH_SIZE + 1}", "summary": batch_summary})

    return merge_document_summaries(batch_merged)


def summarize_document_with_chunking(document_text, file_name):
    """For large documents, splits text into chunks before summarizing
    (avoiding oversized AI requests, e.g. Groq HTTP 413), then merges the
    chunk summaries into one document-level summary via the existing
    hierarchical merge_document_summaries(). Small documents skip
    chunking entirely and behave exactly as summarize_single_document.

    Instrumented with _debug_stage() at every intermediate output
    (extracted text length, each chunk summary, the merged result) so the
    exact stage where real content is lost is visible in the Pipeline
    Debug Trace expander."""
    _debug_stage(f"extracted_text :: {file_name}", document_text)

    if not document_text or not needs_chunking(document_text, threshold=CHUNK_SIZE):
        summary = summarize_single_document(document_text, file_name)
        _debug_stage(f"single_document_summary :: {file_name}", summary)
        return summary

    chunks = chunk_text(document_text, chunk_size=CHUNK_SIZE, overlap=CHUNK_OVERLAP)
    chunk_summaries = []
    for idx, chunk in enumerate(chunks, start=1):
        chunk_label = f"{file_name} (Part {idx}/{len(chunks)})"
        chunk_summary = summarize_single_document(chunk, chunk_label)
        _debug_stage(f"chunk_summary :: {chunk_label}", chunk_summary)
        chunk_summaries.append({"file_name": chunk_label, "summary": chunk_summary})

    merged = merge_document_summaries(chunk_summaries)
    _debug_stage(f"document_merged_summary :: {file_name}", merged)
    return merged


# =============================================================================
# SECTION 5: Timeline extraction
# =============================================================================
def _extract_json_from_ai_response(result, expected_type=dict):
    """Thin backward-compatible alias over core.validation.extract_json,
    which now owns the markdown-fence stripping, error-marker detection,
    and bracket-scan fallback logic (unchanged behavior, single
    implementation shared with any future caller outside this file)."""
    return extract_json(result, expected_type=expected_type)


def extract_timeline_events(ai_narrative):
    """Parses AI narrative to extract structured timeline events."""
    try:
        structuring_prompt = f"""Extract timeline events from this narrative and return as JSON array with objects containing: date (YYYY-MM-DD or YYYY-MM or YYYY), event (string), category (string), impact (string).

Narrative:
{ai_narrative}

Return ONLY valid JSON array, no markdown, no extra text."""

        result = call_ai_with_fallback(structuring_prompt, temperature=0.3)
        return _extract_json_from_ai_response(result, expected_type=list)
    except Exception:
        return []


# =============================================================================
# SECTION 6: Analytics -- Phase 3/5 Institutional Intelligence Modules
# =============================================================================
# All 17 sections below are produced by ONE consolidated AI call
# (run_universal_intelligence_extraction), not 17 separate calls -- this
# directly serves Phase 8 (reduce AI cost / avoid duplicate calls) while
# still covering the full requested breadth.
INTELLIGENCE_MODULES = {
    "executive_summary": {
        "title": "📝 Executive Summary",
        "instruction": "A concise institutional executive summary (250-400 words) covering the business, recent performance, and overall investment context.",
        "expected_type": str,
    },
    "key_metrics": {
        "title": "📌 Key Financial Metrics",
        "instruction": "The key financial metrics (Revenue, EBITDA, PAT, EPS, Debt, Cash, Margins, Growth, Capex, Free Cash Flow, ROE, ROCE, Market Cap, and other material KPIs) are already extracted and verified in the 'financial_data' section of the Source Summary below -- report them as given, with their stated source/confidence, rather than re-deriving them from prose.",
        "expected_type": dict,
    },
    "ratio_analysis": {
        "title": "📐 Ratio Analysis",
        "instruction": "The 'ratios' section of the Source Summary below already contains every ratio (Revenue Growth, PAT Growth, EBITDA Growth, ROE, ROCE, Debt/Equity, Current Ratio, Operating Margin, Net Margin, EPS, Free Cash Flow indicators) that could be verified or calculated deterministically, each tagged with its source (Document or Calculated) and formula where applicable. Report these ratios as given -- do not recalculate them yourself. For any ratio marked 'Cannot Calculate' or absent from that section, use the exact string 'Not computable from provided documents'.",
        "expected_type": dict,
    },
    "financial_performance": {
        "title": "📈 Financial Performance",
        "instruction": "A structured narrative (300-500 words) describing revenue, profitability, and cash flow trends explicitly stated in the source.",
        "expected_type": str,
    },
    "segment_analysis": {
        "title": "🧩 Segment Analysis",
        "instruction": "Each distinct business segment/division mentioned, as an array of objects with 'segment_name', 'performance_summary', 'contribution'.",
        "expected_type": list,
    },
    "sector_analysis": {
        "title": "🏭 Sector Analysis",
        "instruction": "An object with 'sector', 'industry', 'business_model', 'competitors' (list), 'market_position', 'industry_trends', 'peer_context'.",
        "expected_type": dict,
    },
    "competitor_analysis": {
        "title": "⚔️ Competitor Analysis",
        "instruction": "Named competitors/peers mentioned, as an array of objects with 'competitor_name', 'context'.",
        "expected_type": list,
    },
    "swot": {
        "title": "🧭 SWOT Analysis",
        "instruction": "An object with keys 'strengths', 'weaknesses', 'opportunities', 'threats', each a list of short strings grounded in the source.",
        "expected_type": dict,
    },
    "risk_analysis": {
        "title": "⚠️ Risk Analysis",
        "instruction": "A structured risk assessment as an array of objects with 'category' (Business, Financial, Operational, Governance, Regulatory, Macroeconomic, or Investment), 'risk', 'severity' (Low/Medium/High), 'probability' (Low/Medium/High), 'mitigation'.",
        "expected_type": list,
    },
    "controversy_analysis": {
        "title": "🚨 Controversy Analysis",
        "instruction": "Any controversies (litigation, fraud, governance issues, management changes, accounting issues, environmental, political, regulatory actions, negative news) as an array of objects with 'date', 'type', 'description', 'severity', 'source'.",
        "expected_type": list,
    },
    "governance_analysis": {
        "title": "🏛️ Governance Analysis",
        "instruction": "A narrative (150-300 words) on corporate governance details explicitly mentioned (board structure, management changes, related-party matters, audit/compliance notes).",
        "expected_type": str,
    },
    "esg_summary": {
        "title": "🌱 ESG Summary",
        "instruction": "A narrative (150-300 words) summarizing any Environmental, Social, and Governance information explicitly mentioned.",
        "expected_type": str,
    },
    "investment_thesis": {
        "title": "💡 Investment Thesis",
        "instruction": "A concise investment thesis (150-300 words) grounded strictly in the source.",
        "expected_type": str,
    },
    "bull_case": {
        "title": "🐂 Bull Case",
        "instruction": "The strongest bullish arguments supported by the source, as an array of short strings.",
        "expected_type": list,
    },
    "bear_case": {
        "title": "🐻 Bear Case",
        "instruction": "The strongest bearish/risk arguments supported by the source, as an array of short strings.",
        "expected_type": list,
    },
    "catalysts": {
        "title": "⚡ Catalysts",
        "instruction": "Specific upcoming events/dates/triggers mentioned that could move the outlook, as an array of objects with 'catalyst', 'expected_timing'.",
        "expected_type": list,
    },
    "action_points": {
        "title": "✅ Action Points",
        "instruction": "Concrete, specific action points/recommendations an analyst would take next, as an array of short strings.",
        "expected_type": list,
    },
}


def _type_hint(expected_type):
    if expected_type is dict:
        return "a JSON object"
    if expected_type is list:
        return "a JSON array"
    return "a plain text string (no markdown)"


def run_universal_intelligence_extraction(master_summary):
    """Phase 3/5/8: single consolidated AI call that extracts all 17
    institutional intelligence sections defined in INTELLIGENCE_MODULES
    at once. Returns a dict keyed by module key; any module missing from
    the AI's response falls back to an empty value of the expected type
    so downstream rendering never breaks on a missing key.

    BUG FIXES:
    - Now also short-circuits (returning empty_result) when master_summary
      itself is an error string (contains_error_marker(...)), not just
      when it's empty. Previously an error string that slipped through
      from an upstream summarization/merge failure was treated as real
      content and sent to the AI anyway -- with nothing genuine to work
      from, the AI would (correctly, per its own grounding instructions)
      report "insufficient information" for every single field. Combined
      with the _merge_summary_batch fix above, master_summary should now
      only ever be an error string if literally every source failed.
    - Each field's value is now checked against its declared
      `expected_type` before being kept, and a `list` value containing a
      mix of dict and non-dict items (which the model occasionally
      returns) is normalized down to just its dict items. An unfiltered
      mixed list is what caused `pd.DataFrame(...)` to raise
      "ValueError: dictionary update sequence element #0 has length 1; 2
      is required" downstream in rendering/export -- the traceback that
      appeared right after generation.
    """
    empty_result = {key: config["expected_type"]() for key, config in INTELLIGENCE_MODULES.items()}

    _debug_stage("intelligence_extraction_input", master_summary)

    if not master_summary or not master_summary.strip() or contains_error_marker(master_summary):
        return empty_result

    system_prompt = (
        "You are an elite institutional financial research analyst producing "
        "a complete, structured research package. " + GROUNDING_RULE
    )

    field_instructions = "\n".join(
        f'- "{key}" ({_type_hint(config["expected_type"])}): {config["instruction"]}'
        for key, config in INTELLIGENCE_MODULES.items()
    )

    extraction_prompt = f"""Analyze the Source Summary below and produce a complete institutional research package.

Return ONLY a single valid JSON object with exactly these top-level keys:

{field_instructions}

Source Summary:
{master_summary}

Return ONLY the JSON object. No markdown, no extra text, no preamble."""

    result = call_ai_with_fallback(extraction_prompt, system_prompt=system_prompt, temperature=0.2)
    parsed = _extract_json_from_ai_response(result, expected_type=dict)
    if not parsed:
        return empty_result

    final = dict(empty_result)
    for key, config in INTELLIGENCE_MODULES.items():
        if key not in parsed or not parsed[key]:
            continue
        value = parsed[key]
        expected_type = config["expected_type"]

        if expected_type is list and isinstance(value, list):
            dict_items = [item for item in value if isinstance(item, dict)]
            if dict_items and len(dict_items) != len(value):
                # Mixed list (some dict items, some not) -- keep only the
                # structured items so downstream pd.DataFrame(...) calls
                # in rendering/export never see an inconsistent shape.
                value = dict_items
            final[key] = value
        elif isinstance(value, expected_type):
            final[key] = value
        # else: the AI returned a type that doesn't match this field's
        # expected_type at all (e.g. a plain string where a list was
        # expected) -- leave this field at its empty default rather than
        # passing an unexpected shape downstream.

    return final


def render_intelligence_output(key, value):
    """Renders one intelligence module's output using the same visual
    language already used elsewhere in the app (st.subheader, st.dataframe,
    st.markdown) -- no new UI patterns introduced."""
    config = INTELLIGENCE_MODULES[key]
    if not value:
        return
    st.subheader(config["title"])

    if isinstance(value, dict):
        if key == "swot":
            swot_col1, swot_col2 = st.columns(2)
            with swot_col1:
                st.markdown("**Strengths**")
                for item in value.get("strengths") or []:
                    st.markdown(f"- {item}")
                st.markdown("**Opportunities**")
                for item in value.get("opportunities") or []:
                    st.markdown(f"- {item}")
            with swot_col2:
                st.markdown("**Weaknesses**")
                for item in value.get("weaknesses") or []:
                    st.markdown(f"- {item}")
                st.markdown("**Threats**")
                for item in value.get("threats") or []:
                    st.markdown(f"- {item}")
        elif key == "sector_analysis":
            for field_name, field_value in value.items():
                if field_value:
                    label = field_name.replace("_", " ").title()
                    if isinstance(field_value, list):
                        field_value = ", ".join(str(v) for v in field_value)
                    st.markdown(f"**{label}:** {field_value}")
        else:
            st.dataframe(
                pd.DataFrame(list(value.items()), columns=["Field", "Value"]),
                use_container_width=True,
                hide_index=True
            )
    elif isinstance(value, list):
        if value and all(isinstance(item, dict) for item in value):
            st.dataframe(pd.DataFrame(value), use_container_width=True, hide_index=True)
        else:
            for item in value:
                st.markdown(f"- {item}")
    else:
        st.markdown(str(value))


def _render_module3_value(value):
    """Generic, defensive renderer for one Module 3 output section
    (financial_data, ratios, ocr_verification, cross_document_verification,
    confidence, events, timeline, optimized_context). Module 3's internal
    schema for each section is intentionally treated as opaque here --
    this renders any dict as interactive JSON (st.json), any list of dicts
    as a table, any other list as bullets, and anything else as plain
    text, so the UI never breaks regardless of Module 3's exact internal
    shape, and can be refined later without another app.py change."""
    if not value:
        st.caption("No data available for this section.")
        return
    if isinstance(value, dict):
        st.json(value)
    elif isinstance(value, list):
        if value and all(isinstance(item, dict) for item in value):
            st.dataframe(pd.DataFrame(value), use_container_width=True, hide_index=True)
        else:
            for item in value:
                st.markdown(f"- {item}")
    else:
        st.markdown(str(value))
# =============================================================================
# SECTION 7: Export System (DOCX / PDF / JSON / CSV / Excel / Markdown)
# =============================================================================
def _flatten_intelligence_for_text(intelligence_outputs):
    """Shared helper: converts the intelligence_outputs dict into a flat
    list of (title, rendered_text_or_rows) for reuse across DOCX/PDF/MD
    export, avoiding duplicating this formatting logic three times."""
    sections = []
    for key, value in (intelligence_outputs or {}).items():
        if not value:
            continue
        title = INTELLIGENCE_MODULES.get(key, {}).get("title", key)
        sections.append((title, value))
    return sections


def generate_docx_download(text_content, timeline_data=None, intelligence_outputs=None):
    """Compiles the generated AI analysis report into a Word document
    download stream. intelligence_outputs is optional (default None) so
    any existing call site without it still behaves exactly as before."""
    doc = Document()

    doc.add_heading("Institutional Investment Research Memo", level=1)
    doc.add_paragraph("-" * 40)
    doc.add_heading("Executive Summary & Analysis", level=2)

    if text_content:
        clean_text_string = str(text_content)
        for line in clean_text_string.split('\n'):
            if line.strip():
                sanitized_line = "".join(c for c in line if c.isprintable() or c in ['\t', '\n'])
                sanitized_line = sanitized_line.replace('**', '').replace('__', '').replace('```', '')
                if sanitized_line.strip():
                    doc.add_paragraph(sanitized_line.strip())
    else:
        doc.add_paragraph("No report content generated.")

    if timeline_data and len(timeline_data) > 0:
        doc.add_heading("Extracted Timeline Events", level=2)
        for event in timeline_data:
            date_str = "".join(c for c in str(event.get("date", "N/A")) if c.isprintable())
            event_name = "".join(c for c in str(event.get("event", "N/A")) if c.isprintable())
            category = "".join(c for c in str(event.get("category", "N/A")) if c.isprintable())
            impact = "".join(c for c in str(event.get("impact", "N/A")) if c.isprintable())
            doc.add_paragraph(f"📅 {date_str}: {event_name}", style="List Bullet")
            doc.add_paragraph(f"Category: {category} | Impact: {impact}", style="List Bullet 2")

    for title, value in _flatten_intelligence_for_text(intelligence_outputs):
        doc.add_heading(title, level=2)
        if isinstance(value, str):
            doc.add_paragraph(value)
        elif isinstance(value, dict):
            for field_name, field_value in value.items():
                if isinstance(field_value, list):
                    field_value = ", ".join(str(v) for v in field_value)
                doc.add_paragraph(f"{field_name.replace('_', ' ').title()}: {field_value}", style="List Bullet")
        elif isinstance(value, list):
            for item in value:
                if isinstance(item, dict):
                    line = " | ".join(f"{k}: {v}" for k, v in item.items())
                else:
                    line = str(item)
                doc.add_paragraph(line, style="List Bullet")

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


def generate_pdf_download(text_content, timeline_data=None, intelligence_outputs=None):
    """Compiles the report into an institutional-styled PDF using
    reportlab. intelligence_outputs is optional (default None) for
    backward compatibility. All text is XML-escaped since reportlab's
    Paragraph uses an internal XML markup dialect, and emoji are avoided
    in this output (reportlab's built-in fonts lack emoji glyphs, unlike
    python-docx which substitutes fonts automatically)."""
    bio = io.BytesIO()
    doc = SimpleDocTemplate(bio, pagesize=letter, topMargin=0.75 * 72, bottomMargin=0.75 * 72)
    base_styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        "InstitutionalTitle", parent=base_styles["Title"],
        textColor=colors.HexColor("#1a2b4c"), fontSize=20, spaceAfter=4
    )
    meta_style = ParagraphStyle(
        "InstitutionalMeta", parent=base_styles["Normal"],
        textColor=colors.grey, fontSize=8, spaceAfter=14
    )
    heading_style = ParagraphStyle(
        "InstitutionalHeading", parent=base_styles["Heading2"],
        textColor=colors.HexColor("#1a2b4c"), spaceBefore=14, spaceAfter=6
    )
    normal_style = base_styles["Normal"]

    story = [
        Paragraph("Institutional Investment Research Memo", title_style),
        Paragraph(f"Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}", meta_style),
        Paragraph("Executive Summary &amp; Analysis", heading_style),
        Spacer(1, 6),
    ]

    if text_content:
        for line in str(text_content).split('\n'):
            if line.strip():
                sanitized_line = "".join(c for c in line if c.isprintable() or c in ['\t', '\n'])
                sanitized_line = sanitized_line.replace('**', '').replace('__', '').replace('```', '')
                if sanitized_line.strip():
                    story.append(Paragraph(xml_escape(sanitized_line.strip()), normal_style))
                    story.append(Spacer(1, 6))
    else:
        story.append(Paragraph("No report content generated.", normal_style))

    if timeline_data and len(timeline_data) > 0:
        story.append(Paragraph("Extracted Timeline Events", heading_style))
        for event in timeline_data:
            date_str = "".join(c for c in str(event.get("date", "N/A")) if c.isprintable())
            event_name = "".join(c for c in str(event.get("event", "N/A")) if c.isprintable())
            category = "".join(c for c in str(event.get("category", "N/A")) if c.isprintable())
            impact = "".join(c for c in str(event.get("impact", "N/A")) if c.isprintable())
            story.append(Paragraph(xml_escape(f"Date: {date_str} -- {event_name}"), normal_style))
            story.append(Paragraph(xml_escape(f"Category: {category} | Impact: {impact}"), normal_style))
            story.append(Spacer(1, 6))

    for title, value in _flatten_intelligence_for_text(intelligence_outputs):
        story.append(Paragraph(xml_escape(title), heading_style))
        if isinstance(value, str):
            story.append(Paragraph(xml_escape(value), normal_style))
        elif isinstance(value, dict):
            for field_name, field_value in value.items():
                if isinstance(field_value, list):
                    field_value = ", ".join(str(v) for v in field_value)
                line = f"{field_name.replace('_', ' ').title()}: {field_value}"
                story.append(Paragraph(xml_escape(line), normal_style))
        elif isinstance(value, list):
            for item in value:
                if isinstance(item, dict):
                    line = " | ".join(f"{k}: {v}" for k, v in item.items())
                else:
                    line = str(item)
                story.append(Paragraph(xml_escape(line), normal_style))
        story.append(Spacer(1, 6))

    doc.build(story)
    bio.seek(0)
    return bio


def generate_json_export(ai_narrative_result, timeline_events, intelligence_outputs):
    """Phase 4: JSON export -- full structured payload."""
    payload = {
        "investment_memo": ai_narrative_result,
        "timeline_events": timeline_events or [],
        "intelligence": intelligence_outputs or {},
    }
    return json.dumps(payload, indent=2, ensure_ascii=False).encode("utf-8")


def generate_markdown_export(ai_narrative_result, timeline_events, intelligence_outputs):
    """Phase 4: Markdown export."""
    md_lines = ["# Institutional Investment Research Memo", "", "## Investment Memo", "", ai_narrative_result or "", ""]
    if timeline_events:
        md_lines += ["## Timeline Events", ""]
        for event in timeline_events:
            md_lines.append(
                f"- **{event.get('date', 'N/A')}** -- {event.get('event', 'N/A')} "
                f"(_{event.get('category', 'N/A')}_, impact: {event.get('impact', 'N/A')})"
            )
        md_lines.append("")
    for title, value in _flatten_intelligence_for_text(intelligence_outputs):
        md_lines += [f"## {title}", ""]
        if isinstance(value, str):
            md_lines.append(value)
        elif isinstance(value, dict):
            for k, v in value.items():
                if isinstance(v, list):
                    v = ", ".join(str(x) for x in v)
                md_lines.append(f"- **{k}**: {v}")
        elif isinstance(value, list):
            for item in value:
                if isinstance(item, dict):
                    md_lines.append("- " + " | ".join(f"**{k}**: {v}" for k, v in item.items()))
                else:
                    md_lines.append(f"- {item}")
        md_lines.append("")
    return "\n".join(md_lines).encode("utf-8")


def generate_csv_export(timeline_events, intelligence_outputs):
    """Phase 4: CSV export -- stacks the timeline table and any
    list-of-dict intelligence outputs (e.g. risk_analysis) as separate
    labeled sections within one CSV file."""
    buffer = io.StringIO()
    if timeline_events:
        buffer.write("Timeline Events\n")
        pd.DataFrame(timeline_events).to_csv(buffer, index=False)
        buffer.write("\n")
    for key, value in (intelligence_outputs or {}).items():
        if isinstance(value, list) and value and all(isinstance(item, dict) for item in value):
            title = INTELLIGENCE_MODULES.get(key, {}).get("title", key)
            buffer.write(f"{title}\n")
            pd.DataFrame(value).to_csv(buffer, index=False)
            buffer.write("\n")
    return buffer.getvalue().encode("utf-8")


def generate_excel_export(timeline_events, intelligence_outputs, key_metrics=None):
    """Phase 4: Excel export -- one sheet per structured section.

    Bug fix: openpyxl raises `IndexError: At least one sheet must be
    visible` if the workbook ends up with zero worksheets written (e.g.
    when timeline_events, intelligence_outputs, and key_metrics are all
    empty). This guarantees at least one sheet always exists before the
    writer saves: if nothing else was written, a "Summary" sheet is added
    instead of leaving the workbook empty.
    """
    bio = io.BytesIO()
    with pd.ExcelWriter(bio, engine="openpyxl") as writer:
        wrote_any = False
        if timeline_events:
            pd.DataFrame(timeline_events).to_excel(writer, sheet_name="Timeline", index=False)
            wrote_any = True
        if key_metrics:
            pd.DataFrame(list(key_metrics.items()), columns=["Metric", "Value"]).to_excel(
                writer, sheet_name="Key Metrics", index=False
            )
            wrote_any = True
        for key, value in (intelligence_outputs or {}).items():
            if isinstance(value, list) and value and all(isinstance(item, dict) for item in value):
                raw_name = INTELLIGENCE_MODULES.get(key, {}).get("title", key)
                sheet_name = "".join(c for c in raw_name if c not in ['\\', '/', '*', '[', ']', ':', '?']).strip()
                sheet_name = (sheet_name or key)[:31]
                pd.DataFrame(value).to_excel(writer, sheet_name=sheet_name, index=False)
                wrote_any = True
        if not wrote_any:
            pd.DataFrame(["Financial Timeline Engine", "No exportable data available."]).to_excel(
                writer, sheet_name="Summary", index=False, header=False
            )
    bio.seek(0)
    return bio.getvalue()


# =============================================================================
# SECTION 8: Dashboard / Visualization
# =============================================================================
def render_timeline_visualization(timeline_data):
    """Renders a simplified timeline visualization."""
    if not timeline_data or len(timeline_data) == 0:
        st.info("No timeline events extracted yet.")
        return
    st.subheader("📊 Timeline Events")
    df_timeline = pd.DataFrame(timeline_data)
    st.dataframe(df_timeline, use_container_width=True, hide_index=True)


# =============================================================================
# SECTION 9: Future Modules (Phase 11 -- architecture only, not implemented)
# =============================================================================
def get_future_module_status():
    """Returns the roadmap of planned-but-not-yet-implemented modules
    (Valuation, DCF, Comparable Analysis, Forecasting, Earnings Model,
    Financial Modeling, Portfolio Analysis, Watchlist, Company Comparison,
    Stock Scoring, Screening Engine) for display in the UI. This is
    intentionally scaffolding only, per Phase 11's explicit scope."""
    return {key: {"name": name, "status": "planned"} for key, name in FUTURE_MODULES.items()}


# =============================================================================
# SECTION 10: Live Market Intelligence (Phase 6 -- architecture only)
# =============================================================================
def is_live_market_intelligence_enabled():
    return bool(get_secret(LIVE_INTELLIGENCE_API_KEY_NAME, ""))


def fetch_live_market_intelligence(query):
    """Phase 6 architecture stub: a real, working function that gracefully
    degrades when no live-data API key is configured, rather than
    crashing or fabricating results. No specific news/filings/market-data
    vendor is wired in yet (that requires choosing and paying for a
    provider) -- once one is chosen, its client call replaces the body of
    the `if enabled` branch below; the calling code in main() already
    handles both the enabled and disabled cases cleanly."""
    if not is_live_market_intelligence_enabled():
        return {
            "enabled": False,
            "data": None,
            "message": "Live market intelligence is not configured for this deployment.",
        }
    return {
        "enabled": True,
        "data": None,
        "message": "Live market intelligence key detected, but no data provider integration has been selected yet.",
    }
# =============================================================================
# SECTION 10.5: AI Financial Assistant (interactive chatbot)
# =============================================================================
# A conversational UI layer on top of the existing frozen intelligence
# stack (Agentic RAG + verified facts + CalculationSafetyGate + provider
# chain). It deliberately does NOT reimplement any intelligence -- it
# orchestrates backend.chat_assistant.FinancialChatAssistant, which in
# turn composes the existing components.

# Suggested prompts shown in the assistant's empty state.
CHAT_SUGGESTED_PROMPTS = [
    "What was the company's revenue?",
    "Compare revenue across fiscal years.",
    "Why did profitability change?",
    "What are the major financial risks?",
    "Summarize this company's financial position.",
]

# Sprint 3: Intelligence-page suggestion set. The classic dashboard keeps
# CHAT_SUGGESTED_PROMPTS above; these route into the same verified
# Co-Pilot pipeline. Small and focused, and answerable from evidence.
FTE_CHAT_SUGGESTED_PROMPTS = [
    "What are the major financial risks?",
    "Explain the selected metric.",
    "What should I investigate next?",
    "Summarize the strongest verified evidence.",
]


@st.cache_resource(show_spinner=False)
def _build_chat_assistant():
    """Cached singleton for the chat assistant (lazy import keeps app
    startup light and keeps the module unit-testable outside Streamlit)."""
    from backend.chat_assistant import build_chat_assistant
    return build_chat_assistant()


def _chat_documents(extraction_results):
    """Shape ingestion results into the {file_name, financial_facts}
    structure the assistant's document-Q&A path expects."""
    documents = []
    for doc in extraction_results or []:
        documents.append({
            "file_name": doc.get("file_name", "Unknown Document"),
            "financial_facts": doc.get("financial_facts", []) or [],
        })
    return documents


def _render_chat_provenance(metadata):
    """Render concise provenance under a fact-backed assistant answer.

    - BLOCKED / miss intents show the reason instead of a badge.
    - Verified answers get a "✓ Verified" badge plus source / period /
      source-tier line and an expandable Evidence section.
    - Calculation answers get an expandable formula/inputs section.
    Internal hashes / technical objects are never shown.
    """
    intent = metadata.get("intent", "")
    evidence = metadata.get("evidence") or []
    calculation = metadata.get("calculation")

    blocked_intents = (
        "blocked", "calculation_miss", "document_miss", "company_blocked",
        "company_empty", "company_error", "no_provider", "provider_error",
    )
    if intent in blocked_intents:
        st.caption("⚠️ **Analysis limited** — this could not be verified from the current evidence.")
        reason = metadata.get("blocked_reason")
        if reason:
            st.caption(f"Missing input: {reason}")
        return

    if evidence:
        top = evidence[0]
        source = top.get("source") or ""
        period = top.get("period") or ""
        tier = top.get("source_tier")
        document = top.get("document") or ""
        badge_parts = ["✅ **Verified**"]
        if source:
            badge_parts.append(f"Source: {source}")
        if period:
            badge_parts.append(f"Period: {period}")
        if tier is not None:
            badge_parts.append(f"Source tier: {tier}")
        if document:
            badge_parts.append(f"Document: {document}")
        st.caption(" · ".join(badge_parts))
        with st.expander("🔍 Evidence", expanded=False):
            evidence_rows = [
                {
                    "Metric": e.get("metric", ""),
                    "Value": e.get("display", e.get("value", "")),
                    "Period": e.get("period", ""),
                    "Scale": e.get("scale", ""),
                    "Source": e.get("source", ""),
                    "Tier": e.get("source_tier", ""),
                    "Document": e.get("document", ""),
                }
                for e in evidence
            ]
            st.dataframe(pd.DataFrame(evidence_rows), use_container_width=True, hide_index=True)

    if calculation:
        with st.expander("🧮 Calculation", expanded=False):
            st.markdown(f"**{calculation.get('name', 'Calculation')}**")
            if calculation.get("formula"):
                st.markdown(f"Formula: `{calculation['formula']}`")
            if calculation.get("value") is not None:
                st.markdown(f"Value: **{calculation['value']}**")
            if calculation.get("inputs"):
                st.markdown("Inputs: " + ", ".join(
                    f"{k}={v}" for k, v in calculation["inputs"].items()
                ))
            if calculation.get("periods"):
                st.markdown("Periods: " + ", ".join(calculation["periods"]))


def _submit_chat_question(question, extraction_results, provider_health):
    """Run one user turn through the assistant and persist messages +
    bounded conversation context in session state."""
    from backend.chat_assistant import ChatContext

    question = (question or "").strip()
    if not question:
        return

    ctx = ChatContext.from_state(st.session_state.get("chat_context_state"))
    assistant = _build_chat_assistant()
    documents = _chat_documents(extraction_results)

    ctx.add_user(question)
    st.session_state.setdefault("chat_messages", []).append(
        {"role": "user", "content": question}
    )

    with st.spinner("Consulting verified evidence…"):
        response = assistant.answer(
            question,
            context=ctx,
            documents=documents,
            provider_health=provider_health,
        )

    content = response.get("content", "")
    metadata = response.get("metadata", {}) or {}
    ctx.add_assistant(content, metadata)
    st.session_state["chat_context_state"] = ctx.to_state()
    st.session_state.setdefault("chat_messages", []).append(
        {"role": "assistant", "content": content, "metadata": metadata}
    )


def render_financial_assistant(extraction_results, provider_health):
    """AI Financial Assistant section -- st.chat_message / st.chat_input
    conversation with suggested prompts, persistent bounded history, and
    provenance rendering."""
    st.markdown("---")
    st.subheader("🤖 AI Financial Assistant")
    st.caption(
        "Ask questions about your financial documents, verified metrics, "
        "companies, and market data. Every figure is verified against "
        "evidence — the assistant never invents numbers."
    )

    # --- Session state ---
    if "chat_messages" not in st.session_state:
        st.session_state["chat_messages"] = []
    if "chat_context_state" not in st.session_state:
        st.session_state["chat_context_state"] = None

    messages = st.session_state["chat_messages"]

    # --- Empty state with suggested prompts ---
    if not messages:
        st.markdown("#### 💬 Start a conversation")
        cols = st.columns(2)
        for i, prompt in enumerate(CHAT_SUGGESTED_PROMPTS):
            with cols[i % 2]:
                if st.button(prompt, key=f"chat_suggest_{i}", use_container_width=True):
                    _submit_chat_question(prompt, extraction_results, provider_health)
                    st.rerun()
    else:
        # --- Conversation history ---
        for msg in messages:
            with st.chat_message(msg["role"]):
                st.markdown(msg["content"])
                if msg["role"] == "assistant" and msg.get("metadata"):
                    _render_chat_provenance(msg["metadata"])

    # --- Chat input ---
    user_input = st.chat_input(
        "Ask about your financial documents, verified metrics, or companies…"
    )
    if user_input:
        _submit_chat_question(user_input, extraction_results, provider_health)
        st.rerun()


# =============================================================================
# SECTION 10b: Institutional Terminal embed (Phase 0–1 frontend via iframe)
# =============================================================================
# The Phase 0–1 institutional terminal (frontend/index.html, styles.css,
# app.js) is a static SPA served by the FastAPI backend (api.main:app →
# StaticFiles at "/", data at /api/v1/*). It runs on its own origin, so
# Streamlit embeds it with an iframe. URL resolution: FTE_TERMINAL_URL
# secret/env override → http://localhost:5000/ (the dev launcher runs the
# FastAPI service on that port alongside Streamlit).
_TERMINAL_EMBED_HEIGHT = 780
_TERMINAL_PROBE_TTL_SECONDS = 30


def _terminal_base_url() -> str:
    """Resolve the terminal URL: FTE_TERMINAL_URL override, else local FastAPI."""
    override = str(get_secret("FTE_TERMINAL_URL", "") or "").strip().rstrip("/")
    return f"{override}/" if override else "http://localhost:5000/"


def _terminal_reachable(url: str, timeout: float = 2.5) -> bool:
    """Server-side reachability probe (short timeout; never blocks the app)."""
    try:
        resp = requests.get(url, timeout=timeout, headers={"User-Agent": "fte-embed-probe"})
        return resp.status_code < 500
    except requests.RequestException:
        return False


def render_terminal_embed() -> None:
    """Render the institutional terminal inside the Streamlit workspace.

    The financial grid, provenance tray and gateway chrome live in the
    FastAPI-served SPA; Streamlit hosts it in a fixed-height iframe so the
    existing ingestion → analysis workflow stays available below. If the
    terminal service is unreachable the section degrades to a link instead
    of a dead iframe.
    """
    st.subheader("🖥️ Institutional Terminal")
    url = _terminal_base_url()

    probe_key = "fte_terminal_probe"
    cached = st.session_state.get(probe_key)
    now = datetime.now(timezone.utc)
    if cached and (now - cached["at"]).total_seconds() < _TERMINAL_PROBE_TTL_SECONDS:
        reachable = cached["ok"]
    else:
        reachable = _terminal_reachable(url)
        st.session_state[probe_key] = {"ok": reachable, "at": now}

    if reachable:
        st.components.v1.iframe(src=url, height=_TERMINAL_EMBED_HEIGHT, scrolling=True)
        st.caption(f"Embedded terminal — {url}")
    else:
        st.info("The institutional terminal service isn't reachable from this instance right now.")
        st.link_button("Open terminal in new tab", url)

    with st.expander("Terminal connection", expanded=False):
        st.caption(f"Resolved URL: `{url}`")
        st.caption(
            "Set the `FTE_TERMINAL_URL` secret/env var to point at a publicly "
            "deployed FastAPI terminal (e.g. Render/Railway via the Procfile)."
        )


# =============================================================================
# SECTION 10c: Institutional Terminal UI (Streamlit-native, Phase 1.6)
# =============================================================================
# Single-screen terminal rendered directly with Streamlit primitives
# (columns / tabs / containers / dataframe selection / chat). ALL figures
# come from existing backend outputs only — module3 financial_data/ratios,
# intelligence outputs, provider health, and the chat assistant. Nothing is
# ever invented: unknown metrics show "—" (Unanalyzed) and provenance shows
# only fields the pipeline actually provides.
_TERMINAL_CSS = """
:root {
  --fte-bg: #0a0d13;
  --fte-panel: #0f131b;
  --fte-border: #222a38;
  --fte-text: #e6e9ef;
  --fte-muted: #8b93a3;
  --fte-ok: #2fbf71;
  --fte-derived: #d9a13b;
  --fte-conflict: #4f8ef7;
  --fte-blocked: #e05252;
  --fte-unanalyzed: #5a6270;
  --fade-quick: 100ms;
  --fade-provenance: 150ms;
  --hover: 180ms;
}
.stApp { background: var(--fte-bg); }
.fte-header { display: flex; align-items: center; justify-content: space-between; gap: 12px; }
.fte-pill {
  display: inline-flex; align-items: center; gap: 8px;
  padding: 4px 12px; border: 1px solid var(--fte-border); border-radius: 999px;
  background: var(--fte-panel); color: var(--fte-text); font-size: 13px;
  transition: border-color var(--hover) ease;
}
.fte-pill:hover { border-color: #33405a; }
.fte-dot { width: 8px; height: 8px; border-radius: 50%; display: inline-block; }
.fte-dot.ok { background: var(--fte-ok); box-shadow: 0 0 6px rgba(47,191,113,.5); }
.fte-dot.warn { background: var(--fte-derived); }
.fte-dot.off { background: var(--fte-unanalyzed); }
.fte-rail-title {
  font-size: 11px; letter-spacing: .12em; text-transform: uppercase;
  color: var(--fte-muted); margin: 14px 0 6px; font-weight: 600;
}
.fte-source { display: flex; align-items: center; justify-content: space-between; gap: 8px; font-size: 13px; padding: 3px 0; }
.fte-source .nm { color: var(--fte-text); overflow: hidden; text-overflow: ellipsis; white-space: nowrap; }
.fte-source .st { color: var(--fte-muted); flex-shrink: 0; font-size: 12px; }
.fte-count-row { display: flex; align-items: center; justify-content: space-between; font-size: 13px; padding: 3px 0; }
.fte-count-row .k { color: var(--fte-text); }
.fte-count-row .v { color: var(--fte-muted); }
.fte-tray {
  border: 1px solid var(--fte-border); border-radius: 10px;
  background: var(--fte-panel); padding: 14px 16px; margin-top: 10px;
  animation: fteFade var(--fade-provenance) ease-out;
}
.fte-tray-head { font-weight: 600; color: var(--fte-text); margin-bottom: 8px; font-size: 14px; }
.fte-metric-row { display: grid; grid-template-columns: 110px 1fr; gap: 8px; padding: 3px 0; font-size: 13px; }
.fte-metric-row .k { color: var(--fte-muted); }
.fte-metric-row .v { color: var(--fte-text); }
.fte-tray-empty { color: var(--fte-muted); font-size: 13px; }
.fte-cap { color: var(--fte-muted); font-size: 12px; padding: 2px 0; }
.fte-ok-c { color: var(--fte-ok); } .fte-derived-c { color: var(--fte-derived); }
.fte-conflict-c { color: var(--fte-conflict); } .fte-blocked-c { color: var(--fte-blocked); }
.fte-unanalyzed-c { color: var(--fte-unanalyzed); }
[data-testid="stDataFrame"] { border: 1px solid var(--fte-border); border-radius: 8px; }
.fte-limited {
  border: 1px solid rgba(224,82,82,.35); background: rgba(224,82,82,.08);
  color: var(--fte-text); border-radius: 8px; padding: 8px 12px; font-size: 13px;
  animation: fteFade var(--fade-quick) ease-out;
}
.fte-caps {
  color: var(--fte-muted); font-size: 12px; padding: 6px 0 2px;
  border-top: 1px solid var(--fte-border); margin-top: 4px;
}
.fte-pv-avail .v { color: var(--fte-text); }
.fte-pv-missing .v { color: var(--fte-muted); font-style: italic; opacity: .75; }
.fte-pv-missing .k { opacity: .75; }
.fte-block-note {
  margin-top: 8px; padding: 6px 10px; font-size: 12px; color: var(--fte-blocked);
  border-left: 2px solid var(--fte-blocked); background: rgba(224,82,82,.06);
  border-radius: 4px; animation: fteFade var(--fade-provenance) ease-out;
}
.fte-partial { color: var(--fte-derived); font-size: 12px; font-weight: 500; }
[data-testid="stButton"] button { font-size: 13px; min-height: 30px; }
[data-testid="stSidebar"] [data-testid="stButton"] button { font-size: 12px; min-height: 28px; padding: 2px 10px; }
@keyframes fteFade {
  from { opacity: 0; }
  to { opacity: 1; }
}
@media (prefers-reduced-motion: reduce) {
  * { animation: none !important; transition: none !important; }
}
"""

_FTE_CSS = """
/* ============ FT-E entrance / auth / selection / workspace chrome ============ */
.fte-spacer-lg { height: 12vh; }
.fte-spacer-md { height: 7vh; }
.fte-spacer-sm { height: 2.4rem; }
.fte-stage { text-align: center; }
.fte-brand {
  font-weight: 750; letter-spacing: .24em; text-indent: .24em;
  color: var(--fte-text); line-height: 1;
}
.fte-brand-xl { font-size: clamp(2.75rem, 7vw, 4.5rem); }
.fte-moat {
  margin-top: 1.1rem; font-size: .8rem; letter-spacing: .42em; text-indent: .42em;
  color: var(--fte-muted); font-weight: 400;
}
.fte-auth-title { font-size: 1.4rem; font-weight: 650; color: var(--fte-text); letter-spacing: .01em; line-height: 1.2; }
.fte-auth-sub { color: var(--fte-muted); font-size: .92rem; margin-top: .4rem; }
.fte-ws-card {
  border: 1px solid var(--fte-border); border-radius: 12px;
  padding: 1.4rem 1rem 1.1rem; text-align: center; background: transparent;
  transition: border-color var(--hover) ease, background var(--hover) ease,
              transform var(--hover) ease;
}
.fte-ws-card:hover {
  border-color: #33405a; background: rgba(127,127,127,.04); transform: translateY(-1px);
}
.fte-ws-card .ic { line-height: 1; margin-bottom: .6rem; color: var(--fte-text); }
.fte-ws-card .ic svg { display: block; margin: 0 auto; }
.fte-ws-card .nm { font-weight: 600; font-size: .95rem; color: var(--fte-text); letter-spacing: .02em; }
.fte-ws-card .tg { font-size: .76rem; color: var(--fte-muted); margin-top: .3rem; }
.fte-ws-header-brand { font-weight: 750; letter-spacing: .18em; font-size: 1rem; color: var(--fte-text); }
.fte-ws-header-brand span {
  color: var(--fte-muted); font-weight: 400; letter-spacing: .1em;
  font-size: .78rem; margin-left: .7rem;
}
[data-testid="stBaseButton-primary"] {
  min-height: 42px; border-radius: 8px; font-size: .9rem; letter-spacing: .05em;
  transition: transform var(--hover) ease, box-shadow var(--hover) ease,
              background var(--hover) ease;
}
[data-testid="stBaseButton-primary"]:hover { transform: translateY(-1px); }
[data-testid="stButton"] button:focus-visible,
[data-testid="stBaseButton-primary"]:focus-visible,
[data-testid="stBaseButton-secondary"]:focus-visible {
  outline: 2px solid var(--fte-conflict); outline-offset: 2px;
}
[data-testid="stSegmentedControl"] { margin: .5rem 0 .75rem; }
[data-testid="stSegmentedControl"] > div {
  gap: 2px; border: 1px solid var(--fte-border); border-radius: 8px;
  padding: 3px; width: fit-content; background: rgba(127,127,127,.05);
}
[data-testid="stSegmentedControl"] label {
  border-radius: 6px; padding: .32rem .95rem; font-size: .82rem; letter-spacing: .05em;
  color: var(--fte-muted);
  transition: background var(--hover) ease, color var(--hover) ease;
}
[data-testid="stSegmentedControl"] label:hover { color: var(--fte-text); }
[data-testid="stSegmentedControl"] label:has(input:checked) {
  background: rgba(127,127,127,.16); color: var(--fte-text); font-weight: 600;
}
[data-testid="stSegmentedControl"] label:has(input:focus-visible) {
  outline: 2px solid var(--fte-conflict); outline-offset: 1px;
}

/* Sprint 3.2.1: the memo body renders inside a bidirectional custom
   component as ONE continuous HTML document — metric tokens are inline
   spans (never boxed buttons), so sentences stay intact. The doc card
   rule below styles the wrapper; the component styles itself. */
[class*="st-key-fte_memo_doc"] {
  max-width: 860px; margin: .25rem auto 0;
  background: rgba(127,127,127,.04);
}
.fte-memo-para { margin-bottom: .9rem; }

/* Sprint 3: compact Intelligence cards + focused memo document view. */
.fte-intel-card {
  border: 1px solid var(--fte-border); border-radius: 10px;
  padding: .95rem 1.05rem; background: rgba(127,127,127,.04);
  transition: border-color var(--hover) ease, background var(--hover) ease;
}
.fte-intel-card:hover { border-color: rgba(127,127,127,.28); }
.fte-memo-doc {
  max-width: 860px; margin: .25rem auto 0; padding: 1.6rem 2rem;
  border: 1px solid var(--fte-border); border-radius: 10px;
  background: rgba(127,127,127,.04); color: var(--fte-text);
}
.fte-memo-title { font-size: 1.35rem; font-weight: 700; letter-spacing: .01em; }
.fte-memo-context { font-size: .8rem; color: var(--fte-muted); margin-top: .4rem; }
/* Sprint 3.2.4: inline metric links are plain inline text - subtle
   color + dotted underline, no box/pill/button chrome. Clicking one
   opens the same native floating metric dialog (query-param driven). */
.fte-memo-para a.fte-metric-link {
  color: var(--fte-conflict);
  font-weight: 600;
  text-decoration: underline dotted rgba(79, 142, 247, .65);
  text-underline-offset: 3px;
  cursor: pointer;
  border-radius: 2px;
  transition: color 180ms ease, background 180ms ease;
}
.fte-memo-para a.fte-metric-link:hover {
  color: #7aa6f9;
  background: rgba(79, 142, 247, .08);
}

/* Sprint 2.1: native st.dialog size refinement for the metric detail card.
   The built-in dialog component provides the backdrop, × close button,
   rounded corners, shadow and centering — we just constrain its width
   inside the FTE dark theme. */
[data-testid="stDialog"] { width: min(440px, 92vw); }
.fte-demo-note {
  margin: .4rem 0 .9rem; padding: .5rem .8rem; font-size: .8rem;
  color: var(--fte-muted); border: 1px dashed var(--fte-border); border-radius: 6px;
  background: rgba(217,161,59,.05); line-height: 1.55;
}
.fte-demo-note a { color: var(--fte-conflict); text-decoration: none; }
.fte-demo-note a:hover { text-decoration: underline; }
"""

_TERMINAL_METRICS = [
    ("Revenue", "Revenue"),
    ("Net Profit", "Net Income"),
    ("EBITDA", "EBITDA"),
    ("Operating Profit", "Operating Income"),
    ("EPS", "EPS"),
    ("Debt", "Total Debt"),
    ("Assets", "Total Assets"),
    ("Liabilities", "Total Liabilities"),
    ("Equity", "Shareholders' Equity"),
    ("Cash Flow", "Operating Cash Flow"),
    ("Profit Margin", "Profit Margin"),
    ("ROE", "ROE"),
    ("ROA", "ROA"),
    ("Debt to Equity", "Debt / Equity"),
    ("Current Ratio", "Current Ratio"),
    ("CAGR", "CAGR"),
]


def _inject_terminal_css() -> None:
    """Inject the terminal stylesheet once per session."""
    if st.session_state.get("fte_css_injected"):
        return
    st.markdown(f"<style>{_TERMINAL_CSS}\n{_FTE_CSS}</style>", unsafe_allow_html=True)
    st.session_state["fte_css_injected"] = True


def _init_terminal_state() -> None:
    """Session-state defaults for the terminal (selected metric, results)."""
    for key, default in [
        ("fte_selected_metric", None),
        ("fte_module3_result", {}),
        ("fte_module3_key", None),
        ("fte_css_injected", False),
        ("fte_metric_filter", "all"),
        ("fte_route", "entrance"),
        ("fte_page", "Financial Grid"),
        ("fte_workspace", None),
        ("fte_user_name", ""),
        ("fte_user_email", ""),
        ("fte_forgot_hint", False),
        ("fte_memo_draft", ""),
        ("fte_memo_view_open", False),
        ("fte_memo_status", "idle"),
        ("fte_memo_metric_click", None),
        ("fte_last_memo_nonce", None),
        ("fte_overlay_open", False),
        ("fte_prev_selection_rows", ()),
        ("fte_demo_mode", False),
        ("fte_demo_chat_messages", []),
    ]:
        if key not in st.session_state:
            st.session_state[key] = default


def _fmt_num(v) -> str:
    """Compact human formatting for grid values (real value, presentation only)."""
    if v is None:
        return "—"
    try:
        f = float(v)
    except (TypeError, ValueError):
        return str(v)
    if abs(f) >= 1e9:
        return f"{f / 1e9:,.2f}B"
    if abs(f) >= 1e6:
        return f"{f / 1e6:,.2f}M"
    if abs(f) >= 1e3:
        return f"{f / 1e3:,.2f}K"
    if f == int(f):
        return f"{int(f):,}"
    return f"{f:,.2f}".rstrip("0").rstrip(".")


def _parse_gateway_latency():
    """Last recorded provider latency from the activity log, if any."""
    for entry in reversed(st.session_state.get("provider_log", []) or []):
        detail = entry.get("detail") or ""
        idx = detail.find("latency=")
        if idx >= 0:
            num = ""
            for ch in detail[idx + len("latency="):]:
                if ch.isdigit() or ch == ".":
                    num += ch
                else:
                    break
            if num:
                try:
                    return float(num)
                except ValueError:
                    pass
    return None


def render_gateway_pill() -> None:
    """Compact header chrome: '🌐 Gateway · 🟢 Groq · 412ms' with a popover."""
    status = get_canonical_provider_status()
    latency = _parse_gateway_latency()
    active = next((s for s, stt in status.items() if stt == "available"), None)
    if active is None:
        active = next((s for s, stt in status.items() if stt == "configured_unavailable"), None)
    label = (active or "none").replace("_", " ").title()
    if active and status.get(active) == "available":
        dot, icon = "ok", "🟢"
    elif active:
        dot, icon = "warn", "🟡"
    else:
        dot, icon = "off", "⚪"
    latency_txt = f" · {latency:.0f}ms" if latency else ""
    with st.popover(f"🌐 Gateway · {icon} {label}{latency_txt}", use_container_width=True):
        st.caption("Provider status (masked — no secrets shown)")
        for slug, stt in status.items():
            icon_s = {"available": "🟢", "configured_unavailable": "🟡", "not_configured": "⚪"}.get(stt, "⚪")
            st.markdown(f"{icon_s} `{slug}` — {stt.replace('_', ' ')}")
        if latency:
            st.caption(f"Last latency: {latency:.0f}ms")
        st.caption(f"Active provider: {st.session_state.get('ai_provider_used') or '—'}")


def _build_terminal_rows(module3_result):
    """Canonical grid rows from module3 financial_data + ratios ONLY.

    Status semantics (presentation only — backend values are untouched):
      🔵 Conflict   — cross-document verifier reports disagreement
      🔴 Blocked    — required metric listed as missing by the detector
      🟡 Derived    — calculated ratio (source == "Calculated")
      🟢 Verified   — value + source present in the backend output
      ⚪ Unanalyzed — pipeline has not established the metric yet
    A metric never receives a stronger status than its evidence
    supports; missing provenance fields render as '—' and the tray
    distinguishes available from unavailable metadata. Blocked/conflict
    keys reported by the pipeline but absent from the canonical list are
    appended so exception filters can reach them."""
    financial_data = (module3_result or {}).get("financial_data") or {}
    ratios = (module3_result or {}).get("ratios") or {}
    conflicts = _conflict_metrics(module3_result)
    blocked = _blocked_metrics(module3_result)
    rows = []
    seen = set()

    for key, label in _TERMINAL_METRICS:
        seen.add(key)
        fact = financial_data.get(key) or ratios.get(key) or {}
        value = fact.get("value")
        source = fact.get("source") or ""
        period = fact.get("reporting_period") or fact.get("period") or "—"
        if value is not None:
            if key in conflicts:
                rows.append({
                    "metric": key, "Metric": label, "Value": _fmt_num(value),
                    "Period": period, "Source": source or "—", "Status": "🔵 Conflict",
                    "_kind": "conflict", "_fact": fact,
                })
            elif source == "Calculated":
                rows.append({
                    "metric": key, "Metric": label, "Value": _fmt_num(value),
                    "Period": period, "Source": "Calculated", "Status": "🟡 Derived",
                    "_kind": "derived", "_fact": fact,
                })
            else:
                rows.append({
                    "metric": key, "Metric": label, "Value": _fmt_num(value),
                    "Period": period, "Source": source or "Document", "Status": "🟢 Verified",
                    "_kind": "verified", "_fact": fact,
                })
        elif key in blocked:
            rows.append({
                "metric": key, "Metric": label, "Value": "—", "Period": "—",
                "Source": "—", "Status": "🔴 Blocked", "_kind": "blocked",
                "_reason": blocked[key], "_fact": {},
            })
        else:
            rows.append({
                "metric": key, "Metric": label, "Value": "—", "Period": "—",
                "Source": "—", "Status": "⚪ Unanalyzed", "_kind": "unanalyzed",
                "_fact": {},
            })

    # Exceptions reported by the pipeline but outside the canonical list
    # (e.g. PAT, ROCE) — appended as real, honest rows so the grid's
    # exception filters can reach them.
    for key, reason in blocked.items():
        if key not in seen:
            seen.add(key)
            rows.append({
                "metric": key, "Metric": key, "Value": "—", "Period": "—",
                "Source": "—", "Status": "🔴 Blocked", "_kind": "blocked",
                "_reason": reason, "_fact": {},
            })
    for key, info in conflicts.items():
        if key not in seen:
            seen.add(key)
            rows.append({
                "metric": key, "Metric": key, "Value": "—", "Period": "—",
                "Source": "—", "Status": "🔵 Conflict", "_kind": "conflict",
                "_fact": info,
            })
    return rows


def _run_ingestion(uploaded_files):
    """Shared ingestion for the terminal (same pipeline as the classic view)."""
    extraction_results = []
    combined_raw_text = ""
    document_summaries = []
    if uploaded_files:
        extraction_results = extract_multiple(uploaded_files)
        combined_raw_text = merge_document_text(extraction_results)
        for doc in extraction_results:
            file_name = doc["file_name"]
            extracted_text = doc["parsed_document"]["text"]
            cache_key = f"{file_name}:{_hash_text(extracted_text)}"
            summary_text, _ = _cached_call(
                "summary_cache", cache_key,
                lambda et=extracted_text, fn=file_name: summarize_document_with_chunking(et, fn)
            )
            document_summaries.append({"file_name": file_name, "summary": summary_text})
    return extraction_results, combined_raw_text, document_summaries


def _terminal_module3(combined_raw_text, extraction_results):
    """Run Module 3 (cached, same cache key as the classic view)."""
    if not combined_raw_text:
        return {}
    key = _hash_text(combined_raw_text)
    if st.session_state["fte_module3_key"] == key and st.session_state["fte_module3_result"]:
        return st.session_state["fte_module3_result"]
    result, _ = _cached_call(
        "module3_cache", key,
        lambda: run_module3(combined_raw_text, extraction_results)
    )
    st.session_state["fte_module3_result"] = result or {}
    st.session_state["fte_module3_key"] = key
    return result or {}


_BLOCKED_INTENTS = (
    "blocked", "calculation_miss", "document_miss", "company_blocked",
    "company_empty", "company_error", "no_provider", "provider_error",
)


def _conflict_metrics(module3_result) -> dict:
    """Metric keys the cross-document verifier marks as 'Conflict'
    (multiple sources disagree). Reads the real module3 shape (dict of
    {field: {status, verified_value, documents}}) defensively; a legacy
    list shape is also accepted. Presentation only — the verifier's
    output is produced by the backend and is never altered here."""
    x = (module3_result or {}).get("cross_document_verification")
    conflicts = {}
    if isinstance(x, dict):
        for field, info in x.items():
            if isinstance(info, dict) and str(info.get("status", "")).lower() == "conflict":
                conflicts[str(field)] = info
    elif isinstance(x, list):
        for item in x:
            if isinstance(item, dict) and item.get("match") is False and item.get("metric"):
                conflicts[str(item["metric"])] = item
    return conflicts


def _blocked_metrics(module3_result) -> dict:
    """Metric keys the pipeline's missing-data detector lists as required
    but absent (financial_data + ratios sections). Returns {key: reason}.
    Presentation only — the report is produced by the backend
    (backend/missing_data_detector.py)."""
    md = (module3_result or {}).get("missing_data")
    blocked = {}
    if isinstance(md, dict):
        for section, label in (
            ("financial_data", "Required financial evidence is missing"),
            ("ratios", "Required ratio is not available"),
        ):
            items = md.get(section)
            if isinstance(items, list):
                for key in items:
                    if isinstance(key, str) and key.strip():
                        blocked.setdefault(key.strip(), f"{label} ({section})")
            elif isinstance(items, dict):
                for key, present in items.items():
                    if not present:
                        blocked.setdefault(str(key), f"{label} ({section})")
    elif isinstance(md, list):
        for key in md:
            if isinstance(key, str) and key.strip():
                blocked.setdefault(key.strip(), "Listed as missing by the pipeline report")
    return blocked


def _row_counts(rows) -> dict:
    """Taxonomy counts over grid rows (real rows only, nothing invented)."""
    counts = {"verified": 0, "derived": 0, "unanalyzed": 0, "conflict": 0, "blocked": 0}
    for r in rows or []:
        kind = r.get("_kind", "unanalyzed")
        counts[kind] = counts.get(kind, 0) + 1
    return counts


_FILTER_LABELS = {
    "all": "All metrics",
    "verified": "🟢 Verified",
    "derived": "🟡 Derived",
    "unanalyzed": "⚪ Unanalyzed",
    "conflict": "🔵 Conflicts",
    "blocked": "🔴 Blocked",
}

# Plain (un-emoji'd) kind words for the grid status caption.
_METRIC_FILTER_PLAIN = {
    "all": "all",
    "verified": "verified",
    "derived": "derived",
    "unanalyzed": "unanalyzed",
    "conflict": "conflict",
    "blocked": "blocked",
}


def _set_taxonomy_filter(kind: str) -> None:
    """Button callback: set the grid metric filter (runs pre-rerun).
    Fires only on an explicit user click — never on plain reruns."""
    st.session_state["fte_metric_filter"] = kind


def _grid_rows_filtered(rows, kind: str):
    """Filter grid rows by taxonomy state; 'all' returns everything."""
    if not kind or kind == "all":
        return list(rows)
    return [r for r in rows if r.get("_kind") == kind]


def _grid_status_caption(rows, shown, filter_kind) -> str:
    """Caption under the Financial Grid title (presentation only).

    'all' (the default) reads as dataset status — never as a filter:
      "Showing all 20 metrics · 4 verified · 1 derived · 7 blocked · 8 unanalyzed"
    An explicit filter reads as a subset view:
      "Showing 4 verified metrics · 4 of 20"
    """
    if not filter_kind or filter_kind == "all":
        counts = _row_counts(rows)
        return (
            f"Showing all {len(rows)} metrics · {counts['verified']} verified · "
            f"{counts['derived']} derived · {counts['blocked']} blocked · "
            f"{counts['unanalyzed']} unanalyzed"
        )
    plain_label = _METRIC_FILTER_PLAIN.get(filter_kind, filter_kind)
    return f"Showing {len(shown)} {plain_label} metrics · {len(shown)} of {len(rows)}"


def _resolve_grid_selection(shown, selected_metric, sel_rows):
    """Resolve the metric that should be selected after a rerun.

    - A fresh click (sel_rows) wins.
    - Otherwise keep the persisted selected metric while it is visible.
    - If the persisted metric is hidden by an explicit filter, fall back
      to the first visible metric (or None when nothing is visible).
    Selection stays independent of the filter state.
    """
    if sel_rows:
        idx = int(sel_rows[0])
        if 0 <= idx < len(shown):
            return shown[idx]["metric"]
        return None
    if not shown:
        return None
    if selected_metric is None:
        return None
    if all(r["metric"] != selected_metric for r in shown):
        return shown[0]["metric"]
    return selected_metric


def _render_taxonomy_controls(rows, include_exceptions: bool = False) -> None:
    """Actionable taxonomy filter buttons (rail). Clicking a state
    filters the Financial Grid to that state; 'All metrics' restores.
    Presentation only — metric values are never altered."""
    counts = _row_counts(rows)
    current = st.session_state.get("fte_metric_filter", "all")
    if include_exceptions:
        options = [
            ("conflict", "🔵 Conflicts", counts["conflict"]),
            ("blocked", "🔴 Blocked", counts["blocked"]),
        ]
    else:
        options = [
            ("all", "All metrics", 0),
            ("verified", "🟢 Verified", counts["verified"]),
            ("derived", "🟡 Derived", counts["derived"]),
            ("unanalyzed", "⚪ Unanalyzed", counts["unanalyzed"]),
        ]
    for kind, label, count in options:
        label_txt = label if kind == "all" else f"{label} · {count}"
        st.button(
            label_txt,
            key=f"fte_filter_{kind}_{'x' if include_exceptions else 'm'}",
            use_container_width=True,
            type="primary" if current == kind else "secondary",
            on_click=_set_taxonomy_filter,
            args=(kind,),
        )


def _count_conflicts(module3_result) -> int:
    """Conflict count from cross-document verification (real data only)."""
    return len(_conflict_metrics(module3_result))


def _count_blocked(module3_result) -> int:
    """Blocked count from the missing-data detector (real data only)."""
    return len(_blocked_metrics(module3_result))


def _render_source_rail(uploaded_files, extraction_results, module3_result) -> None:
    """Left rail: active sources, taxonomy integrity, exceptions."""
    st.markdown('<div class="fte-rail-title">📁 Active Sources</div>', unsafe_allow_html=True)
    if not uploaded_files:
        st.markdown('<div class="fte-cap">No documents uploaded.</div>', unsafe_allow_html=True)
    else:
        for i, doc in enumerate(extraction_results or []):
            name = doc.get("file_name", "Unknown Document")
            state = "✓ cached" if doc.get("cached") else "✓ parsed"
            st.markdown(
                f'<div class="fte-source"><span class="nm">{html.escape(str(name))}</span>'
                f'<span class="st">{state}</span></div>',
                unsafe_allow_html=True,
            )
        if not extraction_results:
            for f in uploaded_files:
                st.markdown(
                    f'<div class="fte-source"><span class="nm">{html.escape(str(getattr(f, "name", "file")))}</span>'
                    '<span class="st">queued</span></div>',
                    unsafe_allow_html=True,
                )

    rows = _build_terminal_rows(module3_result)

    st.markdown('<div class="fte-rail-title">🧭 Taxonomy Integrity</div>', unsafe_allow_html=True)
    _render_taxonomy_controls(rows, include_exceptions=False)

    st.markdown('<div class="fte-rail-title">🚨 Exceptions</div>', unsafe_allow_html=True)
    _render_taxonomy_controls(rows, include_exceptions=True)


def _render_analysis_limited_banner(rows) -> None:
    """One compact institutional state line above the grid (no repeated
    blocked messaging) with an 'Inspect exceptions' action."""
    counts = _row_counts(rows)
    show = (counts["blocked"] + counts["conflict"] > 0) or (
        counts["unanalyzed"] > 0 and (counts["verified"] + counts["derived"]) > 0
    )
    if not show:
        return
    line = (
        f"⚠️ **Analysis limited** · {counts['blocked']} blocked · "
        f"{counts['unanalyzed']} unanalyzed · {counts['verified']} verified"
    )
    b_col, btn_col = st.columns([4, 1], gap="small")
    with b_col:
        st.markdown(f'<div class="fte-limited">{line}</div>', unsafe_allow_html=True)
    with btn_col:
        st.button(
            "Inspect exceptions",
            key="fte_inspect_exceptions",
            use_container_width=True,
            on_click=_set_taxonomy_filter,
            args=("blocked",),
        )


def _render_financial_grid(module3_result) -> None:
    """Center grid: stable metric table with single-row selection and
    taxonomy filtering. Row dimensions never change between reruns."""
    st.markdown('<div class="fte-rail-title">Financial Grid</div>', unsafe_allow_html=True)
    rows = _build_terminal_rows(module3_result)
    st.session_state["fte_grid_rows"] = rows

    _render_analysis_limited_banner(rows)

    filter_kind = st.session_state.get("fte_metric_filter", "all")
    shown = _grid_rows_filtered(rows, filter_kind)
    st.caption(_grid_status_caption(rows, shown, filter_kind))

    df = pd.DataFrame(
        [{"Metric": r["Metric"], "Value": r["Value"], "Period": r["Period"],
          "Source": r["Source"], "Status": r["Status"]} for r in shown]
    )
    if df.empty:
        if filter_kind != "all":
            st.caption("No metrics match the current filter — choose another taxonomy state or 'All metrics'.")
        else:
            st.caption("Upload a financial document and run the intelligence pipeline to populate the grid.")
        st.session_state["fte_overlay_open"] = False
        return

    st.dataframe(
        df,
        use_container_width=True,
        hide_index=True,
        key="fte_grid_table",
        on_select="rerun",
        selection_mode="single-row",
        column_config={
            "Metric": st.column_config.TextColumn("Metric"),
            "Value": st.column_config.TextColumn("Value", width="small", alignment="right"),
            "Period": st.column_config.TextColumn("Period", width="small", alignment="center"),
            "Source": st.column_config.TextColumn("Source", width="medium"),
            "Status": st.column_config.TextColumn("Status", width="medium"),
        },
    )

    # Selection is independent of the filter: honor a fresh click;
    # otherwise keep the persisted selected metric while it is visible,
    # and fall back to the first visible metric when an explicit filter
    # hides it (grid and Provenance Tray always stay consistent).
    try:
        sel = st.session_state.get("fte_grid_table")
        sel_rows = sel.selection.rows if sel is not None else ()
    except Exception:
        sel_rows = ()
    prev_sel_rows = tuple(st.session_state.get("fte_prev_selection_rows") or ())
    st.session_state["fte_prev_selection_rows"] = tuple(sel_rows) if sel_rows else ()
    fresh_click = bool(sel_rows) and tuple(sel_rows) != prev_sel_rows

    was_selected = st.session_state.get("fte_selected_metric")
    resolved = _resolve_grid_selection(shown, was_selected, sel_rows)
    # A stale selection index (row order changed after a filter) must not
    # leave the grid unselected: fall back to the persisted-metric rules.
    if resolved is None and shown:
        resolved = _resolve_grid_selection(shown, was_selected, ())
    if resolved is not None:
        st.session_state["fte_selected_metric"] = resolved
    if not sel_rows and resolved is not None:
        try:
            widget = st.session_state.get("fte_grid_table")
            restore_idx = next(
                (i for i, r in enumerate(shown) if r["metric"] == resolved), None
            )
            if widget is not None and restore_idx is not None and not widget.selection.rows:
                widget.selection.rows = [restore_idx]
        except Exception:
            pass

    # Sprint 2: focused metric-detail overlay. Opens on a fresh row click;
    # stays open only while its metric remains selected and visible.
    if fresh_click and resolved is not None:
        st.session_state["fte_overlay_open"] = True
    elif st.session_state.get("fte_overlay_open"):
        metric_now = st.session_state.get("fte_selected_metric")
        visible = metric_now is not None and any(
            r["metric"] == metric_now for r in shown
        )
        if not visible or resolved != was_selected:
            st.session_state["fte_overlay_open"] = False

    if st.session_state.get("fte_overlay_open") and st.session_state.get("fte_selected_metric"):
        _metric_detail_dialog(module3_result)


def _provenance_tray_html(rows, module3_result, metric) -> str:
    """Provenance tray content — only fields the pipeline provides; '—'
    otherwise. Available evidence and unavailable metadata are visually
    distinct (fte-pv-avail / fte-pv-missing); blocked metrics get one
    compact limitation note instead of repeated blocked banners."""
    selected = next((r for r in rows if r["metric"] == metric), None)
    if selected is None:
        return '<div class="fte-tray"><div class="fte-tray-empty">Select a metric in the grid to inspect its provenance.</div></div>'
    kind = selected.get("_kind", "unanalyzed")
    fd = (module3_result or {}).get("financial_data") or {}
    rt = (module3_result or {}).get("ratios") or {}
    fact = selected.get("_fact") or fd.get(selected["metric"]) or rt.get(selected["metric"]) or {}

    def g(*keys, default="—"):
        for k in keys:
            v = fact.get(k) if isinstance(fact, dict) else None
            if v not in (None, ""):
                return str(v)
        return default

    status = selected.get("Status") or "—"
    period = selected.get("Period") or "—"
    source = selected.get("Source") or "—"
    value = selected.get("Value") or "—"
    note = None
    if kind == "blocked":
        origin = "Pipeline missing-data report"
        location = "—"
        evidence = "—"
        note = selected.get("_reason") or "Required evidence is not available from the current pipeline."
    elif kind == "conflict":
        origin = "Cross-document verification"
        location = "—"
        conflicts = _conflict_metrics(module3_result)
        info = conflicts.get(selected["metric"]) or {}
        docs = info.get("documents") or []
        if isinstance(docs, list) and docs:
            evidence = " · ".join(
                f"{d.get('document', '?')}={d.get('value', '?')}"
                for d in docs if isinstance(d, dict)
            )
        else:
            evidence = "Multiple sources disagree"
    else:
        origin = "Calculated ratio" if kind == "derived" else ("Document extraction" if kind == "verified" else "—")
        location = g("page", "table_id", "chunk_id", "anchor")
        evidence = g("anchor", "evidence", "context")

    currency = g("unit", "currency", "currency_code")
    scale = g("scale")

    item_rows = [
        ("Metric", selected.get("Metric") or "—"),
        ("Value", value),
        ("Status", status),
        ("Origin", origin),
        ("Period", period),
        ("Source", source),
        ("Location", location),
        ("Currency", currency),
        ("Scale", scale),
        ("Evidence", evidence),
    ]

    def _row(k, v):
        v = "—" if v in (None, "") else str(v)
        cls = "fte-pv-missing" if v == "—" else "fte-pv-avail"
        return f'<div class="fte-metric-row {cls}"><div class="k">{html.escape(str(k))}</div><div class="v">{html.escape(v)}</div></div>'

    body = "".join(_row(k, v) for k, v in item_rows)
    head = f'{html.escape(str(selected.get("Metric") or ""))} · {html.escape(str(selected.get("Status") or ""))}'
    if kind in ("verified", "derived") and (period == "—" or scale == "—"):
        head += ' <span class="fte-partial">· partial provenance</span>'
    if note:
        body += f'<div class="fte-block-note">⚠️ Analysis limited — {html.escape(str(note))}</div>'
    return f'<div class="fte-tray"><div class="fte-tray-head">{head}</div>{body}</div>'


def _render_provenance_tray(module3_result) -> None:
    """Anchored provenance tray below the grid (persistent lower section)."""
    rows = st.session_state.get("fte_grid_rows") or []
    metric = st.session_state.get("fte_selected_metric")
    st.markdown(_provenance_tray_html(rows, module3_result, metric), unsafe_allow_html=True)


# =============================================================================
# Sprint 2: focused metric-detail overlay (presentation only — pipeline is the
# only source of truth; missing fields render "—", never fabricated).
# =============================================================================
# Static, human-written one-line explanations. Educational text only — never
# used to generate values. Metrics without an entry simply omit the section.
_METRIC_EXPLAINER = {
    "revenue": "Money the company earned from selling its products or services during the period.",
    "net profit": "What the company kept after all expenses, interest and taxes.",
    "net income": "What the company kept after all expenses, interest and taxes.",
    "ebitda": "Operating profit before interest, taxes, depreciation and amortisation — a cash-earnings view.",
    "operating profit": "Profit from core operations, before interest and tax.",
    "operating income": "Profit from core operations, before interest and tax.",
    "eps": "Earnings per share — the profit attributable to each share of stock.",
    "debt": "Borrowings the company must repay, including loans and bonds.",
    "total debt": "Borrowings the company must repay, including loans and bonds.",
    "assets": "Everything the company owns or is owed, as recorded on the balance sheet.",
    "total assets": "Everything the company owns or is owed, as recorded on the balance sheet.",
    "liabilities": "What the company owes to others.",
    "total liabilities": "What the company owes to others.",
    "equity": "Shareholders' equity — the owners' stake in the company.",
    "shareholders' equity": "Shareholders' equity — the owners' stake in the company.",
    "cash flow": "Cash generated from operations during the period.",
    "operating cash flow": "Cash generated from core operations during the period.",
    "profit margin": "Profit as a share of revenue — how much of each sale is kept.",
    "operating margin": "Operating profit as a share of revenue — core business efficiency.",
    "net margin": "Net profit as a share of revenue — overall profitability per rupee of sales.",
    "roe": "Return on equity — the profit the company generates per unit of shareholder money.",
    "roa": "Return on assets — the profit the company generates per unit of total assets.",
    "debt to equity": "Borrowings relative to shareholder equity — a measure of financial leverage.",
    "debt/equity": "Borrowings relative to shareholder equity — a measure of financial leverage.",
    "debt / equity": "Borrowings relative to shareholder equity — a measure of financial leverage.",
    "current ratio": "Current assets divided by current liabilities — a short-term solvency check.",
    "cagr": "Compound annual growth rate — the smoothed yearly growth over the period.",
    "pat": "Profit after tax — the bottom-line profit attributable to shareholders.",
    "pat growth": "Year-on-year change in profit after tax, expressed as a percentage.",
    "revenue growth": "Year-on-year change in revenue, expressed as a percentage.",
    "free cash flow": "Cash left after capital spending — funds available to investors.",
    "market cap": "Total market value of the company's shares.",
}


def _metric_explainer(metric: str) -> str:
    """Short human-readable explanation for a metric name (static text)."""
    if not metric:
        return ""
    norm = str(metric).strip().lower()
    return _METRIC_EXPLAINER.get(norm, "")


def _metric_overlay_fields(rows, module3_result, metric) -> dict:
    """Detail fields for the overlay — only what the pipeline provides;
    every missing field is '—'. Never fabricates provenance."""
    out = {
        "metric": metric, "label": metric or "—", "kind": "unanalyzed",
        "value": "—", "status": "—", "origin": "—", "period": "—",
        "source": "—", "location": "—", "currency": "—", "scale": "—",
        "evidence": "—", "note": None, "fragment": None, "fact": None,
    }
    if not metric:
        return out
    selected = next((r for r in rows if r["metric"] == metric), None)
    if selected is None:
        return out
    kind = selected.get("_kind", "unanalyzed")
    out["kind"] = kind
    out["label"] = selected.get("Metric") or metric
    fd = (module3_result or {}).get("financial_data") or {}
    rt = (module3_result or {}).get("ratios") or {}
    fact = selected.get("_fact") or fd.get(metric) or rt.get(metric) or {}
    out["fact"] = fact

    def g(*keys, default="—"):
        for k in keys:
            v = fact.get(k) if isinstance(fact, dict) else None
            if v not in (None, ""):
                return str(v)
        return default

    out["status"] = selected.get("Status") or "—"
    out["period"] = selected.get("Period") or "—"
    out["source"] = selected.get("Source") or "—"
    out["value"] = selected.get("Value") or "—"
    if kind == "blocked":
        out["origin"] = "Pipeline missing-data report"
        out["note"] = selected.get("_reason") or "Required evidence is not available from the current pipeline."
    elif kind == "conflict":
        out["origin"] = "Cross-document verification"
        conflicts = _conflict_metrics(module3_result)
        info = conflicts.get(metric) or {}
        docs = info.get("documents") or []
        if isinstance(docs, list) and docs:
            evidence = " · ".join(
                f"{d.get('document', '?')}={d.get('value', '?')}"
                for d in docs if isinstance(d, dict)
            )
            out["evidence"] = evidence
            out["fragment"] = evidence
        else:
            out["evidence"] = "Multiple sources disagree"
    else:
        out["origin"] = (
            "Calculated ratio" if kind == "derived"
            else ("Document extraction" if kind == "verified" else "—")
        )
        out["location"] = g("page", "table_id", "chunk_id", "anchor")
        out["evidence"] = g("anchor", "evidence", "context")
        if out["evidence"] != "—":
            out["fragment"] = out["evidence"]
    out["currency"] = g("unit", "currency", "currency_code")
    out["scale"] = g("scale")
    return out


def _metric_detail_html(fields: dict, explainer: str) -> str:
    """Compact card body HTML for the st.dialog — no source-fragment
    section (the card stays concise). Only real fields render; '—' otherwise."""
    parts = [
        f'<div style="display:flex;align-items:baseline;justify-content:space-between;gap:10px">'
        f'<span style="font-weight:650;font-size:1.05rem;color:var(--fte-text)">{html.escape(str(fields["label"]))}</span>'
        f'<span style="font-size:.85rem;white-space:nowrap">{html.escape(str(fields["status"]))}</span>'
        f'</div>',
        f'<div style="font-size:1.55rem;font-weight:700;color:var(--fte-text);margin:.55rem 0 .9rem;letter-spacing:.01em">{html.escape(str(fields["value"]))}</div>',
    ]
    if explainer:
        parts.append(
            f'<div style="border-top:1px solid var(--fte-border);margin-top:.7rem;padding-top:.6rem">'
            f'<div style="font-size:11px;letter-spacing:.12em;text-transform:uppercase;color:var(--fte-muted);font-weight:600;margin-bottom:.35rem">What it means</div>'
            f'<div style="font-size:.86rem;color:var(--fte-text);line-height:1.5">{html.escape(explainer)}</div></div>'
        )
    pv_rows = [
        ("Origin", fields["origin"]),
        ("Period", fields["period"]),
        ("Source", fields["source"]),
        ("Evidence", fields["evidence"]),
    ]
    body = ""
    for k, v in pv_rows:
        vtxt = "—" if v in (None, "") else str(v)
        mk = "missing" if vtxt == "—" else ""
        body += (
            f'<div class="{mk}" style="display:grid;grid-template-columns:96px 1fr;gap:5px;font-size:.82rem;padding:2px 0">'
            f'<div style="color:var(--fte-muted)">{html.escape(str(k))}</div>'
            f'<div style="color:var(--fte-text);overflow-wrap:anywhere">{html.escape(vtxt)}</div></div>'
        )
    parts.append(
        f'<div style="border-top:1px solid var(--fte-border);margin-top:.7rem;padding-top:.6rem">'
        f'<div style="font-size:11px;letter-spacing:.12em;text-transform:uppercase;color:var(--fte-muted);font-weight:600;margin-bottom:.35rem">Provenance</div>{body}</div>'
    )
    if fields.get("note"):
        parts.append(
            f'<div style="margin-top:.65rem;padding:.45rem .6rem;font-size:.78rem;color:var(--fte-blocked);border-left:2px solid var(--fte-blocked);background:rgba(224,82,82,.06);border-radius:4px">⚠️ Analysis limited — {html.escape(str(fields["note"]))}</div>'
        )
    return "".join(parts)


def _fte_close_overlay() -> None:
    """Close the metric-detail overlay (selection itself is kept)."""
    st.session_state["fte_overlay_open"] = False


@st.dialog("Metric Detail", width="small", dismissible=True, on_dismiss="rerun")
def _metric_detail_dialog(module3_result) -> None:
    """Native Streamlit dialog for the selected metric detail.
    Small centered floating card; built-in × close; grid stays visible.
    Closes automatically when the user dismisses or selects another metric."""
    metric = st.session_state.get("fte_selected_metric")
    if not metric:
        st.caption("No metric selected.")
        return
    rows = st.session_state.get("fte_grid_rows") or []
    if not any(r["metric"] == metric for r in rows):
        st.caption("The selected metric is no longer visible.")
        return
    fields = _metric_overlay_fields(rows, module3_result, metric)
    explainer = _metric_explainer(metric)
    st.markdown(_metric_detail_html(fields, explainer), unsafe_allow_html=True)
    st.button("Close", key="fte_dlg_close", on_click=_fte_close_overlay)


def _render_intelligence_tab(module3_result, intelligence_outputs) -> None:
    """Intelligence tab: module3 sections + AI intelligence outputs (existing renderers)."""
    st.markdown('<div class="fte-rail-title">Module 3 — Deterministic Intelligence</div>', unsafe_allow_html=True)
    with st.expander("📊 Financial Data", expanded=False):
        _render_module3_value((module3_result or {}).get("financial_data"))
    with st.expander("📐 Ratios", expanded=False):
        _render_module3_value((module3_result or {}).get("ratios"))
    with st.expander("🧾 Verification (OCR / Cross-document)", expanded=False):
        _render_module3_value((module3_result or {}).get("ocr_verification"))
        _render_module3_value((module3_result or {}).get("cross_document_verification"))
    with st.expander("🎯 Confidence", expanded=False):
        _render_module3_value((module3_result or {}).get("confidence"))
    with st.expander("📅 Events / Timeline", expanded=False):
        _render_module3_value((module3_result or {}).get("events"))
        _render_module3_value((module3_result or {}).get("timeline"))
    with st.expander("🗜️ Optimized Context (sent to AI)", expanded=False):
        _render_module3_value((module3_result or {}).get("optimized_context"))

    st.markdown('<div class="fte-rail-title">AI Intelligence Sections</div>', unsafe_allow_html=True)
    if intelligence_outputs:
        for module_key in INTELLIGENCE_MODULES:
            render_intelligence_output(module_key, intelligence_outputs.get(module_key))
    else:
        st.caption("Run the full analysis pipeline (Classic Dashboard → Generate Timeline Report) to populate AI sections.")


def _render_system_tab() -> None:
    """System tab: provider health, roadmap, diagnostics (technical detail lives here)."""
    st.markdown('<div class="fte-rail-title">🌐 Gateway</div>', unsafe_allow_html=True)
    status = get_canonical_provider_status()
    if status:
        for slug, stt in status.items():
            icon_s = {"available": "🟢", "configured_unavailable": "🟡", "not_configured": "⚪"}.get(stt, "⚪")
            st.markdown(f"{icon_s} `{slug}` — {stt.replace('_', ' ')}")
    else:
        st.caption("Gateway unavailable.")

    st.markdown('<div class="fte-rail-title">🔮 Roadmap</div>', unsafe_allow_html=True)
    try:
        roadmap = get_future_module_status()
        st.dataframe(
            pd.DataFrame([{"Module": v["name"], "Status": v["status"].title()} for v in roadmap.values()]),
            use_container_width=True, hide_index=True,
        )
    except Exception:
        st.caption("Roadmap unavailable.")

    st.markdown('<div class="fte-rail-title">🧾 Ingestion Statistics</div>', unsafe_allow_html=True)
    stats = st.session_state.get("ingestion_stats")
    if stats:
        st.caption(str(stats))
    else:
        st.caption("Upload a document to see ingestion statistics.")

    st.markdown('<div class="fte-rail-title">🔍 Pipeline Debug Trace</div>', unsafe_allow_html=True)
    debug_entries = st.session_state.get("pipeline_debug_log", [])
    if debug_entries:
        st.dataframe(
            pd.DataFrame([
                {"Stage": e["stage"], "Length": e["length"],
                 "Error Marker?": "⚠️" if e.get("error_marker_detected") else ""}
                for e in debug_entries
            ]),
            use_container_width=True, hide_index=True,
        )
    else:
        st.caption("No pipeline stages recorded for this run.")


def _copilot_structure(metadata, rows=None) -> list:
    """Structured sections for the Co-Pilot: real metadata fields plus,
    when grid rows are available, an honest evidence summary instead of
    a bare 'BLOCKED / NOT VERIFIED' message."""
    intent = (metadata or {}).get("intent", "")
    evidence = (metadata or {}).get("evidence") or []
    calc = (metadata or {}).get("calculation")
    parts = []
    if intent in _BLOCKED_INTENTS:
        parts.append("⚠️ Analysis limited — not verifiable from current evidence")
    if evidence:
        parts.append(f"🟢 {len(evidence)} verified fact(s)")
    if calc:
        parts.append("🟡 1 calculation")
    if rows:
        counts = _row_counts(rows)
        summary = []
        if counts["verified"]:
            summary.append(f"🟢 {counts['verified']} verified")
        if counts["derived"]:
            summary.append(f"🟡 {counts['derived']} derived")
        if counts["blocked"]:
            summary.append(f"🔴 {counts['blocked']} blocked")
        if counts["unanalyzed"]:
            summary.append(f"⚪ {counts['unanalyzed']} unanalyzed")
        if summary:
            parts.append(" · ".join(summary))
    return parts or ["⚪ No structured facts for this turn"]


def _render_co_pilot(extraction_results, provider_health) -> None:
    """Compact Co-Pilot: capability line, bounded conversation, structured
    sections, and an honest blocked/empty state (no duplicated messages)."""
    st.markdown("---")
    st.markdown('<div class="fte-rail-title">💬 Ask Co-Pilot</div>', unsafe_allow_html=True)

    rows = st.session_state.get("fte_grid_rows") or []

    # Active context (real data only): the uploaded document names.
    active_docs = [html.escape(str(d.get("file_name", ""))) for d in (extraction_results or [])]
    active_docs = [d for d in active_docs if d]
    context_line = " · ".join(active_docs[:3]) if active_docs else "No document loaded"
    st.markdown(f'<div class="fte-cap">📄 Active context: {context_line}</div>', unsafe_allow_html=True)

    # Capability status (real state only).
    m3 = st.session_state.get("fte_module3_result") or {}
    intel = st.session_state.get("intelligence_outputs") or {}
    caps = []
    caps.append("● Financial grid ready" if m3.get("financial_data") else "○ Financial grid awaiting data")
    caps.append("● Intelligence ready" if intel else "○ Intelligence pending full analysis")
    has_key = any((provider_health or {}).values())
    caps.append(
        "● AI provider live"
        if (has_key and st.session_state.get("ai_connected"))
        else ("○ AI provider configured" if has_key else "○ no AI provider configured")
    )
    st.markdown('<div class="fte-cap">' + " · ".join(caps) + "</div>", unsafe_allow_html=True)

    # Bounded conversation (last 4 messages).
    messages = st.session_state.get("chat_messages", []) or []
    last_assistant_intent = ""
    for msg in messages[-4:]:
        with st.chat_message(msg["role"]):
            st.markdown(msg["content"])
            if msg["role"] == "assistant" and msg.get("metadata"):
                metadata = msg["metadata"]
                last_assistant_intent = metadata.get("intent", "")
                structure = _copilot_structure(metadata, rows)
                st.caption(" · ".join(structure))
                _render_chat_provenance(metadata)

    # Honest blocked/empty state for broad questions: one compact line
    # plus what the workspace actually holds — never a second duplicate
    # 'BLOCKED / NOT VERIFIED' banner.
    if last_assistant_intent in _BLOCKED_INTENTS:
        counts = _row_counts(rows) if rows else {"verified": 0, "derived": 0, "unanalyzed": 0, "conflict": 0, "blocked": 0}
        st.markdown(
            '<div class="fte-caps">'
            f"Grid status — {counts['verified']} verified · {counts['derived']} derived · "
            f"{counts['blocked']} blocked · {counts['unanalyzed']} unanalyzed. "
            "Try a specific metric (revenue, ROE, current ratio…), a calculation (CAGR, change), "
            "or a company ticker."
            "</div>",
            unsafe_allow_html=True,
        )

    # Suggested follow-ups: a small set the verified pipeline can actually
    # answer; the chips lead into the same Co-Pilot flow.
    if not messages:
        st.caption("Suggested:")
        cols = st.columns(2)
        for i, prompt in enumerate(FTE_CHAT_SUGGESTED_PROMPTS[:4]):
            with cols[i % 2]:
                if st.button(prompt, key=f"fte_suggest_{i}", use_container_width=True):
                    _submit_chat_question(prompt, extraction_results, provider_health)
                    st.rerun()

    user_input = st.chat_input("Ask about the selected financial evidence…")
    if user_input:
        _submit_chat_question(user_input, extraction_results, provider_health)
        st.rerun()


# =============================================================================
# FT-E account & workspace architecture (frontend-only foundation)
# =============================================================================
# Minimal line icons (monochrome, stroke-based) for the workspace cards.
_ICON_STUDENT = (
    '<svg width="26" height="26" viewBox="0 0 24 24" fill="none" stroke="currentColor" '
    'stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round">'
    '<path d="M2 9.5 L12 4.5 L22 9.5 L12 14.5 Z"/>'
    '<path d="M6 11.8 v3.1 c0 1.7 2.7 3.1 6 3.1 s6 -1.4 6 -3.1 v-3.1"/>'
    '<path d="M12 14.5 v3.5"/></svg>'
)
_ICON_PROFESSIONAL = (
    '<svg width="26" height="26" viewBox="0 0 24 24" fill="none" stroke="currentColor" '
    'stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round">'
    '<path d="M3.5 3.5 v17 h17"/>'
    '<path d="M6.5 14.5 l3.5 -3.5 l3 3 l4.5 -5.5"/></svg>'
)
_ICON_CA = (
    '<svg width="26" height="26" viewBox="0 0 24 24" fill="none" stroke="currentColor" '
    'stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round">'
    '<path d="M6 3 h9 l4 4 v14 H6 Z"/>'
    '<path d="M9.5 12.5 l2 2 l3.5 -4"/></svg>'
)
_ICON_ENTERPRISE = (
    '<svg width="26" height="26" viewBox="0 0 24 24" fill="none" stroke="currentColor" '
    'stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round">'
    '<path d="M5 21 V8.5 h14 V21"/>'
    '<path d="M9 8.5 V4 h6 v4.5"/>'
    '<path d="M8.5 13 h1.5 M8.5 17 h1.5 M14 13 h1.5 M14 17 h1.5"/></svg>'
)


def _fte_goto(route: str) -> None:
    """Move between FT-E views (frontend routing state only)."""
    st.session_state["fte_route"] = route


def _fte_forgot() -> None:
    """Forgot-password hint (UI scaffolding only — real recovery later)."""
    st.session_state["fte_forgot_hint"] = True


def _fte_set_workspace(kind: str) -> None:
    """Persist the chosen workspace and enter the FT-E shell."""
    st.session_state["fte_workspace"] = kind
    st.session_state["fte_page"] = "Financial Grid"
    st.session_state["fte_route"] = "workspace"


def _fte_display_name(email: str) -> str:
    """Derive a display name from the email local part (frontend only)."""
    local = (email or "").split("@", 1)[0].strip()
    words = local.replace("_", " ").replace(".", " ").split()
    return " ".join(w.capitalize() for w in words if w) or "Analyst"


def _render_entrance() -> None:
    """FT-E entrance: centered brand + two entry actions (UI only)."""
    st.markdown('<div class="fte-spacer-lg"></div>', unsafe_allow_html=True)
    l, m, r = st.columns([1, 2, 1], gap="large")
    with m:
        st.markdown(
            '<div class="fte-stage">'
            '<div class="fte-brand fte-brand-xl">FT-E</div>'
            '<div class="fte-moat">(Moat)</div>'
            "</div>",
            unsafe_allow_html=True,
        )
        st.markdown('<div class="fte-spacer-sm"></div>', unsafe_allow_html=True)
        b1, b2 = st.columns(2, gap="small")
        with b1:
            st.button(
                "Sign in", key="fte_btn_signin", use_container_width=True,
                type="primary", on_click=_fte_goto, args=("signin",),
            )
        with b2:
            st.button(
                "Create account", key="fte_btn_create", use_container_width=True,
                on_click=_fte_goto, args=("signup",),
            )
        st.markdown('<div class="fte-spacer-sm"></div>', unsafe_allow_html=True)
        st.button(
            "Try FT-E Demo →", key="fte_btn_demo", use_container_width=True,
            type="secondary", on_click=_fte_goto_demo,
        )
        st.markdown(
            '<div style="text-align:center;font-size:.78rem;color:var(--fte-muted);margin-top:.35rem">'
            "No account, no API key — explore the full workflow on a pre-analyzed sample.</div>",
            unsafe_allow_html=True,
        )


def _render_signin() -> None:
    """Sign-in page — UI only; Continue simulates successful authentication."""
    st.markdown('<div class="fte-spacer-md"></div>', unsafe_allow_html=True)
    l, m, r = st.columns([1, 2, 1], gap="large")
    with m:
        st.markdown('<div class="fte-auth-title">Welcome to</div>', unsafe_allow_html=True)
        st.markdown('<div class="fte-auth-sub">Your financial workspace starts here.</div>', unsafe_allow_html=True)
        st.markdown('<div class="fte-spacer-sm"></div>', unsafe_allow_html=True)
        st.text_input("Email", key="fte_email", placeholder="you@example.com")
        st.text_input("Password", key="fte_password", type="password", placeholder="••••••••")
        if st.button("Continue →", key="fte_btn_continue", use_container_width=True, type="primary"):
            st.session_state["fte_user_email"] = (st.session_state.get("fte_email") or "").strip()
            st.session_state["fte_user_name"] = _fte_display_name(st.session_state.get("fte_email") or "")
            st.session_state["fte_route"] = "select"
            st.rerun()
        st.markdown('<div class="fte-spacer-sm"></div>', unsafe_allow_html=True)
        f1, f2, f3 = st.columns(3, gap="small")
        with f1:
            st.button("← Back", key="fte_btn_back_signin", use_container_width=True,
                      on_click=_fte_goto, args=("entrance",))
        with f2:
            st.button("Forgot password?", key="fte_btn_forgot", use_container_width=True,
                      on_click=_fte_forgot)
        with f3:
            st.button("Create account", key="fte_btn_to_signup", use_container_width=True,
                      on_click=_fte_goto, args=("signup",))
        if st.session_state.get("fte_forgot_hint"):
            st.caption("Password recovery arrives with the full authentication phase — UI scaffolding for now.")


def _render_signup() -> None:
    """Create-account page — UI only; simulates account creation."""
    st.markdown('<div class="fte-spacer-md"></div>', unsafe_allow_html=True)
    l, m, r = st.columns([1, 2, 1], gap="large")
    with m:
        st.markdown('<div class="fte-auth-title">Create your account</div>', unsafe_allow_html=True)
        st.markdown('<div class="fte-auth-sub">Start your financial workspace.</div>', unsafe_allow_html=True)
        st.markdown('<div class="fte-spacer-sm"></div>', unsafe_allow_html=True)
        st.text_input("Name", key="fte_name", placeholder="Your name")
        st.text_input("Email", key="fte_email_signup", placeholder="you@example.com")
        st.text_input("Password", key="fte_password_signup", type="password", placeholder="••••••••")
        if st.button("Create account →", key="fte_btn_signup_submit", use_container_width=True, type="primary"):
            st.session_state["fte_user_name"] = (st.session_state.get("fte_name") or "").strip() or "Analyst"
            st.session_state["fte_user_email"] = (st.session_state.get("fte_email_signup") or "").strip()
            st.session_state["fte_route"] = "select"
            st.rerun()
        st.markdown('<div class="fte-spacer-sm"></div>', unsafe_allow_html=True)
        g1, g2 = st.columns(2, gap="small")
        with g1:
            st.button("← Back", key="fte_btn_back_signup", use_container_width=True,
                      on_click=_fte_goto, args=("entrance",))
        with g2:
            st.button("Sign in", key="fte_btn_to_signin", use_container_width=True,
                      on_click=_fte_goto, args=("signin",))


def _render_workspace_select() -> None:
    """Workspace selection: four minimal line-icon cards (frontend state only)."""
    st.markdown('<div class="fte-spacer-md"></div>', unsafe_allow_html=True)
    st.markdown(
        '<div class="fte-stage">'
        '<div class="fte-auth-title">What are you here to do?</div>'
        '<div class="fte-auth-sub">Choose your workspace</div>'
        "</div>",
        unsafe_allow_html=True,
    )
    st.markdown('<div class="fte-spacer-sm"></div>', unsafe_allow_html=True)
    workspaces = [
        ("student", _ICON_STUDENT, "Student", "Coursework, research and learning"),
        ("professional", _ICON_PROFESSIONAL, "Professional", "Analysis and decision support"),
        ("ca", _ICON_CA, "CA / Auditor", "Assurance, audit and compliance"),
        ("enterprise", _ICON_ENTERPRISE, "Enterprise", "Institutional scale workflows"),
    ]
    for row_start in (0, 2):
        c1, c2 = st.columns(2, gap="medium")
        for col, (key, icon, nm, tag) in zip((c1, c2), workspaces[row_start:row_start + 2]):
            with col:
                st.markdown(
                    f'<div class="fte-ws-card"><div class="ic">{icon}</div>'
                    f'<div class="nm">{nm}</div><div class="tg">{tag}</div></div>',
                    unsafe_allow_html=True,
                )
                st.button(
                    "Select", key=f"fte_ws_{key}", use_container_width=True,
                    on_click=_fte_set_workspace, args=(key,),
                )
    st.markdown('<div class="fte-spacer-sm"></div>', unsafe_allow_html=True)
    st.markdown(
        '<div style="text-align:center;font-size:.8rem;color:var(--fte-muted)">'
        "…or explore the product first.</div>",
        unsafe_allow_html=True,
    )
    st.button(
        "Try FT-E Demo →", key="fte_btn_demo_ws", use_container_width=True,
        type="secondary", on_click=_fte_goto_demo,
    )


def _render_account_popover() -> None:
    """Tiny account control in the workspace header (frontend state only)."""
    name = st.session_state.get("fte_user_name") or "Analyst"
    email = st.session_state.get("fte_user_email") or "user@example.com"
    ws_label = {
        "student": "Student", "professional": "Professional",
        "ca": "CA / Auditor", "enterprise": "Enterprise",
    }.get(st.session_state.get("fte_workspace"), "—")
    with st.popover("Account", use_container_width=True):
        st.markdown(f"**{html.escape(str(name))}**")
        st.caption(html.escape(str(email)))
        st.caption(f"Workspace · {ws_label}")
        if st.button("Sign out", key="fte_btn_signout", use_container_width=True):
            st.session_state["fte_route"] = "entrance"
            st.session_state["fte_workspace"] = None
            st.session_state["fte_user_name"] = ""
            st.session_state["fte_user_email"] = ""
            st.rerun()


def _render_workspace_shell(uploaded_files) -> None:
    """FT-E workspace: shared shell with three separate page views."""
    extraction_results, combined_raw_text, document_summaries = _run_ingestion(uploaded_files)
    module3_result = {}
    try:
        module3_result = _terminal_module3(combined_raw_text, extraction_results)
    except Exception as e:
        st.warning("⚠️ Financial intelligence temporarily unavailable — verified data tools remain below.")
        st.session_state["pipeline_debug_log"] = [{"stage": "module3", "length": 0, "error_marker_detected": True, "preview": str(e)[:300]}]
    provider_health = get_provider_health()

    # Header chrome: brand | gateway | account.
    h_l, h_r = st.columns([3, 2], gap="medium", vertical_alignment="center")
    with h_l:
        st.markdown(
            '<div class="fte-ws-header-brand">FT-E<span>Financial Timeline Engine</span></div>',
            unsafe_allow_html=True,
        )
    with h_r:
        g_c, a_c = st.columns([2, 1], gap="small", vertical_alignment="center")
        with g_c:
            render_gateway_pill()
        with a_c:
            _render_account_popover()

    # Persistent workspace navigation — three separate page views, never a
    # single scrolling page. Page choice lives in fte_page (session state),
    # so switching pages preserves grid selection, filters, chat and memo.
    st.segmented_control(
        "Workspace",
        options=["Financial Grid", "Intelligence", "System"],
        key="fte_page",
        label_visibility="collapsed",
    )
    page = st.session_state["fte_page"]

    if page == "Financial Grid":
        rail_col, center_col = st.columns([1, 3.2], gap="medium")
        with rail_col:
            _render_source_rail(uploaded_files, extraction_results, module3_result)
        with center_col:
            _render_financial_grid(module3_result)
        _render_provenance_tray(module3_result)
    elif page == "Intelligence":
        if st.session_state.get("fte_memo_view_open"):
            _render_memo_focused_view(document_summaries, module3_result)
        else:
            _render_co_pilot(extraction_results, provider_health)
            st.markdown("---")
            _render_selected_metric_analysis(module3_result)
            st.markdown("---")
            _render_memo_generator(document_summaries)
    else:
        _render_system_tab()


def _selected_metric_card_html(fields, explainer) -> str:
    """Compact Intelligence 'Selected Metric' card — name + status, value,
    a short verified explanation, and an evidence line. No big panels, no
    grid duplication, no fabricated provenance ('—' for missing fields)."""
    parts = [
        f'<div style="display:flex;align-items:baseline;justify-content:space-between;gap:10px">'
        f'<span style="font-weight:650;font-size:1.05rem;color:var(--fte-text)">{html.escape(str(fields["label"]))}</span>'
        f'<span style="font-size:.85rem;white-space:nowrap">{html.escape(str(fields["status"]))}</span>'
        f'</div>',
        f'<div style="font-size:1.5rem;font-weight:700;color:var(--fte-text);margin:.5rem 0 .35rem;letter-spacing:.01em">{html.escape(str(fields["value"]))}</div>',
    ]
    if explainer:
        parts.append(
            f'<div style="margin-top:.4rem;font-size:.86rem;color:var(--fte-text);line-height:1.5">'
            f'<span style="color:var(--fte-muted)">What this tells you — </span>{html.escape(explainer)}</div>'
        )
    ev = []
    for label, v in (("Origin", fields["origin"]), ("Source", fields["source"]),
                     ("Period", fields["period"]), ("Evidence", fields["evidence"])):
        if v not in (None, "", "—"):
            ev.append(f"{label}: {v}")
    if ev:
        parts.append(
            f'<div style="margin-top:.6rem;padding-top:.5rem;border-top:1px solid var(--fte-border);font-size:.8rem;color:var(--fte-muted);line-height:1.55">'
            + " · ".join(html.escape(e) for e in ev[:3]) + "</div>"
        )
    if fields.get("note"):
        parts.append(
            f'<div style="margin-top:.6rem;padding:.4rem .55rem;font-size:.78rem;color:var(--fte-blocked);border-left:2px solid var(--fte-blocked);background:rgba(224,82,82,.06);border-radius:4px">⚠️ Analysis limited — {html.escape(str(fields["note"]))}</div>'
        )
    return "".join(parts)


def _render_selected_metric_analysis(module3_result) -> None:
    """Selected Metric Analysis: a compact card for the grid selection only.
    No metric selected → a clean empty state (never a large analysis panel)."""
    st.markdown('<div class="fte-rail-title">🔬 Selected Metric Analysis</div>', unsafe_allow_html=True)
    metric = st.session_state.get("fte_selected_metric")
    rows = st.session_state.get("fte_grid_rows") or _build_terminal_rows(module3_result)
    if not metric or not any(r["metric"] == metric for r in rows):
        st.markdown(
            '<div class="fte-intel-card"><div style="font-size:.86rem;color:var(--fte-muted);line-height:1.55">'
            'Select a metric in <span style="color:var(--fte-text)">Financial Grid</span> to analyze it here.</div></div>',
            unsafe_allow_html=True,
        )
        return
    fields = _metric_overlay_fields(rows, module3_result, metric)
    explainer = _metric_explainer(metric)
    st.markdown(
        '<div class="fte-intel-card">' + _selected_metric_card_html(fields, explainer) + '</div>',
        unsafe_allow_html=True,
    )
    if st.button("View metric detail", key="fte_intel_view_detail", use_container_width=True):
        _metric_detail_dialog(module3_result)


def _render_memo_generator(document_summaries) -> None:
    """Generate Memo: ONE primary action; compact status; the completed
    memo opens in the focused document view (never inline on this page)."""
    st.markdown('<div class="fte-rail-title">📝 Generate Memo</div>', unsafe_allow_html=True)
    memo = st.session_state.get("fte_memo_draft") or ""
    if not document_summaries:
        st.markdown(
            '<div class="fte-intel-card"><div style="font-size:.86rem;color:var(--fte-muted);line-height:1.55">'
            'Upload a financial document to generate an investment memo.</div></div>',
            unsafe_allow_html=True,
        )
        return
    if st.button("Generate Memo", key="fte_btn_memo", use_container_width=True):
        st.session_state["fte_memo_status"] = "generating"
        with st.spinner("Merging document summaries..."):
            merge_cache_key = _hash_text(
                "||".join(f"{d['file_name']}:{d['summary']}" for d in document_summaries)
            )
            master_summary, _ = _cached_call(
                "merge_cache", merge_cache_key,
                lambda: merge_document_summaries(document_summaries),
            )
        if not master_summary or not master_summary.strip() or contains_error_marker(master_summary):
            st.error("❌ Document summarization failed — no real content to analyze.")
            st.session_state["fte_memo_draft"] = ""
            st.session_state["fte_memo_status"] = "failed"
        else:
            with st.spinner("Drafting memo..."):
                memo_system_prompt = (
                    "You are an elite institutional investment research analyst. "
                    "Write the investment memo strictly and exclusively from the "
                    "facts, figures, dates, and events contained in the Document "
                    "Summary provided below. Do not produce generic, templated, or "
                    "boilerplate analysis -- every claim must be traceable to a "
                    "specific fact stated in the Document Summary. If the Document "
                    "Summary lacks information for a requested section, explicitly "
                    "state that the source documents did not provide it rather than "
                    "inventing generic filler."
                )
                prompt = f"""Analyze the Document Summary below carefully. Extract key event milestones, timelines, and potential controversy flags SPECIFIC to this Document Summary. Write a comprehensive multi-paragraph investment memo that identifies, using only facts from the Document Summary:
1. Key financial events and dates
2. Market movements and impacts
3. Risk factors and opportunities
4. Strategic implications

Document Summary:
{master_summary}

Generate a professional investment memo grounded strictly in the Document Summary above. Do not generate generic industry commentary that is not tied to a specific fact in the Document Summary."""
                memo_cache_key = _hash_text(master_summary)
                memo, _ = _cached_call(
                    "memo_cache", memo_cache_key,
                    lambda: call_ai_with_fallback(prompt, system_prompt=memo_system_prompt, temperature=0.3),
                )
                st.session_state["fte_memo_draft"] = memo
                st.session_state["fte_memo_status"] = "ready"
                st.session_state["fte_memo_view_open"] = True
                st.rerun()
    # Compact generation status — honest, one line. The memo itself lives
    # in the focused document view, not inline on the Intelligence hub.
    status = st.session_state.get("fte_memo_status", "idle")
    memo = st.session_state.get("fte_memo_draft") or ""
    if status == "ready" and memo:
        st.caption("✅ Memo complete.")
    elif status == "failed":
        st.caption("Memo generation did not complete — no content was produced.")
    elif status == "generating":
        st.caption("Generating memo — completed material appears as soon as the pipeline produces it.")
    if memo and not st.session_state.get("fte_memo_view_open"):
        if st.button("View memo", key="fte_btn_view_memo", use_container_width=True):
            st.session_state["fte_memo_view_open"] = True
            st.rerun()


# Sprint 3.2.1: bidirectional memo-metric component — the memo body
# renders as ONE continuous HTML document inside the component (true
# inline text flow, no boxed buttons); metric tokens are inline spans
# that report clicks back to Python. Falls back to a plain document if
# the component cannot be declared.
#
# Deployment note (Sprint 3.2.3): the component frontend is served from a
# local directory, so the path MUST be resolved relative to this file
# (__file__) — never relative to the process working directory — because a
# hosted deployment may run Streamlit from a different CWD than the repo
# root. We additionally VERIFY that index.html actually exists at the
# resolved path before registering; if it is missing we degrade gracefully
# to a plain (non-interactive) memo instead of rendering a broken
# component iframe.
#
# Sprint 3.2.4 — TEMPORARILY DISABLED (deployment reliability): the hosted
# Streamlit environment does not serve the local custom-component frontend
# (/component/* static routes), so the memo iframe never became ready and
# the app showed "trouble loading the ... component" on every rerun of the
# memo view. The memo is now rendered with Streamlit-native inline links
# that open the SAME native metric dialog (see _render_memo_focused_view /
# _memo_metric_html). Set _MEMO_COMPONENT_ENABLED = True only once the
# hosting serves local custom-component assets reliably; the path
# resolution below stays relative to this file (__file__), never the
# process working directory.
_MEMO_COMPONENT_ENABLED = False
if _MEMO_COMPONENT_ENABLED:
    try:
        _MEMO_COMPONENT_DIR = None
        _APP_DIR = os.path.dirname(os.path.realpath(__file__))
        for _cand in (
            os.path.join(_APP_DIR, "memo_component"),
            os.path.join(os.path.dirname(os.path.abspath(__file__)), "memo_component"),
        ):
            if os.path.isfile(os.path.join(_cand, "index.html")):
                _MEMO_COMPONENT_DIR = _cand
                break
        if _MEMO_COMPONENT_DIR is not None:
            _memo_metric_component = st.components.v1.declare_component(
                "fte_memo_metric", path=_MEMO_COMPONENT_DIR
            )
        else:
            _memo_metric_component = None
    except Exception:
        _memo_metric_component = None
else:
    _memo_metric_component = None


def _memo_clickable_spans(rows, para) -> list:
    """(start, end, label, metric) spans in a memo paragraph that directly
    represent a pipeline metric: its canonical name, or its exact stored
    value string when that value is unambiguous. Whole-word, case-
    insensitive; the grid rows are the only source of truth — nothing is
    invented, and ordinary prose is never made clickable."""
    spans = []
    if not para:
        return spans
    for r in rows or []:
        name = str(r.get("metric") or r.get("Metric") or "")
        if not name:
            continue
        for m in re.finditer(rf"(?<!\w){re.escape(name)}(?!\w)", para, flags=re.IGNORECASE):
            spans.append((m.start(), m.end(), m.group(0), name))
    # Values: only exact, non-trivial, unambiguous stored value strings.
    value_map = {}
    for r in rows or []:
        v = str(r.get("Value") or "").strip()
        metric = str(r.get("metric") or r.get("Metric") or "")
        if len(v) >= 2 and v != "—" and metric:
            value_map.setdefault(v, []).append(metric)
    for v, names in value_map.items():
        if len(set(names)) != 1:
            continue  # ambiguous — never guess which metric a value belongs to
        for m in re.finditer(rf"(?<!\w){re.escape(v)}(?!\w)", para, flags=re.IGNORECASE):
            spans.append((m.start(), m.end(), m.group(0), names[0]))
    # Order by position; drop overlapping spans (earlier/longer wins).
    spans.sort(key=lambda t: (t[0], -len(t[2])))
    out, last_end = [], 0
    for s, e, label, metric in spans:
        if s < last_end:
            continue
        out.append((s, e, label, metric))
        last_end = e
    return out


def _memo_metric_html(rows, memo) -> str:
    """Memo body as ONE continuous HTML document. Pipeline metric names
    and unambiguous exact stored values become INLINE NATIVE LINKS
    (?fte_metric=<name> query param); ordinary prose is never wrapped, so
    sentences stay intact. No custom component is required — clicking a
    link reruns the app and opens the existing native metric dialog."""
    paragraphs = [p.strip() for p in memo.replace("\r\n", "\n").split("\n\n") if p.strip()]
    out = []
    for para in paragraphs:
        spans = _memo_clickable_spans(rows, para)
        if not spans:
            out.append(f'<div class="fte-memo-para">{html.escape(para).replace(chr(10), "<br>")}</div>')
            continue
        body, pos = [], 0
        for s, e, label, metric in spans:
            if s > pos:
                body.append(html.escape(para[pos:s]).replace(chr(10), "<br>"))
            # Demo Mode links carry a demo-only marker so a full page
            # navigation can be reconstructed as the demo session instead of
            # falling back to the Entrance default. Real (non-demo) memo
            # links never include it.
            demo_suffix = "&fte_demo=1" if st.session_state.get("fte_demo_mode") else ""
            body.append(
                f'<a class="fte-metric-link" href="?fte_metric={urllib.parse.quote(metric)}{demo_suffix}" '
                f'title="Inspect {html.escape(metric, quote=True)}">{html.escape(label)}</a>'
            )
            pos = e
        if pos < len(para):
            body.append(html.escape(para[pos:]).replace(chr(10), "<br>"))
        out.append(f'<div class="fte-memo-para">{"".join(body)}</div>')
    return "".join(out)


def _clear_memo_metric_click() -> None:
    """Close the memo metric card (also wired to × / ESC / backdrop)."""
    st.session_state["fte_memo_metric_click"] = None


@st.dialog("Metric Detail", width="small", dismissible=True, on_dismiss=_clear_memo_metric_click)
def _memo_metric_dialog(module3_result) -> None:
    """Floating evidence card for a metric clicked inside the memo. Same
    compact card as the grid overlay; pipeline records only — no AI, no
    API key required. Blocked metrics show the honest limitation."""
    metric = st.session_state.get("fte_memo_metric_click")
    if not metric:
        st.caption("No metric selected.")
        return
    rows = st.session_state.get("fte_grid_rows") or _build_terminal_rows(module3_result)
    if not any(r["metric"] == metric for r in rows):
        st.caption("This metric is not available from the current pipeline.")
        return
    fields = _metric_overlay_fields(rows, module3_result, metric)
    explainer = _metric_explainer(metric)
    st.markdown(_metric_detail_html(fields, explainer), unsafe_allow_html=True)
    st.button("Close", key="fte_memo_metric_close", on_click=_clear_memo_metric_click)


def _render_memo_focused_view(document_summaries, module3_result) -> None:
    """Focused document-style memo view: back control, memo title, company /
    document context, the memo body, and export. Pipeline metrics mentioned
    in the memo are highlighted and clickable — clicking opens ONE floating
    evidence card ON TOP of the memo (native dialog); the document itself
    is never resized or pushed down. Keeps the memo out of the Intelligence
    hub; uses whatever the existing pipeline produced — never invented."""
    if st.button("← Back to Intelligence", key="fte_btn_memo_back"):
        st.session_state["fte_memo_view_open"] = False
        st.session_state["fte_memo_metric_click"] = None
        st.rerun()
    memo = st.session_state.get("fte_memo_draft") or ""
    if not memo:
        st.markdown(
            '<div class="fte-memo-doc"><div style="color:var(--fte-muted);font-size:.9rem">'
            'No memo has been generated yet.</div></div>',
            unsafe_allow_html=True,
        )
        return
    context = " · ".join(
        html.escape(str(d.get("file_name", ""))) for d in (document_summaries or []) if d.get("file_name")
    )[:180]
    rows = st.session_state.get("fte_grid_rows") or _build_terminal_rows(module3_result)
    memo_html = _memo_metric_html(rows, memo)
    with st.container(key="fte_memo_doc", border=True):
        st.markdown('<div class="fte-memo-title">Investment Memo</div>', unsafe_allow_html=True)
        st.markdown(f'<div class="fte-memo-context">📄 {context or "No document context"}</div>', unsafe_allow_html=True)
        st.markdown('<hr style="border:none;border-top:1px solid var(--fte-border);margin:1.1rem 0">', unsafe_allow_html=True)
        # The memo body is ONE continuous HTML document (true inline
        # flow). Metric tokens are inline NATIVE links (still plain
        # inline text — subtle color + dotted underline). Clicking one
        # navigates to ?fte_metric=<name>; the rerun opens the SAME
        # floating native metric dialog. No custom component, no iframe,
        # no external frontend — reliable on any Streamlit hosting.
        st.markdown(memo_html, unsafe_allow_html=True)
    # Metric clicks arrive via the URL query param (?fte_metric=...)
    # instead of a custom component; consume it and open the dialog.
    qmetric = st.query_params.get("fte_metric")
    if qmetric:
        st.session_state["fte_memo_metric_click"] = str(qmetric)
        del st.query_params["fte_metric"]
    if st.session_state.get("fte_memo_metric_click"):
        _memo_metric_dialog(module3_result)
    dl1, dl2 = st.columns(2)
    with dl1:
        st.download_button(
            "📥 Download as Word Document",
            data=generate_docx_download(memo),
            file_name="FT_E_Investment_Memo.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )
    with dl2:
        st.download_button(
            "📄 Download PDF",
            data=generate_pdf_download(memo),
            file_name="FT_E_Investment_Memo.pdf",
            mime="application/pdf",
            use_container_width=True,
        )


# =============================================================================
# SECTION 10e: Static Demo Mode (Sprint 4)
# =============================================================================
# Deployment-safe demo: ONE pre-analyzed sample (Microsoft FY2025 10-K) is
# fed through the EXISTING Financial Grid / Intelligence / memo components.
# Zero AI calls, zero ingestion, zero computation, zero uploads, zero user
# storage — deterministic and effectively zero-cost per visitor. Every value
# is explicitly labelled "Demo / Pre-analyzed sample" and is never presented
# as live financial information. The real pipeline is completely untouched.
_FTE_DEMO_LABEL = "Demo / Pre-analyzed sample"
_FTE_DEMO_COMPANY = "Microsoft Corporation"
_FTE_DEMO_TICKER = "MSFT"
_FTE_DEMO_DOC_NAME = "Microsoft Corp — Form 10-K, FY2025 (Demo / Pre-analyzed sample)"
_FTE_DEMO_SUMMARY = (
    "Microsoft Corporation (NASDAQ: MSFT) reported fiscal-year 2025 revenue of $281.70B, "
    "operating income of $125.50B and net income of $98.30B (diluted EPS $13.05). Operating "
    "cash flow was $127.80B. The balance sheet closed with $512.20B of assets, $96.60B of "
    "total debt and $268.50B of shareholders' equity. Cloud and AI services continued to "
    "anchor growth, with Azure driving a meaningful share of the revenue increase."
)


def _demo_fact(value, source, period="FY2025", **extra):
    """One static demo fact, shaped exactly like a module3 pipeline fact so
    every existing presentation component consumes it unchanged."""
    fact = {"value": value, "source": source, "reporting_period": period, "unit": "USD"}
    fact.update(extra)
    return fact


def _demo_module3_result() -> dict:
    """The complete static demo sample as a module3-shaped result dict.
    Pure function — no computation, no I/O, no AI. Values are internally
    consistent (ratios derive from the stated figures)."""
    fd = {
        "Revenue": _demo_fact(281700000000, "10-K FY2025 · Income Statement",
                              evidence="Consolidated Statements of Income, p. 26", page="26", scale="B"),
        "Net Profit": _demo_fact(98300000000, "10-K FY2025 · Income Statement",
                                 evidence="Consolidated Statements of Income, p. 26", page="26", scale="B"),
        "Operating Profit": _demo_fact(125500000000, "10-K FY2025 · Income Statement",
                                       evidence="Consolidated Statements of Income, p. 26", page="26", scale="B"),
        "EPS": _demo_fact(13.05, "10-K FY2025 · Income Statement",
                          evidence="Diluted earnings per share, Consolidated Statements of Income, p. 26", page="26"),
        "Debt": _demo_fact(96600000000, "10-K FY2025 · Balance Sheet",
                           evidence="Consolidated Balance Sheets, p. 27", page="27", scale="B"),
        "Assets": _demo_fact(512200000000, "10-K FY2025 · Balance Sheet",
                             evidence="Consolidated Balance Sheets, p. 27", page="27", scale="B"),
        "Liabilities": _demo_fact(243700000000, "10-K FY2025 · Balance Sheet",
                                  evidence="Consolidated Balance Sheets, p. 27", page="27", scale="B"),
        "Equity": _demo_fact(268500000000, "10-K FY2025 · Balance Sheet",
                             evidence="Consolidated Balance Sheets, p. 27", page="27", scale="B"),
        "Cash Flow": _demo_fact(127800000000, "10-K FY2025 · Cash Flow Statement",
                                evidence="Consolidated Statements of Cash Flows, p. 28", page="28", scale="B"),
    }
    rt = {
        "EBITDA": _demo_fact(161000000000, "Calculated",
                             evidence="Operating income + depreciation & amortisation (10-K FY2025)", scale="B"),
        "Profit Margin": _demo_fact(0.349, "Calculated",
                                    evidence="Net income ÷ revenue (10-K FY2025)"),
        "ROE": _demo_fact(0.366, "Calculated",
                          evidence="Net income ÷ shareholders' equity (10-K FY2025)"),
        "ROA": _demo_fact(0.192, "Calculated",
                          evidence="Net income ÷ total assets (10-K FY2025)"),
        "Debt to Equity": _demo_fact(0.36, "Calculated",
                                     evidence="Total debt ÷ shareholders' equity (10-K FY2025)"),
        "Current Ratio": _demo_fact(1.40, "Calculated",
                                    evidence="Current assets ÷ current liabilities (10-K FY2025)"),
    }
    return {
        "financial_data": fd,
        "ratios": rt,
        "missing_data": {"financial_data": ["Segment Gross Margin"], "ratios": []},
        "demo": True,
    }


def _demo_extraction_results() -> list:
    """Demo source rail entry (presentation only — no file is processed)."""
    return [{"file_name": _FTE_DEMO_DOC_NAME, "cached": True}]


def _demo_document_summaries() -> list:
    """Demo document summaries — static text only; memo context line."""
    return [{"file_name": _FTE_DEMO_DOC_NAME, "summary": _FTE_DEMO_SUMMARY}]


class _DemoFileStub:
    """Minimal stand-in so the source rail renders the demo source name."""
    name = _FTE_DEMO_DOC_NAME


_FTE_DEMO_MEMO = """EXECUTIVE SUMMARY
Microsoft Corporation closed fiscal 2025 with Revenue of $281.70B, lifted by continued cloud and AI demand. Net Profit reached $98.30B (diluted EPS of $13.05), and Operating Cash Flow of $127.80B continued to fund buybacks, dividends and data-centre investment. The balance sheet stays conservative: Debt of $96.60B against Equity of $268.50B keeps Debt to Equity at 0.36.

KEY FINANCIAL EVENTS
Revenue growth was led by cloud and AI services. Operating Profit of $125.50B supported EBITDA of $161.00B. The company added substantial capital-spending capacity while preserving a Current Ratio of 1.40 and ROA of 0.19.

FINANCIAL PERFORMANCE
Profitability remained strong: ROE of 0.37 and Profit Margin of 0.35 reflect efficient conversion of $281.70B of revenue into $98.30B of net income. Assets of $512.20B are balanced by Liabilities of $243.70B, leaving Equity of $268.50B. Cash generation is the standout: Operating Cash Flow of $127.80B comfortably exceeds Net Profit.

RISKS & OPPORTUNITIES
Concentration in AI-infrastructure spending is the primary watch item: capital intensity is rising even as Revenue growth accelerates. Segment Gross Margin is not captured in the demo sample, so margin mix across cloud, PC and productivity segments is unanalyzed. The opportunity: continued monetisation of AI services through the installed enterprise base.

STRATEGIC IMPLICATIONS
The model is shifting toward AI-linked recurring revenue, with cloud economics improving as scale grows. The strong balance sheet (Debt to Equity of 0.36, Current Ratio of 1.40) gives management financial flexibility for organic capex and selective M&A.

RECOMMENDATIONS
Monitor quarterly operating leverage and Azure growth. Track capital intensity against Operating Cash Flow of $127.80B. Confirm segment margin disclosures before extrapolating the profit mix beyond the verified company-level figures.

— Demo memo · Pre-analyzed sample · No AI generation used —"""


def _fte_goto_demo() -> None:
    """Enter Static Demo Mode (frontend routing state only)."""
    st.session_state["fte_route"] = "demo"
    st.session_state["fte_demo_mode"] = True
    st.session_state["fte_page"] = "Financial Grid"
    st.session_state["fte_memo_view_open"] = False
    st.session_state["fte_memo_metric_click"] = None


def _reconstruct_demo_from_query() -> None:
    """Reconstruct Demo-Mode session state when a demo metric click arrives
    via the URL (?fte_demo=1&fte_metric=...) after a full page navigation.

    Hosted/preview environments do not always restore session state on a
    full reload, so fte_route would otherwise fall back to the 'entrance'
    default before main() routes — sending the user to the start screen.
    This runs BEFORE any routing: it rebuilds the demo route, memo view and
    static memo, transfers fte_metric into fte_memo_metric_click (the ONLY
    role of that param), then consumes both params. It never touches real
    (non-demo) sessions because their memo links never carry fte_demo=1.
    """
    if not st.query_params.get("fte_demo"):
        return
    st.session_state["fte_route"] = "demo"
    st.session_state["fte_demo_mode"] = True
    st.session_state["fte_page"] = "Intelligence"
    st.session_state["fte_memo_view_open"] = True
    st.session_state["fte_memo_draft"] = _FTE_DEMO_MEMO
    st.session_state["fte_memo_status"] = "ready"
    qm = st.query_params.get("fte_metric")
    if qm:
        st.session_state["fte_memo_metric_click"] = str(qm)
    # Consume the URL trigger after transferring it — never reset the route.
    for _p in ("fte_demo", "fte_metric"):
        if _p in st.query_params:
            del st.query_params[_p]


def _demo_copilot_answer(question, rows):
    """Deterministic Demo-Mode Co-Pilot: predefined, evidence-backed answers
    built ONLY from the static demo dataset. Never calls an AI provider."""
    q = (question or "").lower()
    has = lambda *ks: any(k in q for k in ks)

    if has("strongest verified", "strongest evidence", "best evidence", "what is verified"):
        content = (
            "The strongest verified evidence in the demo sample is the income statement: "
            "Revenue of $281.70B, Net Profit of $98.30B and diluted EPS of $13.05, each "
            "drawn from the FY2025 10-K (Consolidated Statements of Income, p. 26). The three "
            "figures are mutually consistent — $98.30B of net income on $281.70B of revenue "
            "implies a Profit Margin of 0.35. Operating Cash Flow of $127.80B adds a cash "
            "quality check on those earnings."
        )
        intent = "verified_evidence"
    elif has("risk", "threat", "concern", "downside"):
        content = (
            "Key risks in the demo sample: (1) AI-infrastructure capital intensity — spending "
            "is rising even as Revenue of $281.70B accelerates; (2) segment mix is unanalyzed — "
            "Segment Gross Margin is blocked because segment disclosures sit outside the "
            "extracted sample; (3) Debt of $96.60B against Equity of $268.50B (Debt to Equity "
            "0.36) is manageable today, but leverage deserves watching if buyback-funded."
        )
        intent = "risks"
    elif has("investigate", "what next", "next step", "watch", "monitor"):
        content = (
            "What to investigate next: (1) segment-level revenue and gross margin — the sample "
            "lacks segment disclosures (Segment Gross Margin is blocked); (2) cash conversion — "
            "compare Operating Cash Flow of $127.80B with Net Profit of $98.30B across quarters; "
            "(3) whether the Current Ratio of 1.40 and ROA of 0.19 hold as capital intensity rises."
        )
        intent = "investigate_next"
    else:
        metric = st.session_state.get("fte_selected_metric")
        if metric and any(r.get("metric") == metric for r in (rows or [])):
            fields = _metric_overlay_fields(rows, {}, metric)
            expl = _metric_explainer(metric)
            content = (
                f"Selected metric — {fields['label']}: {fields['value']} "
                f"(status {fields['status']}). {expl or 'See the evidence card for context.'} "
                f"Evidence: {fields['evidence']}."
            )
            intent = "selected_metric"
        else:
            content = (
                "I can answer a limited set of questions in Demo Mode. Try asking about the "
                "selected metric, strongest verified evidence, key risks, or what to "
                "investigate next."
            )
            intent = "demo_limited"
    return content, {"intent": intent, "demo": True, "demo_label": _FTE_DEMO_LABEL}


def _demo_submit_question(question) -> None:
    """One demo Co-Pilot turn — deterministic, no AI provider."""
    question = (question or "").strip()
    if not question:
        return
    rows = st.session_state.get("fte_grid_rows") or _build_terminal_rows(_demo_module3_result())
    content, metadata = _demo_copilot_answer(question, rows)
    st.session_state.setdefault("fte_demo_chat_messages", []).append(
        {"role": "user", "content": question}
    )
    st.session_state.setdefault("fte_demo_chat_messages", []).append(
        {"role": "assistant", "content": content, "metadata": metadata}
    )


def _render_demo_co_pilot() -> None:
    """Demo Ask Co-Pilot — deterministic predefined answers only."""
    st.markdown("---")
    st.markdown('<div class="fte-rail-title">💬 Ask Co-Pilot</div>', unsafe_allow_html=True)
    st.markdown(
        f'<div class="fte-cap">📄 Active context: {html.escape(_FTE_DEMO_DOC_NAME)}</div>',
        unsafe_allow_html=True,
    )
    st.markdown(
        f'<div class="fte-cap">🎯 Demo Mode · deterministic answers · no AI provider · '
        f'{html.escape(_FTE_DEMO_LABEL)}</div>',
        unsafe_allow_html=True,
    )
    messages = st.session_state.get("fte_demo_chat_messages") or []
    for msg in messages[-4:]:
        with st.chat_message(msg["role"]):
            st.markdown(msg["content"])
    if not messages:
        st.caption("Suggested (Demo):")
        cols = st.columns(2)
        prompts = [
            "Summarize the strongest verified evidence.",
            "What are the key risks?",
            "What should I investigate next?",
            "Explain the selected metric.",
        ]
        for i, prompt in enumerate(prompts):
            with cols[i % 2]:
                if st.button(prompt, key=f"fte_demo_suggest_{i}", use_container_width=True):
                    _demo_submit_question(prompt)
                    st.rerun()
    user_input = st.chat_input("Ask about the demo evidence…")
    if user_input:
        _demo_submit_question(user_input)
        st.rerun()


def _render_demo_memo_generator() -> None:
    """Demo Generate Memo — opens the SAME focused memo view with the static
    pre-analyzed memo. No AI pipeline, no document processing, no cost."""
    st.markdown('<div class="fte-rail-title">📝 Generate Memo</div>', unsafe_allow_html=True)
    if st.button("Generate Demo Memo", key="fte_btn_demo_memo", use_container_width=True):
        st.session_state["fte_memo_draft"] = _FTE_DEMO_MEMO
        st.session_state["fte_memo_status"] = "ready"
        st.session_state["fte_memo_view_open"] = True
        st.rerun()
    st.caption("Demo Memo · Pre-analyzed sample · No AI generation used")
    if st.session_state.get("fte_memo_status") == "ready" and st.session_state.get("fte_memo_view_open"):
        st.caption("✓ Demo memo ready")


def _render_demo_workspace() -> None:
    """Demo workspace: the same shell and pages as the real workspace, fed
    ONLY the static pre-analyzed sample. No uploads, no ingestion, no AI,
    no user storage — every presentation component is the existing one."""
    st.session_state["fte_demo_mode"] = True
    if st.query_params.get("fte_exit_demo"):
        st.session_state["fte_demo_mode"] = False
        st.session_state["fte_route"] = "entrance"
        st.session_state["fte_workspace"] = None
        del st.query_params["fte_exit_demo"]
        st.rerun()
        return

    module3_result = _demo_module3_result()
    extraction_results = _demo_extraction_results()
    document_summaries = _demo_document_summaries()

    h_l, h_r = st.columns([3, 2], gap="medium", vertical_alignment="center")
    with h_l:
        st.markdown(
            '<div class="fte-ws-header-brand">FT-E<span>Financial Timeline Engine</span></div>',
            unsafe_allow_html=True,
        )
    with h_r:
        g_c, a_c = st.columns([2, 1], gap="small", vertical_alignment="center")
        with g_c:
            st.markdown(
                '<span class="fte-pill"><span class="fte-dot warn"></span>DEMO MODE</span>',
                unsafe_allow_html=True,
            )
        with a_c:
            _render_account_popover()

    st.markdown(
        f'<div class="fte-demo-note">Demo / Pre-analyzed sample — static illustrative data, '
        f"not live financial information. "
        f'<a href="?fte_exit_demo=1">Create account / Sign in to analyze your own documents →</a></div>',
        unsafe_allow_html=True,
    )

    st.segmented_control(
        "Workspace",
        options=["Financial Grid", "Intelligence", "System"],
        key="fte_page",
        label_visibility="collapsed",
    )
    page = st.session_state["fte_page"]

    if page == "Financial Grid":
        rail_col, center_col = st.columns([1, 3.2], gap="medium")
        with rail_col:
            _render_source_rail([_DemoFileStub()], extraction_results, module3_result)
        with center_col:
            _render_financial_grid(module3_result)
        _render_provenance_tray(module3_result)
    elif page == "Intelligence":
        if st.session_state.get("fte_memo_view_open"):
            _render_memo_focused_view(document_summaries, module3_result)
        else:
            _render_demo_co_pilot()
            st.markdown("---")
            _render_selected_metric_analysis(module3_result)
            st.markdown("---")
            _render_demo_memo_generator()
    else:
        _render_system_tab()


# =============================================================================
# SECTION 11: Main App / UI
# =============================================================================
def main():
    _reconstruct_demo_from_query()
    _inject_terminal_css()
    _init_terminal_state()

    # FT-E entrance flow (UI only — no real auth/storage in this phase).
    route = st.session_state.get("fte_route", "entrance")
    if route == "entrance":
        _render_entrance()
        return
    if route == "signin":
        _render_signin()
        return
    if route == "signup":
        _render_signup()
        return
    if route == "select":
        _render_workspace_select()
        return
    if route == "demo":
        _render_demo_workspace()
        return

    # Workspace shell: shared sidebar (legacy view switch + ingestion).
    fte_view = st.sidebar.radio(
        "Workspace",
        ["Institutional Terminal", "Classic Dashboard"],
        index=0,
        key="fte_view_radio",
    )
    uploaded_files = st.sidebar.file_uploader(
        "Upload Financial Documents (.txt, .csv, .xlsx, .docx, .pdf)",
        type=["txt", "csv", "xlsx", "docx", "pdf"],
        accept_multiple_files=True
    )

    if fte_view == "Classic Dashboard":
        _render_classic_dashboard(uploaded_files)
    else:
        _render_workspace_shell(uploaded_files)


def _render_classic_dashboard(uploaded_files):
    # --- AI status (canonical provider health via backend/gateway) ---
    # Status states: 🔴 no eligible provider / 🟢 live generation verified /
    # 🟡 configured but no live call yet. Per-provider detail comes from the
    # canonical ProviderManager (available / configured_unavailable /
    # not_configured), never from a bare key presence check.
    health = get_provider_health()
    has_any_key = any(health.values())
    if not has_any_key:
        st.error(f"🔴 AI Status: Offline — {NO_ELIGIBLE_PROVIDER_MESSAGE}")
    elif st.session_state["ai_connected"]:
        provider = st.session_state.get("ai_provider_used", "AI Provider")
        st.success(f"🟢 AI Status: Live ({provider})")
    else:
        st.info("🟡 AI Status: Provider(s) configured — awaiting first live generation")

    with st.expander("🩺 Provider Health & Activity Log", expanded=False):
        provider_display = {
            "google": "Google AI Studio", "groq": "Groq", "openrouter": "OpenRouter",
            "nvidia": "NVIDIA", "rapidapi": "RapidAPI", "sambanova": "SambaNova",
            "github": "GitHub Models", "cerebras": "Cerebras", "cohere": "Cohere",
        }
        status_icon = {"available": "🟢", "configured_unavailable": "🟡", "not_configured": "⚪"}
        canonical_status = get_canonical_provider_status()
        health_rows = []
        for slug, state in canonical_status.items():
            health_rows.append({
                "Provider": provider_display.get(slug, slug.replace("_", " ").title()),
                "Status": f"{status_icon.get(state, '⚪')} {state.replace('_', ' ')}",
            })
        # Fall back to key-presence rows only if the gateway is unavailable.
        if not health_rows:
            health_rows = [
                {"Provider": k, "Status": ("🟢 available" if v else "⚪ not configured")}
                for k, v in health.items()
            ]
        health_df = pd.DataFrame(health_rows)
        st.dataframe(health_df, use_container_width=True, hide_index=True)
        if st.session_state["provider_log"]:
            st.dataframe(pd.DataFrame(st.session_state["provider_log"][-20:]), use_container_width=True, hide_index=True)
        else:
            st.caption("No AI provider activity yet.")

    with st.expander("🔮 Roadmap (Planned Modules)", expanded=False):
        roadmap = get_future_module_status()
        st.dataframe(
            pd.DataFrame([{"Module": v["name"], "Status": v["status"].title()} for v in roadmap.values()]),
            use_container_width=True, hide_index=True
        )

    if is_live_market_intelligence_enabled():
        st.caption("🌐 Live Market Intelligence: key detected (provider integration pending).")

    combined_raw_text = ""
    document_summaries = []
    extraction_results = []

    if uploaded_files:
        # Single ingestion pipeline: parsing, caching, chunking, and
        # statistics all happen inside ingestion.extract_multiple(), which
        # returns one result dict per file (each with "file_name",
        # "parsed_document", "chunks", "statistics", "cached").
        extraction_results = extract_multiple(uploaded_files)

        # merge_document_text() combines every file's raw extracted text
        # into one delimited blob, replacing the old manual
        # "combined_raw_text +=" loop.
        combined_raw_text = merge_document_text(extraction_results)

        for doc in extraction_results:
            file_name = doc["file_name"]
            extracted_text = doc["parsed_document"]["text"]

            # Phase 8: hash-based cache -- re-uploading the same file
            # content doesn't re-trigger an AI summarization call.
            cache_key = f"{file_name}:{_hash_text(extracted_text)}"
            summary_text, _ = _cached_call(
                "summary_cache", cache_key,
                lambda et=extracted_text, fn=file_name: summarize_document_with_chunking(et, fn)
            )
            document_summaries.append({"file_name": file_name, "summary": summary_text})

    st.subheader("📊 Ingested Data Summary")
    col1, col2 = st.columns(2)
    col1.metric(label="📄 Files Processed", value=len(uploaded_files) if uploaded_files else 0)
    col2.metric(label="📊 Extracted Characters", value=len(combined_raw_text))

    if extraction_results:
        with st.expander("🧾 Ingestion Statistics (pages, tables, chunks, tokens)", expanded=False):
            st.text(print_statistics(document_statistics(extraction_results)))

    # --- AI Financial Assistant (interactive chatbot) ---
    # Runs after ingestion so the assistant can answer questions against
    # the uploaded documents' extracted financial facts. Passes provider
    # health so it degrades gracefully when no AI key is configured.
    render_financial_assistant(extraction_results, health)

    st.markdown("---")
    st.subheader("🔬 AI Analysis Engine")

    if st.button("🚀 Generate Timeline Report"):
        if not uploaded_files:
            st.warning("Please upload at least one financial document before generating a report.")
        else:
            # Reset the pipeline debug trace for this run so it only shows
            # entries from the current click, not accumulated across every
            # click in the session.
            st.session_state["pipeline_debug_log"] = []

            with st.spinner("Merging document summaries..."):
                merge_cache_key = _hash_text(
                    "||".join(f"{d['file_name']}:{d['summary']}" for d in document_summaries)
                )
                master_summary, _ = _cached_call(
                    "merge_cache", merge_cache_key,
                    lambda: merge_document_summaries(document_summaries)
                )
                _debug_stage("master_summary (final, pre-memo)", master_summary)

            # Bug fix: if summarization/merging failed for every uploaded
            # file/chunk, master_summary itself is an error string. Detect
            # that HERE (in addition to the existing ai_narrative_result
            # check below) so the wasted downstream memo/timeline/
            # intelligence AI calls are skipped and the user gets a clear,
            # specific reason instead of 17 sections each independently
            # reporting "insufficient information".
            master_summary_failed = (
                not master_summary or not master_summary.strip() or contains_error_marker(master_summary)
            )
            if master_summary_failed:
                st.error(
                    "❌ Document summarization failed for every uploaded file/chunk -- "
                    "there is no real content to analyze. Check the Provider Health & "
                    "Activity Log above for the underlying error, then try again."
                )

            with st.spinner("Generating investment memo..."):
                memo_system_prompt = (
                    "You are an elite institutional investment research analyst. "
                    "Write the investment memo strictly and exclusively from the "
                    "facts, figures, dates, and events contained in the Document "
                    "Summary provided below. Do not produce generic, templated, or "
                    "boilerplate analysis -- every claim must be traceable to a "
                    "specific fact stated in the Document Summary. If the Document "
                    "Summary lacks information for a requested section, explicitly "
                    "state that the source documents did not provide it rather than "
                    "inventing generic filler."
                )
                prompt = f"""Analyze the Document Summary below carefully. Extract key event milestones, timelines, and potential controversy flags SPECIFIC to this Document Summary. Write a comprehensive multi-paragraph investment memo that identifies, using only facts from the Document Summary:
1. Key financial events and dates
2. Market movements and impacts
3. Risk factors and opportunities
4. Strategic implications

Document Summary:
{master_summary}

Generate a professional investment memo grounded strictly in the Document Summary above. Do not generate generic industry commentary that is not tied to a specific fact in the Document Summary."""

                memo_cache_key = _hash_text(master_summary)
                ai_narrative_result, _ = _cached_call(
                    "memo_cache", memo_cache_key,
                    lambda: call_ai_with_fallback(prompt, system_prompt=memo_system_prompt, temperature=0.3)
                )

            st.markdown("### 📝 Generated Investment Memo")
            st.write(ai_narrative_result)

            is_error = (
                master_summary_failed
                or ("❌" in ai_narrative_result) or ("🔴" in ai_narrative_result) or ("⚠️" in ai_narrative_result)
            )

            timeline_events = []
            intelligence_outputs = {}
            if not is_error:
                with st.spinner("Extracting timeline events..."):
                    timeline_cache_key = _hash_text(ai_narrative_result)
                    timeline_events, _ = _cached_call(
                        "timeline_cache", timeline_cache_key,
                        lambda: extract_timeline_events(ai_narrative_result)
                    )
                    st.session_state["timeline_data"] = timeline_events

                if timeline_events:
                    render_timeline_visualization(timeline_events)

                # --- Module 3: deterministic financial intelligence ---
                # Runs on the raw extracted text/documents from Module 2
                # (not on the AI-generated master_summary) so its facts and
                # ratios are grounded in the actual source figures rather
                # than a paraphrase of them. Summarization and the
                # investment memo above are unaffected -- Module 3 only
                # changes what Universal Intelligence Extraction receives,
                # below.
                with st.spinner("Running Module 3 financial intelligence pipeline..."):
                    module3_cache_key = _hash_text(combined_raw_text)
                    module3_result, _ = _cached_call(
                        "module3_cache", module3_cache_key,
                        lambda: run_module3(combined_raw_text, extraction_results)
                    )

                st.markdown("### 🧮 Module 3 — Financial Intelligence")
                with st.expander("📊 Financial Data", expanded=False):
                    _render_module3_value(module3_result.get("financial_data"))
                with st.expander("📐 Ratios", expanded=False):
                    _render_module3_value(module3_result.get("ratios"))
                with st.expander("🔍 OCR Verification", expanded=False):
                    _render_module3_value(module3_result.get("ocr_verification"))
                with st.expander("🧾 Cross Document Verification", expanded=False):
                    _render_module3_value(module3_result.get("cross_document_verification"))
                with st.expander("🎯 Confidence Scores", expanded=False):
                    _render_module3_value(module3_result.get("confidence"))
                with st.expander("📅 Events", expanded=False):
                    _render_module3_value(module3_result.get("events"))
                with st.expander("🕒 Module 3 Timeline", expanded=False):
                    _render_module3_value(module3_result.get("timeline"))
                with st.expander("🗜️ Optimized Context (sent to AI)", expanded=False):
                    _render_module3_value(module3_result.get("optimized_context"))

                # The AI receives ONLY the compressed optimized_context --
                # never the raw extracted document text -- which is the
                # entire point of Module 3's token-reduction pass.
                optimized_context = module3_result.get("optimized_context") or {}
                optimized_context_text = json.dumps(optimized_context, indent=2, ensure_ascii=False, default=str)

                with st.spinner("Running institutional intelligence extraction..."):
                    intelligence_cache_key = _hash_text(optimized_context_text)
                    intelligence_outputs, _ = _cached_call(
                        "intelligence_cache", intelligence_cache_key,
                        lambda: run_universal_intelligence_extraction(optimized_context_text)
                    )
                    st.session_state["intelligence_outputs"] = intelligence_outputs
                    st.session_state["key_metrics"] = intelligence_outputs.get("key_metrics", {})
                    st.session_state["sector_analysis"] = intelligence_outputs.get("sector_analysis", {})
                    st.session_state["risk_analysis"] = intelligence_outputs.get("risk_analysis", [])
                    st.session_state["controversy_analysis"] = intelligence_outputs.get("controversy_analysis", [])

                for module_key in INTELLIGENCE_MODULES:
                    render_intelligence_output(module_key, intelligence_outputs.get(module_key))

                docx_file_stream = generate_docx_download(ai_narrative_result, timeline_events, intelligence_outputs)
                pdf_file_stream = generate_pdf_download(ai_narrative_result, timeline_events, intelligence_outputs)

                export_col1, export_col2 = st.columns(2)
                with export_col1:
                    st.download_button(
                        label="📥 Download as Word Document",
                        data=docx_file_stream,
                        file_name="Financial_Timeline_Investment_Memo.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
                with export_col2:
                    st.download_button(
                        label="📄 Download PDF",
                        data=pdf_file_stream,
                        file_name="Financial_Timeline_Investment_Memo.pdf",
                        mime="application/pdf",
                        use_container_width=True
                    )

                with st.expander("📤 More Export Formats (JSON / CSV / Excel / Markdown)", expanded=False):
                    json_bytes = generate_json_export(ai_narrative_result, timeline_events, intelligence_outputs)
                    md_bytes = generate_markdown_export(ai_narrative_result, timeline_events, intelligence_outputs)
                    csv_bytes = generate_csv_export(timeline_events, intelligence_outputs)
                    excel_bytes = generate_excel_export(
                        timeline_events, intelligence_outputs, intelligence_outputs.get("key_metrics", {})
                    )

                    more_col1, more_col2 = st.columns(2)
                    with more_col1:
                        st.download_button(
                            "🧾 Download JSON", data=json_bytes,
                            file_name="Financial_Timeline_Report.json", mime="application/json",
                            use_container_width=True
                        )
                        st.download_button(
                            "📊 Download CSV", data=csv_bytes,
                            file_name="Financial_Timeline_Report.csv", mime="text/csv",
                            use_container_width=True
                        )
                    with more_col2:
                        st.download_button(
                            "📗 Download Excel", data=excel_bytes,
                            file_name="Financial_Timeline_Report.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            use_container_width=True
                        )
                        st.download_button(
                            "📃 Download Markdown", data=md_bytes,
                            file_name="Financial_Timeline_Report.md", mime="text/markdown",
                            use_container_width=True
                        )
            else:
                st.warning("AI generation encountered an error. Please review the message above and try again.")

            with st.expander("🔍 Pipeline Debug Trace", expanded=False):
                debug_entries = st.session_state.get("pipeline_debug_log", [])
                if debug_entries:
                    st.dataframe(
                        pd.DataFrame([
                            {
                                "Stage": e["stage"],
                                "Length": e["length"],
                                "Error Marker?": "⚠️" if e["error_marker_detected"] else "",
                                "Preview (first 1000 chars)": e["preview"],
                            }
                            for e in debug_entries
                        ]),
                        use_container_width=True, hide_index=True
                    )
                else:
                    st.caption("No pipeline stages recorded for this run.")


# =============================================================================
# SECTION 12: Auth
# =============================================================================
if __name__ == "__main__":
    main()
