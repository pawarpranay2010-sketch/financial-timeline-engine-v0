# App.py
"""
Financial Timeline Engine -- Institutional Financial Intelligence Platform

Section map (Phase 12 -- modular organization within a single file):
  1. Config & Session State
  2. Parsing            (file ingestion: PDF/DOCX/XLSX/CSV/TXT)
  3. AI Providers        (Google AI Studio -> Groq -> OpenRouter fallback chain)
  4. Document Processing  (chunking, summarization, hierarchical merge, caching)
  5. Timeline             (event extraction)
  6. Analytics            (Phase 3/5 institutional intelligence modules)
  7. Export               (DOCX / PDF / JSON / CSV / Excel / Markdown)
  8. Dashboard             (visualization)
  9. Future Modules        (Phase 11 architecture scaffold)
 10. Live Market Intelligence (Phase 6 architecture scaffold)
 11. Main App / UI
 12. Auth
"""
import streamlit as st
import requests
import io
import json
from datetime import datetime, timezone
import pandas as pd
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from xml.sax.saxutils import escape as xml_escape
from google import genai
from google.genai import types

from core import (
    SETTINGS,
    FUTURE_MODULES,
    GROUNDING_RULE,
    DEFAULT_SESSION_STATE,
    get_secret,
    get_provider_health,
    provider_logger,
    hash_text,
    CacheManager,
    retry,
    extract_json,
    contains_error_marker,
)

# AI Executive Gateway -- replaces the hard-coded provider chain with
# workload-aware routing, deterministic failover, and Redis quota tracking.
# Initialised lazily so the app still starts if the gateway module has an
# import-time issue (unlikely, but keeps startup robust).
_ai_executive = None
from ingestion import (
    extract_multiple,
    merge_document_text,
    chunk_text,
    needs_chunking,
    document_statistics,
    print_statistics,
)
# TODO: verify this matches your actual backend/module3_controller.py --
# assumed path per your latest instructions (flat backend/ package, no
# module3/ subfolder) and that it exports a top-level `run_module3`.
from backend.module3_controller import run_module3

# =============================================================================
# SECTION 1: Config & Session State
# =============================================================================
# All tunable constants (model IDs, timeouts, retry policy, chunking sizes,
# the future-module roadmap) now live in core/config.py as `SETTINGS`, a
# single typed EngineSettings instance -- see that module's docstring for
# why. The names below are kept as module-level aliases so every existing
# reference further down in this file (PRIMARY_MODEL, CHUNK_SIZE, etc.)
# continues to work completely unchanged.
st.set_page_config(page_title="Financial Timeline Engine", layout="centered")

PRIMARY_MODEL = SETTINGS.primary_model
FALLBACK_MODEL = SETTINGS.fallback_model
OPENROUTER_MODELS = list(SETTINGS.openrouter_models)
GROQ_MODELS = list(SETTINGS.groq_models)

GROQ_TIMEOUT_SECONDS = SETTINGS.groq_timeout_seconds
OPENROUTER_TIMEOUT_SECONDS = SETTINGS.openrouter_timeout_seconds
PROVIDER_RETRY_ATTEMPTS = SETTINGS.provider_retry_attempts
PROVIDER_RETRY_DELAY_SECONDS = SETTINGS.provider_retry_delay_seconds

CHUNK_SIZE = SETTINGS.chunk_size
CHUNK_OVERLAP = SETTINGS.chunk_overlap
MERGE_BATCH_SIZE = SETTINGS.merge_batch_size

LIVE_INTELLIGENCE_API_KEY_NAME = SETTINGS.live_intelligence_api_key_name

# --- Session state init ---
# Default shape now lives in core/constants.py (DEFAULT_SESSION_STATE) so
# it can be reused by future non-Streamlit entry points (tests, backend
# warm-up) without importing this UI module.
for _key, _default in DEFAULT_SESSION_STATE.items():
    if _key not in st.session_state:
        st.session_state[_key] = _default


def _hash_text(text):
    """Thin backward-compatible alias for core.utilities.hash_text, kept
    so every call site below (`_hash_text(...)`) needs no further edits."""
    return hash_text(text)


def _cached_call(cache_name, cache_key, compute_fn):
    """Thin backward-compatible alias over core.utilities.CacheManager,
    preserving the exact (cache_name, cache_key, compute_fn) call
    signature used throughout this file. Builds a fresh CacheManager over
    the relevant st.session_state[cache_name] dict on each call -- cheap,
    since CacheManager itself holds no state beyond the mapping reference."""
    cache_store = st.session_state.setdefault(cache_name, {})
    return CacheManager(cache_store).get_or_compute(cache_key, compute_fn)


def _log_provider_event(stage, provider, status, detail=""):
    """Thin backward-compatible alias over the shared
    core.logging.provider_logger instance."""
    return provider_logger.log(stage, provider, status, detail)


def _retry(fn, attempts=PROVIDER_RETRY_ATTEMPTS, delay=PROVIDER_RETRY_DELAY_SECONDS):
    """Thin backward-compatible alias over core.utilities.retry."""
    return retry(fn, attempts=attempts, delay=delay)


def _debug_stage(label, text):
    """Pipeline debug instrumentation (added while diagnosing the
    "always insufficient information" + export traceback bugs): records
    the length, a first-1000-character preview, and whether an error
    marker is present for one intermediate pipeline output. Entries
    accumulate in st.session_state['pipeline_debug_log'] for the current
    "Generate Timeline Report" run (reset at the top of that button
    handler) and are rendered in a "Pipeline Debug Trace" expander so the
    exact stage where real content is lost -- or silently replaced by an
    error/placeholder -- is visible directly in the UI, without needing
    server log access."""
    text = text or ""
    st.session_state.setdefault("pipeline_debug_log", []).append({
        "stage": label,
        "length": len(text),
        "error_marker_detected": contains_error_marker(text),
        "preview": text[:1000],
    })

# =============================================================================
# SECTION 2: Parsing (file ingestion)
# =============================================================================
# All document parsing (PDF/DOCX/XLSX/CSV/TXT), chunking, caching, and
# per-file statistics now live in the ingestion/ package (parser.py,
# chunking.py, cache.py, statistics.py, extraction.py) and are imported
# above. There is exactly one document ingestion pipeline --
# ingestion.extract_multiple() / ingestion.merge_document_text() -- and no
# parsing logic is duplicated here. The legacy in-file extraction function
# that used to live in this section has been removed.


# =============================================================================
# SECTION 3: AI Providers (canonical gateway — backend/gateway/)
# =============================================================================
# backend/gateway/ is the SINGLE canonical AI provider/fallback system for
# the whole application (workload-aware routing, capability detection,
# provider health, deterministic failover, normalized responses, graceful
# no-provider behavior, env-vars -> Streamlit-secrets configuration).
#
# The four legacy functions below are retained ONLY as thin compatibility
# wrappers so existing callers/tests keep working; none of them implements
# its own provider chain anymore. Every real generation request flows
# through call_ai_with_fallback() -> AIExecutive.generate() ->
# ProviderManager/router/adapters. A missing optional provider (e.g.
# OpenRouter) is skipped by the gateway's failover and can NEVER surface as
# a raw provider-specific error string.

NO_ELIGIBLE_PROVIDER_MESSAGE = (
    "No eligible AI provider is configured. Add an AI provider key to continue."
)


def call_google_ai_studio(prompt_text, system_prompt=None, temperature=None):
    """Thin compatibility wrapper — generation now flows through the
    canonical AI gateway (backend/gateway). Kept only so existing
    callers/tests keep working; checks the Google key, then delegates
    to `call_ai_with_fallback` (the app's single provider chain)."""
    api_key = get_secret("GOOGLE_API_KEY", "")
    if not api_key:
        raise ValueError("Missing Google Key")
    return call_ai_with_fallback(prompt_text, system_prompt=system_prompt, temperature=temperature)


def call_groq_engine(prompt_text, system_prompt=None, temperature=None):
    """Thin compatibility wrapper — generation now flows through the
    canonical AI gateway (backend/gateway). Kept only so existing
    callers/tests keep working; checks the Groq key, then delegates
    to `call_ai_with_fallback` (the app's single provider chain)."""
    api_key = get_secret("GROQ_API_KEY", "")
    if not api_key:
        raise ValueError("Missing Groq Key")
    return call_ai_with_fallback(prompt_text, system_prompt=system_prompt, temperature=temperature)


def _openrouter_request(prompt_text, model_id, system_prompt=None, temperature=None):
    """Thin compatibility wrapper — generation now flows through the
    canonical AI gateway (backend/gateway), which skips unavailable
    providers deterministically. Returns (success: bool, content_or_error).
    A missing OpenRouter key yields the shared graceful no-provider
    message — never a raw provider-specific error string."""
    api_key = get_secret("OPENROUTER_API_KEY", "")
    if not api_key:
        return False, NO_ELIGIBLE_PROVIDER_MESSAGE
    result = call_ai_with_fallback(prompt_text, system_prompt=system_prompt, temperature=temperature)
    if result.startswith(("❌", "🔴", "⚠️")):
        return False, result
    return True, result


def call_openrouter_engine(prompt_text, system_prompt=None, temperature=None):
    """Thin compatibility wrapper — generation now flows through the
    canonical AI gateway (backend/gateway). Kept only so existing
    callers/tests keep working."""
    if system_prompt is None:
        system_prompt = (
            "You are an elite Wall Street financial research analyst. Generate "
            "structured multi-section corporate reports with key dates, events, "
            "and milestones."
        )
    _, result = _openrouter_request(prompt_text, PRIMARY_MODEL, system_prompt=system_prompt, temperature=temperature)
    return result


def _get_ai_executive():
    """Lazy singleton: initialises the AI Executive gateway on first use.
    Returns None if the module isn't available (graceful no-provider
    behavior downstream)."""
    global _ai_executive
    if _ai_executive is None:
        try:
            from backend.gateway import AIExecutive
            _ai_executive = AIExecutive()
        except Exception:
            _ai_executive = None
    return _ai_executive


def get_canonical_provider_status():
    """Per-provider status from the canonical ProviderManager
    (backend/gateway). Returns {slug: "available" |
    "configured_unavailable" | "not_configured"} for every registered
    provider, so the UI can distinguish 🟢 available / 🟡 configured but
    unavailable / ⚪ not configured instead of claiming live connectivity
    from a key alone."""
    executive = _get_ai_executive()
    if executive is None:
        return {}
    pm = executive.provider_manager
    status = {}
    for name in pm.DEFAULT_PRIORITY:
        keyed = pm.key_status().get(name, False)
        adapter = pm.get(name)
        if keyed and adapter is not None and adapter.health_check():
            status[name] = "available"
        elif keyed:
            status[name] = "configured_unavailable"
        else:
            status[name] = "not_configured"
    return status


def _detect_task_type(system_prompt, prompt_text):
    """Heuristic: guess the workload type from prompt content so the AI
    Executive router can select the best provider. Returns a task type
    string compatible with backend.gateway.router.* constants.

    Rules (checked in order):
    - system_prompt mentions JSON / array / object → "structured"
    - prompt asks for "investment memo" / "financial" → "financial"
    - prompt is very short (< 200 chars) → "simple"
    - otherwise → "financial" (safe default for this application)
    """
    combined = ((system_prompt or "") + " " + (prompt_text or "")).lower()
    json_keywords = ["json", "return only valid json", "return a json", "as json",
                     "return only a single valid json object", "expected_type"]
    if any(kw in combined for kw in json_keywords):
        return "structured"
    memo_keywords = ["investment memo", "investment research", "analyze the document summary",
                     "financial analysis", "financial statements", "analyze financial",
                     "institutional", "key financial events", "market movements",
                     "rag", "retrieved", "retrieved data", "retrieved documents",
                     "based on the retrieved", "based only on the following",
                     "source documents", "provided context"]
    if any(kw in combined for kw in memo_keywords):
        return "financial"
    if len(prompt_text or "") < 200:
        return "simple"
    return "financial"


def call_ai_with_fallback(prompt_text, system_prompt=None, temperature=None):
    """SINGLE canonical AI entry point for the entire application.

    Delegates exclusively to the AI Executive gateway (backend/gateway/):
        prompt -> AIExecutive.generate -> Router (workload-aware)
               -> ProviderManager + adapters (Google -> Groq -> OpenRouter
                  -> NVIDIA -> ... deterministic failover)
               -> normalized response -> content

    There is no second provider chain anywhere in app.py. Missing optional
    providers are skipped by the gateway's failover, so a missing
    OPENROUTER_API_KEY can never block an available Groq key and can never
    surface as a raw error string. If no eligible provider exists (or the
    gateway is unavailable), the shared graceful message is returned.
    """
    executive = _get_ai_executive()
    if executive is not None:
        try:
            task_type = _detect_task_type(system_prompt, prompt_text)
            response = executive.generate(
                prompt=prompt_text,
                system_prompt=system_prompt or "",
                temperature=temperature or 0.3,
                task_type=task_type,
            )
            if response.success:
                st.session_state["ai_connected"] = True
                st.session_state["ai_provider_used"] = f"{response.provider}/{response.model}"
                _log_provider_event("call_ai_with_fallback", response.provider, "success",
                                    f"task={task_type}, model={response.model}, latency={response.latency_ms}ms")
                return response.content
            # Gateway returned an error -- log and return the graceful
            # no-provider message (never a raw provider error string).
            _log_provider_event("call_ai_with_fallback", "gateway", "failed",
                                f"task={task_type}, error={(response.error or 'unknown')[:100]}")
        except Exception as e:
            _log_provider_event("call_ai_with_fallback", "gateway", "failed",
                                f"exception={type(e).__name__}: {str(e)[:100]}")

    # No eligible provider / gateway unavailable: graceful and machine-
    # detectable (⚠️ is in ERROR_RESPONSE_MARKERS) so downstream pipeline
    # stages degrade instead of treating this as generated content.
    st.session_state["ai_connected"] = False
    st.session_state["ai_provider_used"] = ""
    return f"⚠️ {NO_ELIGIBLE_PROVIDER_MESSAGE}"
  
# =============================================================================
# SECTION 4: Document Processing (chunking, summarization, hierarchical merge)
# =============================================================================
# GROUNDING_RULE now lives in core/constants.py and is imported above; kept
# as a plain reference here (no reassignment) so every use further down in
# this file is unchanged.


# =============================================================================
# SECTION 4: Document Processing (chunking, summarization, hierarchical merge)
# =============================================================================
# GROUNDING_RULE now lives in core/constants.py and is imported above; kept
# as a plain reference here (no reassignment) so every use further down in
# this file is unchanged.
#
# chunk_text() and needs_chunking() are no longer defined here -- they are
# imported from the ingestion package above (ingestion/chunking.py), which
# is now the single source of chunking logic for the whole application.
# This eliminates the duplicate chunking implementation that used to exist
# in both app.py and ingestion/chunking.py.


def summarize_single_document(document_text, file_name):
    """Summarizes one document/chunk into a concise institutional
    financial summary (500-1000 words), explicitly instructed to
    preserve numbers, tables, dates, timeline events, management
    commentary, and financial-statement figures rather than smoothing
    them into vague prose (Phase 2)."""
    if not document_text or not document_text.strip():
        return f"⚠️ No extractable text found in '{file_name}'. Nothing to summarize."

    system_prompt = (
        "You are an elite institutional financial analyst. Produce concise, "
        "precise, and professional document summaries suitable for a "
        "buy-side investment research desk. Focus on material facts: exact "
        "figures, tables (rendered as clear text), dates, timeline events, "
        "management commentary, and financial statement line items. Avoid "
        "filler language and avoid restating the obvious. " + GROUNDING_RULE
    )

    summarization_prompt = f"""Summarize the following document into a single, coherent institutional financial summary.

Requirements:
- Length: 500 to 1000 words.
- Tone: professional, analytical, institutional-grade.
- Preserve exact figures, tables, dates, timeline events, management commentary, and financial statement data verbatim -- do not round or paraphrase numbers.
- Do not use markdown headers or bullet lists; write in clear prose paragraphs (tables may be rendered as simple "Label: Value" lines within the prose).

Source Document Name: {file_name}

Source Document Text:
{document_text}

Return only the summary text, with no preamble or meta-commentary."""

    return call_ai_with_fallback(summarization_prompt, system_prompt=system_prompt, temperature=0.3)


def _merge_summary_batch(document_summaries):
    """Merges a single batch (<= MERGE_BATCH_SIZE) of summaries into one
    coherent master summary. Internal helper for merge_document_summaries;
    see that function for the hierarchical/recursive batching logic.

    ROOT-CAUSE FIX for "every intelligence module says insufficient
    information": an individual chunk/document summarization call can
    fail and return one of this app's own rendered error strings (e.g.
    "❌ OpenRouter Connection Failed...", see call_ai_with_fallback /
    summarize_single_document) instead of raising. Previously that error
    string was indistinguishable from a real summary by the time it
    reached here, so it was fed straight into the merge prompt as if it
    were genuine document content. When enough (or all) inputs to a merge
    were error strings, the merge model had nothing real to work with and,
    per its own grounding instructions, correctly reported that no
    information was available -- and *that* message became the
    master_summary, so every downstream intelligence field also
    (correctly, given that garbage input) reported insufficient
    information. This was the exact stage where real extracted text was
    being replaced by a fallback message.

    Fix: error-marked inputs are filtered out before building the merge
    prompt (the prompt text itself is completely unchanged), and if
    literally everything in this batch failed, a clear error is returned
    immediately instead of asking the AI to "merge" garbage.
    """
    valid_summaries = [d for d in document_summaries if not contains_error_marker(d.get("summary", ""))]
    failed_summaries = [d for d in document_summaries if contains_error_marker(d.get("summary", ""))]

    for doc in failed_summaries:
        _log_provider_event(
            "_merge_summary_batch", "pipeline", "excluded",
            f"Excluded '{doc.get('file_name')}' from merge -- its summary was an "
            f"error message, not real content: {str(doc.get('summary', ''))[:200]}"
        )

    if not valid_summaries:
        error_result = (
            "❌ Unable to generate a summary: every source document/chunk in this "
            "batch failed AI summarization. Check the Provider Health & Activity "
            "Log above for the underlying provider errors, then try again."
        )
        _debug_stage("merge_batch_result (all inputs failed)", error_result)
        return error_result

    system_prompt = (
        "You are an elite institutional financial analyst. Merge multiple "
        "per-document summaries into a single, coherent master summary "
        "suitable for a buy-side investment research desk. Reconcile "
        "overlapping information rather than repeating it, but preserve "
        "every distinct financial metric, table, date, timeline event, "
        "management commentary, and strategic insight found across the "
        "source summaries. Do not drop material information for brevity. "
        + GROUNDING_RULE
    )

    combined_summaries_text = ""
    for doc in valid_summaries:
        file_name = doc.get("file_name", "Unknown Document")
        summary = doc.get("summary", "")
        combined_summaries_text += f"\n--- Summary of: {file_name} ---\n{summary}\n"

    merge_prompt = f"""Below are individual summaries of separate financial documents (or document sections). Merge them into one coherent institutional financial master summary.

Requirements:
- Reconcile and consolidate overlapping information; do not repeat the same fact twice.
- Preserve all distinct financial metrics, tables, dates/timelines, management commentary, risk factors, controversies, and strategic implications mentioned across the summaries.
- Organize the merged summary in clear prose paragraphs.
- Attribute conflicting figures or claims across documents where relevant, rather than silently picking one.
- Do not fabricate information not present in the source summaries.

Individual Summaries:
{combined_summaries_text}

Return only the merged master summary text, with no preamble or meta-commentary."""

    result = call_ai_with_fallback(merge_prompt, system_prompt=system_prompt, temperature=0.3)
    _debug_stage("merge_batch_result", result)
    return result


def merge_document_summaries(document_summaries):
    """Merges per-document (or per-chunk) summaries into one coherent
    institutional financial master summary.

    Phase 2/7 -- hierarchical/recursive merge: if there are more than
    MERGE_BATCH_SIZE summaries, they are merged in batches first, and the
    batch-level merged summaries are then recursively merged again. This
    lets the pipeline handle an arbitrarily large number of documents or
    chunks without ever sending an oversized single merge request to any
    provider (which is what caused Groq's earlier HTTP 413s)."""
    if not document_summaries or len(document_summaries) == 0:
        return "⚠️ No document summaries available to merge."

    if len(document_summaries) <= MERGE_BATCH_SIZE:
        return _merge_summary_batch(document_summaries)

    batch_merged = []
    for i in range(0, len(document_summaries), MERGE_BATCH_SIZE):
        batch = document_summaries[i:i + MERGE_BATCH_SIZE]
        batch_summary = _merge_summary_batch(batch)
        batch_merged.append({"file_name": f"Batch {i // MERGE_BATCH_SIZE + 1}", "summary": batch_summary})

    return merge_document_summaries(batch_merged)


def summarize_document_with_chunking(document_text, file_name):
    """For large documents, splits text into chunks before summarizing
    (avoiding oversized AI requests, e.g. Groq HTTP 413), then merges the
    chunk summaries into one document-level summary via the existing
    hierarchical merge_document_summaries(). Small documents skip
    chunking entirely and behave exactly as summarize_single_document.

    Instrumented with _debug_stage() at every intermediate output
    (extracted text length, each chunk summary, the merged result) so the
    exact stage where real content is lost is visible in the Pipeline
    Debug Trace expander."""
    _debug_stage(f"extracted_text :: {file_name}", document_text)

    if not document_text or not needs_chunking(document_text, threshold=CHUNK_SIZE):
        summary = summarize_single_document(document_text, file_name)
        _debug_stage(f"single_document_summary :: {file_name}", summary)
        return summary

    chunks = chunk_text(document_text, chunk_size=CHUNK_SIZE, overlap=CHUNK_OVERLAP)
    chunk_summaries = []
    for idx, chunk in enumerate(chunks, start=1):
        chunk_label = f"{file_name} (Part {idx}/{len(chunks)})"
        chunk_summary = summarize_single_document(chunk, chunk_label)
        _debug_stage(f"chunk_summary :: {chunk_label}", chunk_summary)
        chunk_summaries.append({"file_name": chunk_label, "summary": chunk_summary})

    merged = merge_document_summaries(chunk_summaries)
    _debug_stage(f"document_merged_summary :: {file_name}", merged)
    return merged


# =============================================================================
# SECTION 5: Timeline extraction
# =============================================================================
def _extract_json_from_ai_response(result, expected_type=dict):
    """Thin backward-compatible alias over core.validation.extract_json,
    which now owns the markdown-fence stripping, error-marker detection,
    and bracket-scan fallback logic (unchanged behavior, single
    implementation shared with any future caller outside this file)."""
    return extract_json(result, expected_type=expected_type)


def extract_timeline_events(ai_narrative):
    """Parses AI narrative to extract structured timeline events."""
    try:
        structuring_prompt = f"""Extract timeline events from this narrative and return as JSON array with objects containing: date (YYYY-MM-DD or YYYY-MM or YYYY), event (string), category (string), impact (string).

Narrative:
{ai_narrative}

Return ONLY valid JSON array, no markdown, no extra text."""

        result = call_ai_with_fallback(structuring_prompt, temperature=0.3)
        return _extract_json_from_ai_response(result, expected_type=list)
    except Exception:
        return []


# =============================================================================
# SECTION 6: Analytics -- Phase 3/5 Institutional Intelligence Modules
# =============================================================================
# All 17 sections below are produced by ONE consolidated AI call
# (run_universal_intelligence_extraction), not 17 separate calls -- this
# directly serves Phase 8 (reduce AI cost / avoid duplicate calls) while
# still covering the full requested breadth.
INTELLIGENCE_MODULES = {
    "executive_summary": {
        "title": "📝 Executive Summary",
        "instruction": "A concise institutional executive summary (250-400 words) covering the business, recent performance, and overall investment context.",
        "expected_type": str,
    },
    "key_metrics": {
        "title": "📌 Key Financial Metrics",
        "instruction": "The key financial metrics (Revenue, EBITDA, PAT, EPS, Debt, Cash, Margins, Growth, Capex, Free Cash Flow, ROE, ROCE, Market Cap, and other material KPIs) are already extracted and verified in the 'financial_data' section of the Source Summary below -- report them as given, with their stated source/confidence, rather than re-deriving them from prose.",
        "expected_type": dict,
    },
    "ratio_analysis": {
        "title": "📐 Ratio Analysis",
        "instruction": "The 'ratios' section of the Source Summary below already contains every ratio (Revenue Growth, PAT Growth, EBITDA Growth, ROE, ROCE, Debt/Equity, Current Ratio, Operating Margin, Net Margin, EPS, Free Cash Flow indicators) that could be verified or calculated deterministically, each tagged with its source (Document or Calculated) and formula where applicable. Report these ratios as given -- do not recalculate them yourself. For any ratio marked 'Cannot Calculate' or absent from that section, use the exact string 'Not computable from provided documents'.",
        "expected_type": dict,
    },
    "financial_performance": {
        "title": "📈 Financial Performance",
        "instruction": "A structured narrative (300-500 words) describing revenue, profitability, and cash flow trends explicitly stated in the source.",
        "expected_type": str,
    },
    "segment_analysis": {
        "title": "🧩 Segment Analysis",
        "instruction": "Each distinct business segment/division mentioned, as an array of objects with 'segment_name', 'performance_summary', 'contribution'.",
        "expected_type": list,
    },
    "sector_analysis": {
        "title": "🏭 Sector Analysis",
        "instruction": "An object with 'sector', 'industry', 'business_model', 'competitors' (list), 'market_position', 'industry_trends', 'peer_context'.",
        "expected_type": dict,
    },
    "competitor_analysis": {
        "title": "⚔️ Competitor Analysis",
        "instruction": "Named competitors/peers mentioned, as an array of objects with 'competitor_name', 'context'.",
        "expected_type": list,
    },
    "swot": {
        "title": "🧭 SWOT Analysis",
        "instruction": "An object with keys 'strengths', 'weaknesses', 'opportunities', 'threats', each a list of short strings grounded in the source.",
        "expected_type": dict,
    },
    "risk_analysis": {
        "title": "⚠️ Risk Analysis",
        "instruction": "A structured risk assessment as an array of objects with 'category' (Business, Financial, Operational, Governance, Regulatory, Macroeconomic, or Investment), 'risk', 'severity' (Low/Medium/High), 'probability' (Low/Medium/High), 'mitigation'.",
        "expected_type": list,
    },
    "controversy_analysis": {
        "title": "🚨 Controversy Analysis",
        "instruction": "Any controversies (litigation, fraud, governance issues, management changes, accounting issues, environmental, political, regulatory actions, negative news) as an array of objects with 'date', 'type', 'description', 'severity', 'source'.",
        "expected_type": list,
    },
    "governance_analysis": {
        "title": "🏛️ Governance Analysis",
        "instruction": "A narrative (150-300 words) on corporate governance details explicitly mentioned (board structure, management changes, related-party matters, audit/compliance notes).",
        "expected_type": str,
    },
    "esg_summary": {
        "title": "🌱 ESG Summary",
        "instruction": "A narrative (150-300 words) summarizing any Environmental, Social, and Governance information explicitly mentioned.",
        "expected_type": str,
    },
    "investment_thesis": {
        "title": "💡 Investment Thesis",
        "instruction": "A concise investment thesis (150-300 words) grounded strictly in the source.",
        "expected_type": str,
    },
    "bull_case": {
        "title": "🐂 Bull Case",
        "instruction": "The strongest bullish arguments supported by the source, as an array of short strings.",
        "expected_type": list,
    },
    "bear_case": {
        "title": "🐻 Bear Case",
        "instruction": "The strongest bearish/risk arguments supported by the source, as an array of short strings.",
        "expected_type": list,
    },
    "catalysts": {
        "title": "⚡ Catalysts",
        "instruction": "Specific upcoming events/dates/triggers mentioned that could move the outlook, as an array of objects with 'catalyst', 'expected_timing'.",
        "expected_type": list,
    },
    "action_points": {
        "title": "✅ Action Points",
        "instruction": "Concrete, specific action points/recommendations an analyst would take next, as an array of short strings.",
        "expected_type": list,
    },
}


def _type_hint(expected_type):
    if expected_type is dict:
        return "a JSON object"
    if expected_type is list:
        return "a JSON array"
    return "a plain text string (no markdown)"


def run_universal_intelligence_extraction(master_summary):
    """Phase 3/5/8: single consolidated AI call that extracts all 17
    institutional intelligence sections defined in INTELLIGENCE_MODULES
    at once. Returns a dict keyed by module key; any module missing from
    the AI's response falls back to an empty value of the expected type
    so downstream rendering never breaks on a missing key.

    BUG FIXES:
    - Now also short-circuits (returning empty_result) when master_summary
      itself is an error string (contains_error_marker(...)), not just
      when it's empty. Previously an error string that slipped through
      from an upstream summarization/merge failure was treated as real
      content and sent to the AI anyway -- with nothing genuine to work
      from, the AI would (correctly, per its own grounding instructions)
      report "insufficient information" for every single field. Combined
      with the _merge_summary_batch fix above, master_summary should now
      only ever be an error string if literally every source failed.
    - Each field's value is now checked against its declared
      `expected_type` before being kept, and a `list` value containing a
      mix of dict and non-dict items (which the model occasionally
      returns) is normalized down to just its dict items. An unfiltered
      mixed list is what caused `pd.DataFrame(...)` to raise
      "ValueError: dictionary update sequence element #0 has length 1; 2
      is required" downstream in rendering/export -- the traceback that
      appeared right after generation.
    """
    empty_result = {key: config["expected_type"]() for key, config in INTELLIGENCE_MODULES.items()}

    _debug_stage("intelligence_extraction_input", master_summary)

    if not master_summary or not master_summary.strip() or contains_error_marker(master_summary):
        return empty_result

    system_prompt = (
        "You are an elite institutional financial research analyst producing "
        "a complete, structured research package. " + GROUNDING_RULE
    )

    field_instructions = "\n".join(
        f'- "{key}" ({_type_hint(config["expected_type"])}): {config["instruction"]}'
        for key, config in INTELLIGENCE_MODULES.items()
    )

    extraction_prompt = f"""Analyze the Source Summary below and produce a complete institutional research package.

Return ONLY a single valid JSON object with exactly these top-level keys:

{field_instructions}

Source Summary:
{master_summary}

Return ONLY the JSON object. No markdown, no extra text, no preamble."""

    result = call_ai_with_fallback(extraction_prompt, system_prompt=system_prompt, temperature=0.2)
    parsed = _extract_json_from_ai_response(result, expected_type=dict)
    if not parsed:
        return empty_result

    final = dict(empty_result)
    for key, config in INTELLIGENCE_MODULES.items():
        if key not in parsed or not parsed[key]:
            continue
        value = parsed[key]
        expected_type = config["expected_type"]

        if expected_type is list and isinstance(value, list):
            dict_items = [item for item in value if isinstance(item, dict)]
            if dict_items and len(dict_items) != len(value):
                # Mixed list (some dict items, some not) -- keep only the
                # structured items so downstream pd.DataFrame(...) calls
                # in rendering/export never see an inconsistent shape.
                value = dict_items
            final[key] = value
        elif isinstance(value, expected_type):
            final[key] = value
        # else: the AI returned a type that doesn't match this field's
        # expected_type at all (e.g. a plain string where a list was
        # expected) -- leave this field at its empty default rather than
        # passing an unexpected shape downstream.

    return final


def render_intelligence_output(key, value):
    """Renders one intelligence module's output using the same visual
    language already used elsewhere in the app (st.subheader, st.dataframe,
    st.markdown) -- no new UI patterns introduced."""
    config = INTELLIGENCE_MODULES[key]
    if not value:
        return
    st.subheader(config["title"])

    if isinstance(value, dict):
        if key == "swot":
            swot_col1, swot_col2 = st.columns(2)
            with swot_col1:
                st.markdown("**Strengths**")
                for item in value.get("strengths") or []:
                    st.markdown(f"- {item}")
                st.markdown("**Opportunities**")
                for item in value.get("opportunities") or []:
                    st.markdown(f"- {item}")
            with swot_col2:
                st.markdown("**Weaknesses**")
                for item in value.get("weaknesses") or []:
                    st.markdown(f"- {item}")
                st.markdown("**Threats**")
                for item in value.get("threats") or []:
                    st.markdown(f"- {item}")
        elif key == "sector_analysis":
            for field_name, field_value in value.items():
                if field_value:
                    label = field_name.replace("_", " ").title()
                    if isinstance(field_value, list):
                        field_value = ", ".join(str(v) for v in field_value)
                    st.markdown(f"**{label}:** {field_value}")
        else:
            st.dataframe(
                pd.DataFrame(list(value.items()), columns=["Field", "Value"]),
                use_container_width=True,
                hide_index=True
            )
    elif isinstance(value, list):
        if value and all(isinstance(item, dict) for item in value):
            st.dataframe(pd.DataFrame(value), use_container_width=True, hide_index=True)
        else:
            for item in value:
                st.markdown(f"- {item}")
    else:
        st.markdown(str(value))


def _render_module3_value(value):
    """Generic, defensive renderer for one Module 3 output section
    (financial_data, ratios, ocr_verification, cross_document_verification,
    confidence, events, timeline, optimized_context). Module 3's internal
    schema for each section is intentionally treated as opaque here --
    this renders any dict as interactive JSON (st.json), any list of dicts
    as a table, any other list as bullets, and anything else as plain
    text, so the UI never breaks regardless of Module 3's exact internal
    shape, and can be refined later without another app.py change."""
    if not value:
        st.caption("No data available for this section.")
        return
    if isinstance(value, dict):
        st.json(value)
    elif isinstance(value, list):
        if value and all(isinstance(item, dict) for item in value):
            st.dataframe(pd.DataFrame(value), use_container_width=True, hide_index=True)
        else:
            for item in value:
                st.markdown(f"- {item}")
    else:
        st.markdown(str(value))
# =============================================================================
# SECTION 7: Export System (DOCX / PDF / JSON / CSV / Excel / Markdown)
# =============================================================================
def _flatten_intelligence_for_text(intelligence_outputs):
    """Shared helper: converts the intelligence_outputs dict into a flat
    list of (title, rendered_text_or_rows) for reuse across DOCX/PDF/MD
    export, avoiding duplicating this formatting logic three times."""
    sections = []
    for key, value in (intelligence_outputs or {}).items():
        if not value:
            continue
        title = INTELLIGENCE_MODULES.get(key, {}).get("title", key)
        sections.append((title, value))
    return sections


def generate_docx_download(text_content, timeline_data=None, intelligence_outputs=None):
    """Compiles the generated AI analysis report into a Word document
    download stream. intelligence_outputs is optional (default None) so
    any existing call site without it still behaves exactly as before."""
    doc = Document()

    doc.add_heading("Institutional Investment Research Memo", level=1)
    doc.add_paragraph("-" * 40)
    doc.add_heading("Executive Summary & Analysis", level=2)

    if text_content:
        clean_text_string = str(text_content)
        for line in clean_text_string.split('\n'):
            if line.strip():
                sanitized_line = "".join(c for c in line if c.isprintable() or c in ['\t', '\n'])
                sanitized_line = sanitized_line.replace('**', '').replace('__', '').replace('```', '')
                if sanitized_line.strip():
                    doc.add_paragraph(sanitized_line.strip())
    else:
        doc.add_paragraph("No report content generated.")

    if timeline_data and len(timeline_data) > 0:
        doc.add_heading("Extracted Timeline Events", level=2)
        for event in timeline_data:
            date_str = "".join(c for c in str(event.get("date", "N/A")) if c.isprintable())
            event_name = "".join(c for c in str(event.get("event", "N/A")) if c.isprintable())
            category = "".join(c for c in str(event.get("category", "N/A")) if c.isprintable())
            impact = "".join(c for c in str(event.get("impact", "N/A")) if c.isprintable())
            doc.add_paragraph(f"📅 {date_str}: {event_name}", style="List Bullet")
            doc.add_paragraph(f"Category: {category} | Impact: {impact}", style="List Bullet 2")

    for title, value in _flatten_intelligence_for_text(intelligence_outputs):
        doc.add_heading(title, level=2)
        if isinstance(value, str):
            doc.add_paragraph(value)
        elif isinstance(value, dict):
            for field_name, field_value in value.items():
                if isinstance(field_value, list):
                    field_value = ", ".join(str(v) for v in field_value)
                doc.add_paragraph(f"{field_name.replace('_', ' ').title()}: {field_value}", style="List Bullet")
        elif isinstance(value, list):
            for item in value:
                if isinstance(item, dict):
                    line = " | ".join(f"{k}: {v}" for k, v in item.items())
                else:
                    line = str(item)
                doc.add_paragraph(line, style="List Bullet")

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


def generate_pdf_download(text_content, timeline_data=None, intelligence_outputs=None):
    """Compiles the report into an institutional-styled PDF using
    reportlab. intelligence_outputs is optional (default None) for
    backward compatibility. All text is XML-escaped since reportlab's
    Paragraph uses an internal XML markup dialect, and emoji are avoided
    in this output (reportlab's built-in fonts lack emoji glyphs, unlike
    python-docx which substitutes fonts automatically)."""
    bio = io.BytesIO()
    doc = SimpleDocTemplate(bio, pagesize=letter, topMargin=0.75 * 72, bottomMargin=0.75 * 72)
    base_styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        "InstitutionalTitle", parent=base_styles["Title"],
        textColor=colors.HexColor("#1a2b4c"), fontSize=20, spaceAfter=4
    )
    meta_style = ParagraphStyle(
        "InstitutionalMeta", parent=base_styles["Normal"],
        textColor=colors.grey, fontSize=8, spaceAfter=14
    )
    heading_style = ParagraphStyle(
        "InstitutionalHeading", parent=base_styles["Heading2"],
        textColor=colors.HexColor("#1a2b4c"), spaceBefore=14, spaceAfter=6
    )
    normal_style = base_styles["Normal"]

    story = [
        Paragraph("Institutional Investment Research Memo", title_style),
        Paragraph(f"Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}", meta_style),
        Paragraph("Executive Summary &amp; Analysis", heading_style),
        Spacer(1, 6),
    ]

    if text_content:
        for line in str(text_content).split('\n'):
            if line.strip():
                sanitized_line = "".join(c for c in line if c.isprintable() or c in ['\t', '\n'])
                sanitized_line = sanitized_line.replace('**', '').replace('__', '').replace('```', '')
                if sanitized_line.strip():
                    story.append(Paragraph(xml_escape(sanitized_line.strip()), normal_style))
                    story.append(Spacer(1, 6))
    else:
        story.append(Paragraph("No report content generated.", normal_style))

    if timeline_data and len(timeline_data) > 0:
        story.append(Paragraph("Extracted Timeline Events", heading_style))
        for event in timeline_data:
            date_str = "".join(c for c in str(event.get("date", "N/A")) if c.isprintable())
            event_name = "".join(c for c in str(event.get("event", "N/A")) if c.isprintable())
            category = "".join(c for c in str(event.get("category", "N/A")) if c.isprintable())
            impact = "".join(c for c in str(event.get("impact", "N/A")) if c.isprintable())
            story.append(Paragraph(xml_escape(f"Date: {date_str} -- {event_name}"), normal_style))
            story.append(Paragraph(xml_escape(f"Category: {category} | Impact: {impact}"), normal_style))
            story.append(Spacer(1, 6))

    for title, value in _flatten_intelligence_for_text(intelligence_outputs):
        story.append(Paragraph(xml_escape(title), heading_style))
        if isinstance(value, str):
            story.append(Paragraph(xml_escape(value), normal_style))
        elif isinstance(value, dict):
            for field_name, field_value in value.items():
                if isinstance(field_value, list):
                    field_value = ", ".join(str(v) for v in field_value)
                line = f"{field_name.replace('_', ' ').title()}: {field_value}"
                story.append(Paragraph(xml_escape(line), normal_style))
        elif isinstance(value, list):
            for item in value:
                if isinstance(item, dict):
                    line = " | ".join(f"{k}: {v}" for k, v in item.items())
                else:
                    line = str(item)
                story.append(Paragraph(xml_escape(line), normal_style))
        story.append(Spacer(1, 6))

    doc.build(story)
    bio.seek(0)
    return bio


def generate_json_export(ai_narrative_result, timeline_events, intelligence_outputs):
    """Phase 4: JSON export -- full structured payload."""
    payload = {
        "investment_memo": ai_narrative_result,
        "timeline_events": timeline_events or [],
        "intelligence": intelligence_outputs or {},
    }
    return json.dumps(payload, indent=2, ensure_ascii=False).encode("utf-8")


def generate_markdown_export(ai_narrative_result, timeline_events, intelligence_outputs):
    """Phase 4: Markdown export."""
    md_lines = ["# Institutional Investment Research Memo", "", "## Investment Memo", "", ai_narrative_result or "", ""]
    if timeline_events:
        md_lines += ["## Timeline Events", ""]
        for event in timeline_events:
            md_lines.append(
                f"- **{event.get('date', 'N/A')}** -- {event.get('event', 'N/A')} "
                f"(_{event.get('category', 'N/A')}_, impact: {event.get('impact', 'N/A')})"
            )
        md_lines.append("")
    for title, value in _flatten_intelligence_for_text(intelligence_outputs):
        md_lines += [f"## {title}", ""]
        if isinstance(value, str):
            md_lines.append(value)
        elif isinstance(value, dict):
            for k, v in value.items():
                if isinstance(v, list):
                    v = ", ".join(str(x) for x in v)
                md_lines.append(f"- **{k}**: {v}")
        elif isinstance(value, list):
            for item in value:
                if isinstance(item, dict):
                    md_lines.append("- " + " | ".join(f"**{k}**: {v}" for k, v in item.items()))
                else:
                    md_lines.append(f"- {item}")
        md_lines.append("")
    return "\n".join(md_lines).encode("utf-8")


def generate_csv_export(timeline_events, intelligence_outputs):
    """Phase 4: CSV export -- stacks the timeline table and any
    list-of-dict intelligence outputs (e.g. risk_analysis) as separate
    labeled sections within one CSV file."""
    buffer = io.StringIO()
    if timeline_events:
        buffer.write("Timeline Events\n")
        pd.DataFrame(timeline_events).to_csv(buffer, index=False)
        buffer.write("\n")
    for key, value in (intelligence_outputs or {}).items():
        if isinstance(value, list) and value and all(isinstance(item, dict) for item in value):
            title = INTELLIGENCE_MODULES.get(key, {}).get("title", key)
            buffer.write(f"{title}\n")
            pd.DataFrame(value).to_csv(buffer, index=False)
            buffer.write("\n")
    return buffer.getvalue().encode("utf-8")


def generate_excel_export(timeline_events, intelligence_outputs, key_metrics=None):
    """Phase 4: Excel export -- one sheet per structured section.

    Bug fix: openpyxl raises `IndexError: At least one sheet must be
    visible` if the workbook ends up with zero worksheets written (e.g.
    when timeline_events, intelligence_outputs, and key_metrics are all
    empty). This guarantees at least one sheet always exists before the
    writer saves: if nothing else was written, a "Summary" sheet is added
    instead of leaving the workbook empty.
    """
    bio = io.BytesIO()
    with pd.ExcelWriter(bio, engine="openpyxl") as writer:
        wrote_any = False
        if timeline_events:
            pd.DataFrame(timeline_events).to_excel(writer, sheet_name="Timeline", index=False)
            wrote_any = True
        if key_metrics:
            pd.DataFrame(list(key_metrics.items()), columns=["Metric", "Value"]).to_excel(
                writer, sheet_name="Key Metrics", index=False
            )
            wrote_any = True
        for key, value in (intelligence_outputs or {}).items():
            if isinstance(value, list) and value and all(isinstance(item, dict) for item in value):
                raw_name = INTELLIGENCE_MODULES.get(key, {}).get("title", key)
                sheet_name = "".join(c for c in raw_name if c not in ['\\', '/', '*', '[', ']', ':', '?']).strip()
                sheet_name = (sheet_name or key)[:31]
                pd.DataFrame(value).to_excel(writer, sheet_name=sheet_name, index=False)
                wrote_any = True
        if not wrote_any:
            pd.DataFrame(["Financial Timeline Engine", "No exportable data available."]).to_excel(
                writer, sheet_name="Summary", index=False, header=False
            )
    bio.seek(0)
    return bio.getvalue()


# =============================================================================
# SECTION 8: Dashboard / Visualization
# =============================================================================
def render_timeline_visualization(timeline_data):
    """Renders a simplified timeline visualization."""
    if not timeline_data or len(timeline_data) == 0:
        st.info("No timeline events extracted yet.")
        return
    st.subheader("📊 Timeline Events")
    df_timeline = pd.DataFrame(timeline_data)
    st.dataframe(df_timeline, use_container_width=True, hide_index=True)


# =============================================================================
# SECTION 9: Future Modules (Phase 11 -- architecture only, not implemented)
# =============================================================================
def get_future_module_status():
    """Returns the roadmap of planned-but-not-yet-implemented modules
    (Valuation, DCF, Comparable Analysis, Forecasting, Earnings Model,
    Financial Modeling, Portfolio Analysis, Watchlist, Company Comparison,
    Stock Scoring, Screening Engine) for display in the UI. This is
    intentionally scaffolding only, per Phase 11's explicit scope."""
    return {key: {"name": name, "status": "planned"} for key, name in FUTURE_MODULES.items()}


# =============================================================================
# SECTION 10: Live Market Intelligence (Phase 6 -- architecture only)
# =============================================================================
def is_live_market_intelligence_enabled():
    return bool(get_secret(LIVE_INTELLIGENCE_API_KEY_NAME, ""))


def fetch_live_market_intelligence(query):
    """Phase 6 architecture stub: a real, working function that gracefully
    degrades when no live-data API key is configured, rather than
    crashing or fabricating results. No specific news/filings/market-data
    vendor is wired in yet (that requires choosing and paying for a
    provider) -- once one is chosen, its client call replaces the body of
    the `if enabled` branch below; the calling code in main() already
    handles both the enabled and disabled cases cleanly."""
    if not is_live_market_intelligence_enabled():
        return {
            "enabled": False,
            "data": None,
            "message": "Live market intelligence is not configured for this deployment.",
        }
    return {
        "enabled": True,
        "data": None,
        "message": "Live market intelligence key detected, but no data provider integration has been selected yet.",
    }
# =============================================================================
# SECTION 10.5: AI Financial Assistant (interactive chatbot)
# =============================================================================
# A conversational UI layer on top of the existing frozen intelligence
# stack (Agentic RAG + verified facts + CalculationSafetyGate + provider
# chain). It deliberately does NOT reimplement any intelligence -- it
# orchestrates backend.chat_assistant.FinancialChatAssistant, which in
# turn composes the existing components.

# Suggested prompts shown in the assistant's empty state.
CHAT_SUGGESTED_PROMPTS = [
    "What was the company's revenue?",
    "Compare revenue across fiscal years.",
    "Why did profitability change?",
    "What are the major financial risks?",
    "Summarize this company's financial position.",
]


@st.cache_resource(show_spinner=False)
def _build_chat_assistant():
    """Cached singleton for the chat assistant (lazy import keeps app
    startup light and keeps the module unit-testable outside Streamlit)."""
    from backend.chat_assistant import build_chat_assistant
    return build_chat_assistant()


def _chat_documents(extraction_results):
    """Shape ingestion results into the {file_name, financial_facts}
    structure the assistant's document-Q&A path expects."""
    documents = []
    for doc in extraction_results or []:
        documents.append({
            "file_name": doc.get("file_name", "Unknown Document"),
            "financial_facts": doc.get("financial_facts", []) or [],
        })
    return documents


def _render_chat_provenance(metadata):
    """Render concise provenance under a fact-backed assistant answer.

    - BLOCKED / miss intents show the reason instead of a badge.
    - Verified answers get a "✓ Verified" badge plus source / period /
      source-tier line and an expandable Evidence section.
    - Calculation answers get an expandable formula/inputs section.
    Internal hashes / technical objects are never shown.
    """
    intent = metadata.get("intent", "")
    evidence = metadata.get("evidence") or []
    calculation = metadata.get("calculation")

    blocked_intents = (
        "blocked", "calculation_miss", "document_miss", "company_blocked",
        "company_empty", "company_error", "no_provider", "provider_error",
    )
    if intent in blocked_intents:
        st.caption("🚫 **BLOCKED / NOT VERIFIED** — no fabricated figures.")
        reason = metadata.get("blocked_reason")
        if reason:
            st.caption(f"Reason: {reason}")
        return

    if evidence:
        top = evidence[0]
        source = top.get("source") or ""
        period = top.get("period") or ""
        tier = top.get("source_tier")
        document = top.get("document") or ""
        badge_parts = ["✅ **Verified**"]
        if source:
            badge_parts.append(f"Source: {source}")
        if period:
            badge_parts.append(f"Period: {period}")
        if tier is not None:
            badge_parts.append(f"Source tier: {tier}")
        if document:
            badge_parts.append(f"Document: {document}")
        st.caption(" · ".join(badge_parts))
        with st.expander("🔍 Evidence", expanded=False):
            evidence_rows = [
                {
                    "Metric": e.get("metric", ""),
                    "Value": e.get("display", e.get("value", "")),
                    "Period": e.get("period", ""),
                    "Scale": e.get("scale", ""),
                    "Source": e.get("source", ""),
                    "Tier": e.get("source_tier", ""),
                    "Document": e.get("document", ""),
                }
                for e in evidence
            ]
            st.dataframe(pd.DataFrame(evidence_rows), use_container_width=True, hide_index=True)

    if calculation:
        with st.expander("🧮 Calculation", expanded=False):
            st.markdown(f"**{calculation.get('name', 'Calculation')}**")
            if calculation.get("formula"):
                st.markdown(f"Formula: `{calculation['formula']}`")
            if calculation.get("value") is not None:
                st.markdown(f"Value: **{calculation['value']}**")
            if calculation.get("inputs"):
                st.markdown("Inputs: " + ", ".join(
                    f"{k}={v}" for k, v in calculation["inputs"].items()
                ))
            if calculation.get("periods"):
                st.markdown("Periods: " + ", ".join(calculation["periods"]))


def _submit_chat_question(question, extraction_results, provider_health):
    """Run one user turn through the assistant and persist messages +
    bounded conversation context in session state."""
    from backend.chat_assistant import ChatContext

    question = (question or "").strip()
    if not question:
        return

    ctx = ChatContext.from_state(st.session_state.get("chat_context_state"))
    assistant = _build_chat_assistant()
    documents = _chat_documents(extraction_results)

    ctx.add_user(question)
    st.session_state.setdefault("chat_messages", []).append(
        {"role": "user", "content": question}
    )

    with st.spinner("Consulting verified evidence…"):
        response = assistant.answer(
            question,
            context=ctx,
            documents=documents,
            provider_health=provider_health,
        )

    content = response.get("content", "")
    metadata = response.get("metadata", {}) or {}
    ctx.add_assistant(content, metadata)
    st.session_state["chat_context_state"] = ctx.to_state()
    st.session_state.setdefault("chat_messages", []).append(
        {"role": "assistant", "content": content, "metadata": metadata}
    )


def render_financial_assistant(extraction_results, provider_health):
    """AI Financial Assistant section -- st.chat_message / st.chat_input
    conversation with suggested prompts, persistent bounded history, and
    provenance rendering."""
    st.markdown("---")
    st.subheader("🤖 AI Financial Assistant")
    st.caption(
        "Ask questions about your financial documents, verified metrics, "
        "companies, and market data. Every figure is verified against "
        "evidence — the assistant never invents numbers."
    )

    # --- Session state ---
    if "chat_messages" not in st.session_state:
        st.session_state["chat_messages"] = []
    if "chat_context_state" not in st.session_state:
        st.session_state["chat_context_state"] = None

    messages = st.session_state["chat_messages"]

    # --- Empty state with suggested prompts ---
    if not messages:
        st.markdown("#### 💬 Start a conversation")
        cols = st.columns(2)
        for i, prompt in enumerate(CHAT_SUGGESTED_PROMPTS):
            with cols[i % 2]:
                if st.button(prompt, key=f"chat_suggest_{i}", use_container_width=True):
                    _submit_chat_question(prompt, extraction_results, provider_health)
                    st.rerun()
    else:
        # --- Conversation history ---
        for msg in messages:
            with st.chat_message(msg["role"]):
                st.markdown(msg["content"])
                if msg["role"] == "assistant" and msg.get("metadata"):
                    _render_chat_provenance(msg["metadata"])

    # --- Chat input ---
    user_input = st.chat_input(
        "Ask about your financial documents, verified metrics, or companies…"
    )
    if user_input:
        _submit_chat_question(user_input, extraction_results, provider_health)
        st.rerun()


# =============================================================================
# SECTION 11: Main App / UI
# =============================================================================
def main():
    st.title("📈 Financial Timeline Engine")

    # --- AI status (canonical provider health via backend/gateway) ---
    # Status states: 🔴 no eligible provider / 🟢 live generation verified /
    # 🟡 configured but no live call yet. Per-provider detail comes from the
    # canonical ProviderManager (available / configured_unavailable /
    # not_configured), never from a bare key presence check.
    health = get_provider_health()
    has_any_key = any(health.values())
    if not has_any_key:
        st.error(f"🔴 AI Status: Offline — {NO_ELIGIBLE_PROVIDER_MESSAGE}")
    elif st.session_state["ai_connected"]:
        provider = st.session_state.get("ai_provider_used", "AI Provider")
        st.success(f"🟢 AI Status: Live ({provider})")
    else:
        st.info("🟡 AI Status: Provider(s) configured — awaiting first live generation")

    with st.expander("🩺 Provider Health & Activity Log", expanded=False):
        provider_display = {
            "google": "Google AI Studio", "groq": "Groq", "openrouter": "OpenRouter",
            "nvidia": "NVIDIA", "rapidapi": "RapidAPI", "sambanova": "SambaNova",
            "github": "GitHub Models", "cerebras": "Cerebras", "cohere": "Cohere",
        }
        status_icon = {"available": "🟢", "configured_unavailable": "🟡", "not_configured": "⚪"}
        canonical_status = get_canonical_provider_status()
        health_rows = []
        for slug, state in canonical_status.items():
            health_rows.append({
                "Provider": provider_display.get(slug, slug.replace("_", " ").title()),
                "Status": f"{status_icon.get(state, '⚪')} {state.replace('_', ' ')}",
            })
        # Fall back to key-presence rows only if the gateway is unavailable.
        if not health_rows:
            health_rows = [
                {"Provider": k, "Status": ("🟢 available" if v else "⚪ not configured")}
                for k, v in health.items()
            ]
        health_df = pd.DataFrame(health_rows)
        st.dataframe(health_df, use_container_width=True, hide_index=True)
        if st.session_state["provider_log"]:
            st.dataframe(pd.DataFrame(st.session_state["provider_log"][-20:]), use_container_width=True, hide_index=True)
        else:
            st.caption("No AI provider activity yet.")

    with st.expander("🔮 Roadmap (Planned Modules)", expanded=False):
        roadmap = get_future_module_status()
        st.dataframe(
            pd.DataFrame([{"Module": v["name"], "Status": v["status"].title()} for v in roadmap.values()]),
            use_container_width=True, hide_index=True
        )

    if is_live_market_intelligence_enabled():
        st.caption("🌐 Live Market Intelligence: key detected (provider integration pending).")

    # --- Sidebar Document Ingestion ---
    st.sidebar.header("📁 Document Ingestion")
    uploaded_files = st.sidebar.file_uploader(
        "Upload Financial Documents (.txt, .csv, .xlsx, .docx, .pdf)",
        type=["txt", "csv", "xlsx", "docx", "pdf"],
        accept_multiple_files=True
    )

    combined_raw_text = ""
    document_summaries = []
    extraction_results = []

    if uploaded_files:
        # Single ingestion pipeline: parsing, caching, chunking, and
        # statistics all happen inside ingestion.extract_multiple(), which
        # returns one result dict per file (each with "file_name",
        # "parsed_document", "chunks", "statistics", "cached").
        extraction_results = extract_multiple(uploaded_files)

        # merge_document_text() combines every file's raw extracted text
        # into one delimited blob, replacing the old manual
        # "combined_raw_text +=" loop.
        combined_raw_text = merge_document_text(extraction_results)

        for doc in extraction_results:
            file_name = doc["file_name"]
            extracted_text = doc["parsed_document"]["text"]

            # Phase 8: hash-based cache -- re-uploading the same file
            # content doesn't re-trigger an AI summarization call.
            cache_key = f"{file_name}:{_hash_text(extracted_text)}"
            summary_text, _ = _cached_call(
                "summary_cache", cache_key,
                lambda et=extracted_text, fn=file_name: summarize_document_with_chunking(et, fn)
            )
            document_summaries.append({"file_name": file_name, "summary": summary_text})

    st.subheader("📊 Ingested Data Summary")
    col1, col2 = st.columns(2)
    col1.metric(label="📄 Files Processed", value=len(uploaded_files) if uploaded_files else 0)
    col2.metric(label="📊 Extracted Characters", value=len(combined_raw_text))

    if extraction_results:
        with st.expander("🧾 Ingestion Statistics (pages, tables, chunks, tokens)", expanded=False):
            st.text(print_statistics(document_statistics(extraction_results)))

    # --- AI Financial Assistant (interactive chatbot) ---
    # Runs after ingestion so the assistant can answer questions against
    # the uploaded documents' extracted financial facts. Passes provider
    # health so it degrades gracefully when no AI key is configured.
    render_financial_assistant(extraction_results, health)

    st.markdown("---")
    st.subheader("🔬 AI Analysis Engine")

    if st.button("🚀 Generate Timeline Report"):
        if not uploaded_files:
            st.warning("Please upload at least one financial document before generating a report.")
        else:
            # Reset the pipeline debug trace for this run so it only shows
            # entries from the current click, not accumulated across every
            # click in the session.
            st.session_state["pipeline_debug_log"] = []

            with st.spinner("Merging document summaries..."):
                merge_cache_key = _hash_text(
                    "||".join(f"{d['file_name']}:{d['summary']}" for d in document_summaries)
                )
                master_summary, _ = _cached_call(
                    "merge_cache", merge_cache_key,
                    lambda: merge_document_summaries(document_summaries)
                )
                _debug_stage("master_summary (final, pre-memo)", master_summary)

            # Bug fix: if summarization/merging failed for every uploaded
            # file/chunk, master_summary itself is an error string. Detect
            # that HERE (in addition to the existing ai_narrative_result
            # check below) so the wasted downstream memo/timeline/
            # intelligence AI calls are skipped and the user gets a clear,
            # specific reason instead of 17 sections each independently
            # reporting "insufficient information".
            master_summary_failed = (
                not master_summary or not master_summary.strip() or contains_error_marker(master_summary)
            )
            if master_summary_failed:
                st.error(
                    "❌ Document summarization failed for every uploaded file/chunk -- "
                    "there is no real content to analyze. Check the Provider Health & "
                    "Activity Log above for the underlying error, then try again."
                )

            with st.spinner("Generating investment memo..."):
                memo_system_prompt = (
                    "You are an elite institutional investment research analyst. "
                    "Write the investment memo strictly and exclusively from the "
                    "facts, figures, dates, and events contained in the Document "
                    "Summary provided below. Do not produce generic, templated, or "
                    "boilerplate analysis -- every claim must be traceable to a "
                    "specific fact stated in the Document Summary. If the Document "
                    "Summary lacks information for a requested section, explicitly "
                    "state that the source documents did not provide it rather than "
                    "inventing generic filler."
                )
                prompt = f"""Analyze the Document Summary below carefully. Extract key event milestones, timelines, and potential controversy flags SPECIFIC to this Document Summary. Write a comprehensive multi-paragraph investment memo that identifies, using only facts from the Document Summary:
1. Key financial events and dates
2. Market movements and impacts
3. Risk factors and opportunities
4. Strategic implications

Document Summary:
{master_summary}

Generate a professional investment memo grounded strictly in the Document Summary above. Do not generate generic industry commentary that is not tied to a specific fact in the Document Summary."""

                memo_cache_key = _hash_text(master_summary)
                ai_narrative_result, _ = _cached_call(
                    "memo_cache", memo_cache_key,
                    lambda: call_ai_with_fallback(prompt, system_prompt=memo_system_prompt, temperature=0.3)
                )

            st.markdown("### 📝 Generated Investment Memo")
            st.write(ai_narrative_result)

            is_error = (
                master_summary_failed
                or ("❌" in ai_narrative_result) or ("🔴" in ai_narrative_result) or ("⚠️" in ai_narrative_result)
            )

            timeline_events = []
            intelligence_outputs = {}
            if not is_error:
                with st.spinner("Extracting timeline events..."):
                    timeline_cache_key = _hash_text(ai_narrative_result)
                    timeline_events, _ = _cached_call(
                        "timeline_cache", timeline_cache_key,
                        lambda: extract_timeline_events(ai_narrative_result)
                    )
                    st.session_state["timeline_data"] = timeline_events

                if timeline_events:
                    render_timeline_visualization(timeline_events)

                # --- Module 3: deterministic financial intelligence ---
                # Runs on the raw extracted text/documents from Module 2
                # (not on the AI-generated master_summary) so its facts and
                # ratios are grounded in the actual source figures rather
                # than a paraphrase of them. Summarization and the
                # investment memo above are unaffected -- Module 3 only
                # changes what Universal Intelligence Extraction receives,
                # below.
                with st.spinner("Running Module 3 financial intelligence pipeline..."):
                    module3_cache_key = _hash_text(combined_raw_text)
                    module3_result, _ = _cached_call(
                        "module3_cache", module3_cache_key,
                        lambda: run_module3(combined_raw_text, extraction_results)
                    )

                st.markdown("### 🧮 Module 3 — Financial Intelligence")
                with st.expander("📊 Financial Data", expanded=False):
                    _render_module3_value(module3_result.get("financial_data"))
                with st.expander("📐 Ratios", expanded=False):
                    _render_module3_value(module3_result.get("ratios"))
                with st.expander("🔍 OCR Verification", expanded=False):
                    _render_module3_value(module3_result.get("ocr_verification"))
                with st.expander("🧾 Cross Document Verification", expanded=False):
                    _render_module3_value(module3_result.get("cross_document_verification"))
                with st.expander("🎯 Confidence Scores", expanded=False):
                    _render_module3_value(module3_result.get("confidence"))
                with st.expander("📅 Events", expanded=False):
                    _render_module3_value(module3_result.get("events"))
                with st.expander("🕒 Module 3 Timeline", expanded=False):
                    _render_module3_value(module3_result.get("timeline"))
                with st.expander("🗜️ Optimized Context (sent to AI)", expanded=False):
                    _render_module3_value(module3_result.get("optimized_context"))

                # The AI receives ONLY the compressed optimized_context --
                # never the raw extracted document text -- which is the
                # entire point of Module 3's token-reduction pass.
                optimized_context = module3_result.get("optimized_context") or {}
                optimized_context_text = json.dumps(optimized_context, indent=2, ensure_ascii=False, default=str)

                with st.spinner("Running institutional intelligence extraction..."):
                    intelligence_cache_key = _hash_text(optimized_context_text)
                    intelligence_outputs, _ = _cached_call(
                        "intelligence_cache", intelligence_cache_key,
                        lambda: run_universal_intelligence_extraction(optimized_context_text)
                    )
                    st.session_state["intelligence_outputs"] = intelligence_outputs
                    st.session_state["key_metrics"] = intelligence_outputs.get("key_metrics", {})
                    st.session_state["sector_analysis"] = intelligence_outputs.get("sector_analysis", {})
                    st.session_state["risk_analysis"] = intelligence_outputs.get("risk_analysis", [])
                    st.session_state["controversy_analysis"] = intelligence_outputs.get("controversy_analysis", [])

                for module_key in INTELLIGENCE_MODULES:
                    render_intelligence_output(module_key, intelligence_outputs.get(module_key))

                docx_file_stream = generate_docx_download(ai_narrative_result, timeline_events, intelligence_outputs)
                pdf_file_stream = generate_pdf_download(ai_narrative_result, timeline_events, intelligence_outputs)

                export_col1, export_col2 = st.columns(2)
                with export_col1:
                    st.download_button(
                        label="📥 Download as Word Document",
                        data=docx_file_stream,
                        file_name="Financial_Timeline_Investment_Memo.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
                with export_col2:
                    st.download_button(
                        label="📄 Download PDF",
                        data=pdf_file_stream,
                        file_name="Financial_Timeline_Investment_Memo.pdf",
                        mime="application/pdf",
                        use_container_width=True
                    )

                with st.expander("📤 More Export Formats (JSON / CSV / Excel / Markdown)", expanded=False):
                    json_bytes = generate_json_export(ai_narrative_result, timeline_events, intelligence_outputs)
                    md_bytes = generate_markdown_export(ai_narrative_result, timeline_events, intelligence_outputs)
                    csv_bytes = generate_csv_export(timeline_events, intelligence_outputs)
                    excel_bytes = generate_excel_export(
                        timeline_events, intelligence_outputs, intelligence_outputs.get("key_metrics", {})
                    )

                    more_col1, more_col2 = st.columns(2)
                    with more_col1:
                        st.download_button(
                            "🧾 Download JSON", data=json_bytes,
                            file_name="Financial_Timeline_Report.json", mime="application/json",
                            use_container_width=True
                        )
                        st.download_button(
                            "📊 Download CSV", data=csv_bytes,
                            file_name="Financial_Timeline_Report.csv", mime="text/csv",
                            use_container_width=True
                        )
                    with more_col2:
                        st.download_button(
                            "📗 Download Excel", data=excel_bytes,
                            file_name="Financial_Timeline_Report.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            use_container_width=True
                        )
                        st.download_button(
                            "📃 Download Markdown", data=md_bytes,
                            file_name="Financial_Timeline_Report.md", mime="text/markdown",
                            use_container_width=True
                        )
            else:
                st.warning("AI generation encountered an error. Please review the message above and try again.")

            with st.expander("🔍 Pipeline Debug Trace", expanded=False):
                debug_entries = st.session_state.get("pipeline_debug_log", [])
                if debug_entries:
                    st.dataframe(
                        pd.DataFrame([
                            {
                                "Stage": e["stage"],
                                "Length": e["length"],
                                "Error Marker?": "⚠️" if e["error_marker_detected"] else "",
                                "Preview (first 1000 chars)": e["preview"],
                            }
                            for e in debug_entries
                        ]),
                        use_container_width=True, hide_index=True
                    )
                else:
                    st.caption("No pipeline stages recorded for this run.")


# =============================================================================
# SECTION 12: Auth
# =============================================================================
def check_login():
    if "authenticated" not in st.session_state:
        st.session_state["authenticated"] = False
    if not st.session_state["authenticated"]:
        st.markdown("🔐 Institutional Terminal Access")
        col_l1, col_l2, col_l3 = st.columns([1, 2, 1])
        with col_l2:
            input_user = st.text_input("Username")
            input_pass = st.text_input("Password", type="password")
            if st.button("🚀 Log In", use_container_width=True):
                if input_user == "admin" and input_pass == "financial_terminal_2026":
                    st.session_state["authenticated"] = True
                    st.rerun()
                else:
                    st.error("❌ Invalid Credentials")
        return False
    return True


if __name__ == "__main__":
    if check_login():
        main()
