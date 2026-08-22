"""
Platrixa
Module 2 - Universal Document Parser

Supported:
- PDF
- DOCX
- XLSX
- CSV
- TXT
- HTML (including SEC Inline XBRL)

High fidelity extraction while preserving ordering,
tables (structured `table_data`), page sequence and metadata.

API compatibility:
  parse_document(uploaded_file) -> {
      "type", "pages", "tables" (int count), "text",
      "table_data" (list of structured table dicts),   # NEW (additive)
      "xbrl_facts" (list of XBRL fact dicts),           # NEW (html/xbrl only)
  }
The existing "type"/"pages"/"tables"/"text" keys are unchanged.
"""

from __future__ import annotations

import io
import os
import pandas as pd
from docx import Document
from pypdf import PdfReader


# ============================================================
# Supported Extensions
# ============================================================

SUPPORTED_TYPES = {
    ".pdf",
    ".docx",
    ".xlsx",
    ".csv",
    ".txt",
    ".html",
    ".htm",
}


# ============================================================
# File Type Detection
# ============================================================

def detect_file_type(filename: str) -> str:
    ext = os.path.splitext(filename.lower())[1]

    if ext not in SUPPORTED_TYPES:
        raise ValueError(f"Unsupported file type: {ext}")

    return ext


# ============================================================
# Table rendering helper (shared with ingestion.chunking)
# ============================================================

def render_table_block(table: dict, table_id: str, location: str = "") -> str:
    """
    Render a structured table dict as an atomic, marker-delimited text
    block. Used so table content is preserved in the document text and
    chunking keeps the whole table together.

    Format:
        === TABLE <id> (<location>) ===
        header1 | header2
        rowlabel | cell1 | cell2
        === END TABLE ===
    """
    lines = [f"=== TABLE {table_id} ({location}) ==="]
    headers = table.get("headers") or []
    if headers:
        lines.append(" | ".join(str(h) for h in headers))
    for row in table.get("rows") or []:
        if isinstance(row, dict):
            label = str(row.get("label", ""))
            cells = [str(c) for c in (row.get("cells") or [])]
        else:
            cells = [str(c) for c in row]
            label = cells[0] if cells else ""
            cells = cells[1:]
        lines.append(" | ".join([label] + cells))
    lines.append("=== END TABLE ===")
    return "\n".join(lines)


# ============================================================
# TXT
# ============================================================

def parse_txt(uploaded_file) -> dict:

    uploaded_file.seek(0)

    text = uploaded_file.read().decode(
        "utf-8",
        errors="ignore"
    )

    return {
        "type": "txt",
        "pages": 1,
        "tables": 0,
        "table_data": [],
        "xbrl_facts": [],
        "text": text
    }


# ============================================================
# CSV
# ============================================================

def parse_csv(uploaded_file) -> dict:

    uploaded_file.seek(0)

    df = pd.read_csv(uploaded_file)

    table_data = [_df_to_table(
        df,
        table_id="csv_table_1",
        location="csv",
    )]

    return {
        "type": "csv",
        "pages": 1,
        "tables": 1,
        "table_data": table_data,
        "xbrl_facts": [],
        "text": df.to_string(index=False)
    }


# ============================================================
# Excel
# ============================================================

def parse_excel(uploaded_file) -> dict:

    uploaded_file.seek(0)

    sheets = pd.read_excel(
        uploaded_file,
        sheet_name=None
    )

    output = []

    table_data = []

    table_count = 0

    for name, df in sheets.items():

        table_count += 1

        output.append(
            f"\n========== SHEET : {name} ==========\n"
        )

        output.append(
            df.to_string(index=False)
        )

        table_data.append(
            _df_to_table(
                df,
                table_id=f"sheet_{name}",
                location=f"xlsx sheet {name}",
            )
        )

    return {
        "type": "xlsx",
        "pages": len(sheets),
        "tables": table_count,
        "table_data": table_data,
        "xbrl_facts": [],
        "text": "\n".join(output)
    }


def _df_to_table(df, table_id: str, location: str) -> dict:
    """Convert a pandas DataFrame into a structured table dict."""
    headers = [str(c) for c in df.columns] if df.columns is not None else []
    rows = []
    for record in df.itertuples(index=False):
        cells = ["" if pd.isna(v) else str(v) for v in record]
        if not cells:
            continue
        rows.append({"label": cells[0], "cells": cells[1:]})
    return {
        "table_id": table_id,
        "headers": headers,
        "rows": rows,
        "source_location": location,
    }


# ============================================================
# DOCX
# ============================================================

def parse_docx(uploaded_file) -> dict:

    uploaded_file.seek(0)

    document = Document(uploaded_file)

    paragraphs = []

    for para in document.paragraphs:

        txt = para.text.strip()

        if txt:

            paragraphs.append(txt)

    table_data = []

    text_parts = list(paragraphs)

    for i, tbl in enumerate(document.tables):

        rows = []

        for row in tbl.rows:

            rows.append([
                cell.text.strip()
                for cell in row.cells
            ])

        if not rows:
            continue

        table_id = f"docx_table_{i + 1}"

        table_data.append({
            "table_id": table_id,
            "headers": rows[0],
            "rows": [
                {"label": r[0], "cells": r[1:]}
                for r in rows[1:]
            ],
            "source_location": f"docx table #{i + 1}",
        })

        # Preserve table content in the text stream (marked, atomic)
        text_parts.append(
            render_table_block(
                table_data[-1],
                table_id=table_id,
                location=f"docx table #{i + 1}",
            )
        )

    return {
        "type": "docx",
        "pages": 1,
        "tables": len(document.tables),
        "table_data": table_data,
        "xbrl_facts": [],
        "text": "\n".join(text_parts)
    }


# ============================================================
# PDF
# ============================================================

def parse_pdf(uploaded_file) -> dict:

    uploaded_file.seek(0)

    reader = PdfReader(uploaded_file)

    pages = []

    page_count = 0

    for i, page in enumerate(reader.pages):

        page_count += 1

        text = page.extract_text()

        if text:

            pages.append(
                f"\n========== PAGE {i+1} ==========\n"
            )

            pages.append(text)

    joined = "\n".join(pages)

    table_data = _extract_pdf_tables(joined)

    return {
        "type": "pdf",
        "pages": page_count,
        "tables": len(table_data),
        "table_data": table_data,
        "xbrl_facts": [],
        "text": joined
    }


def _extract_pdf_tables(text: str) -> list:
    """Layout-aware table detection from PDF page text (best effort)."""
    try:
        from backend.extraction2.table_extractor import TableExtractor
        tables = TableExtractor().extract_text_tables(text, source="pdf")
        return [t.to_dict() for t in tables]
    except Exception:
        return []


# ============================================================
# HTML (including SEC Inline XBRL)
# ============================================================

def parse_html(uploaded_file) -> dict:

    uploaded_file.seek(0)

    raw = uploaded_file.read().decode(
        "utf-8",
        errors="replace"
    )

    # ---- text extraction (BeautifulSoup) ----
    text = ""
    table_data = []
    xbrl_facts = []

    try:
        from bs4 import BeautifulSoup

        soup = BeautifulSoup(raw, "lxml")

        text = soup.get_text(" ", strip=True)

        from backend.extraction2.table_extractor import TableExtractor

        tables = TableExtractor().extract_html_tables(raw)

        table_data = [t.to_dict() for t in tables]

    except Exception:
        # Last resort: crude tag strip
        import re
        text = re.sub(r"<[^>]+>", " ", raw)
        text = re.sub(r"\s+", " ", text)

    # ---- Inline XBRL / XBRL structured facts ----
    try:
        from backend.extraction2.xbrl_extractor import XbrlExtractor
        xbrl_facts = [f.to_dict() for f in XbrlExtractor().extract(raw)]
    except Exception:
        xbrl_facts = []

    return {
        "type": "html",
        "pages": 1,
        "tables": len(table_data),
        "table_data": table_data,
        "xbrl_facts": xbrl_facts,
        "text": text
    }


# ============================================================
# Universal Parser
# ============================================================

def parse_document(uploaded_file):

    file_type = detect_file_type(uploaded_file.name)

    if file_type == ".pdf":
        return parse_pdf(uploaded_file)

    if file_type == ".docx":
        return parse_docx(uploaded_file)

    if file_type == ".xlsx":
        return parse_excel(uploaded_file)

    if file_type == ".csv":
        return parse_csv(uploaded_file)

    if file_type == ".txt":
        return parse_txt(uploaded_file)

    if file_type in (".html", ".htm"):
        return parse_html(uploaded_file)

    raise RuntimeError("Parser not found")
